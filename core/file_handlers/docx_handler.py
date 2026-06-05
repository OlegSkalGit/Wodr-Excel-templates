import os
import re
from docx import Document
from docxtpl import DocxTemplate, Listing
from lxml import etree
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def is_hf_defined(section, attr):
    ref_type = 'first' if 'first_page' in attr else ('even' if 'even_page' in attr else 'default')
    tag = 'headerReference' if 'header' in attr else 'footerReference'
    
    for ref in section._sectPr.findall(qn(f'w:{tag}')):
        w_type = ref.get(qn('w:type'))
        if w_type is None:
            w_type = 'default'
        if w_type == ref_type:
            return True
    return False


def get_run_props(r):
    if r._r.rPr is not None:
        return etree.tostring(r._r.rPr).decode('utf-8')
    return None

def apply_run_props(r, props):
    if props is not None:
        try:
            rPr = parse_xml(props)
            r._r.insert(0, rPr)
        except Exception:
            pass

def get_formatting_chunks(p):
    char_to_props = []
    for r in p.runs:
        props = get_run_props(r)
        for _ in range(len(r.text)):
            char_to_props.append(props)
    chunks = []
    if not p.text: return chunks
    cur_text = p.text[0]
    cur_props = char_to_props[0]
    for i in range(1, len(p.text)):
        if char_to_props[i] == cur_props:
            cur_text += p.text[i]
        else:
            chunks.append((cur_text, cur_props))
            cur_text = p.text[i]
            cur_props = char_to_props[i]
    chunks.append((cur_text, cur_props))
    return chunks

def consolidate_paragraph_tags(p):
    text = p.text
    if '{{' not in text or '}}' not in text: return
    matches = list(re.finditer(r'\{\{.*?\}\}', text))
    if not matches: return
    char_to_run = []
    for r_idx, r in enumerate(p.runs):
        for _ in range(len(r.text)):
            char_to_run.append((r_idx, r))
    if len(char_to_run) != len(text): return
    for m in matches:
        start_idx, end_idx = m.start(), m.end()
        tag_text = m.group(0)
        span_runs = [char_to_run[idx] for idx in range(start_idx, end_idx)]
        unique_run_indices = []
        unique_runs = []
        for r_idx, r in span_runs:
            if r_idx not in unique_run_indices:
                unique_run_indices.append(r_idx)
                unique_runs.append(r)
        if len(unique_runs) > 1:
            first_run_idx = unique_run_indices[0]
            last_run_idx = unique_run_indices[-1]
            first_run = unique_runs[0]
            last_run = unique_runs[-1]
            
            first_run_start = sum(len(p.runs[i].text) for i in range(first_run_idx))
            last_run_start = sum(len(p.runs[i].text) for i in range(last_run_idx))
            
            prefix = first_run.text[:start_idx - first_run_start]
            suffix = last_run.text[end_idx - last_run_start:]
            
            first_run.text = prefix + tag_text
            for other_idx in range(first_run_idx + 1, last_run_idx):
                p.runs[other_idx].text = ""
            last_run.text = suffix
            char_to_run = []
            for r_idx, r in enumerate(p.runs):
                for _ in range(len(r.text)):
                    char_to_run.append((r_idx, r))
            if len(char_to_run) != len(text): break

def consolidate_jinja_tags(doc):
    try:
        for p in doc.paragraphs:
            if p is not None: consolidate_paragraph_tags(p)
    except: pass
    try:
        for table in doc.tables:
            if table is None: continue
            for row in table.rows:
                if row is None: continue
                for cell in row.cells:
                    if cell is None: continue
                    for p in cell.paragraphs:
                        if p is not None: consolidate_paragraph_tags(p)
    except: pass
    try:
        for section in doc.sections:
            if section is None: continue
            for attr in ['header', 'footer', 'first_page_header', 'first_page_footer', 'even_page_header', 'even_page_footer']:
                if not is_hf_defined(section, attr): continue
                h_f = getattr(section, attr, None)
                if h_f is not None:
                    try:
                        for p in h_f.paragraphs:
                            if p is not None: consolidate_paragraph_tags(p)
                    except: pass
                    try:
                        for table in h_f.tables:
                            if table is None: continue
                            for row in table.rows:
                                if row is None: continue
                                for cell in row.cells:
                                    if cell is None: continue
                                    for p in cell.paragraphs:
                                        if p is not None: consolidate_paragraph_tags(p)
                    except: pass
    except: pass

def process_word(template_path, output_path, variables):
    try:
        doc = DocxTemplate(template_path)
        consolidate_jinja_tags(doc)
        processed_vars = {}
        for key, value in variables.items():
            if isinstance(value, str) and '\n' in value: processed_vars[key] = Listing(value)
            else: processed_vars[key] = value
        doc.render(processed_vars)
        doc.save(output_path)
        return True, None
    except Exception as e:
        err = f"[Помилка Word]: {e}"
        print(f"  {err}")
        return False, err
