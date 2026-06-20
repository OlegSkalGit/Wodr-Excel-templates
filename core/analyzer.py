import os
import re
import difflib
import openpyxl
from datetime import datetime
from openpyxl.styles import Alignment
from docx import Document

from core.text_processor import get_norm_key, tokenize, get_unique_path, render_string_template
from core.file_handlers.docx_handler import get_formatting_chunks, get_run_props, apply_run_props
from core.excel_styles import check_cell_style_diffs, serialize_color, serialize_fill, serialize_alignment, serialize_border

def find_all_vars_in_slot(strings, diff_map, v_idx, char_to_props=None):
    s0 = str(strings[0] or "")
    tok0 = tokenize(s0)
    
    tok_char_idx = []
    curr_idx = 0
    for tok in tok0:
        tok_char_idx.append(curr_idx)
        curr_idx += len(tok)
    tok_char_idx.append(curr_idx)

    is_var = [False] * len(tok0)
    for si in strings[1:]:
        if s0 == si: continue
        toki = tokenize(str(si or ""))
        sm = difflib.SequenceMatcher(None, tok0, toki)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag != 'equal':
                diff_0 = "".join(tok0[i1:i2])
                diff_i = "".join(toki[j1:j2])
                if not diff_0.strip() and not diff_i.strip():
                    continue
                if tag == 'insert':
                    if i1 < len(tok0): is_var[i1] = True
                    elif i1 > 0: is_var[i1 - 1] = True
                for k in range(i1, i2): is_var[k] = True
            elif tag == 'equal':
                equal_str = "".join(tok0[i1:i2])
                if len(equal_str) <= 3 and not equal_str.strip():
                    for k in range(i1, i2): is_var[k] = True

    # Pass 2: Bridge short gaps between True values
    i = 0
    while i < len(tok0):
        if not is_var[i]:
            start_false = i
            while i < len(tok0) and not is_var[i]:
                i += 1
            end_false = i
            
            if start_false > 0 and end_false < len(tok0):
                gap_str = "".join(tok0[start_false:end_false])
                if len(gap_str) <= 4:
                    for k in range(start_false, end_false):
                        is_var[k] = True
        else:
            i += 1
            
    var_ranges = []
    def extract_key_formatting(prop_xml):
        if not prop_xml: return ""
        keys = []
        for tag in ['w:b', 'w:i', 'w:u', 'w:strike', 'w:sz', 'w:color', 'w:highlight', 'w:vertAlign']:
            m = re.search(r'<' + tag + r'(?: [^>]*?)?/>', prop_xml)
            if m: keys.append(m.group(0))
            m_cs = re.search(r'<' + tag + r'Cs(?: [^>]*?)?/>', prop_xml)
            if m_cs: keys.append(m_cs.group(0))
        return "".join(keys)

    i = 0
    while i < len(tok0):
        if is_var[i]:
            start = i
            i += 1
            if char_to_props is not None:
                start_char = tok_char_idx[start]
                start_fmt = extract_key_formatting(char_to_props[start_char])
                while i < len(tok0) and is_var[i]:
                    curr_char = tok_char_idx[i]
                    next_char = tok_char_idx[i+1] if i+1 < len(tok_char_idx) else len(char_to_props)
                    slice_props = char_to_props[curr_char:next_char]
                    if slice_props:
                        curr_fmts = {extract_key_formatting(p) for p in slice_props}
                        if len(curr_fmts) > 1 or list(curr_fmts)[0] != start_fmt:
                            break
                    i += 1
            else:
                while i < len(tok0) and is_var[i]: i += 1
            
            while start < i and not tok0[start].strip():
                is_var[start] = False
                start += 1
            while i > start and not tok0[i-1].strip():
                i -= 1
                is_var[i] = False
                
            if start >= i:
                continue
                
            var_text = "".join(tok0[start:i])
            if len(var_text) <= 2 and not any(c.isalnum() for c in var_text):
                for k in range(start, i):
                    is_var[k] = False
                continue
            
            var_ranges.append((start, i))
        else: i += 1
        
    valid_var_ranges = []
    template_parts = []
    last_idx = 0
    all_values_list = [[] for _ in range(len(strings))]
    consumed_j_list = [set() for _ in range(len(strings))]
    
    for v_start, v_end in var_ranges:
        template_parts.append("".join(tok0[last_idx:v_start]))
        vals = []
        for f_idx, si in enumerate(strings):
            if f_idx == 0: val = "".join(tok0[v_start:v_end])
            else:
                toki = tokenize(str(si or ""))
                sm = difflib.SequenceMatcher(None, tok0, toki)
                val_bits = []
                for tag, i1, i2, j1, j2 in sm.get_opcodes():
                    if tag == 'insert':
                        if v_start <= i1 <= v_end:
                            unconsumed = [toki[j] for j in range(j1, j2) if j not in consumed_j_list[f_idx]]
                            val_bits.append("".join(unconsumed))
                            consumed_j_list[f_idx].update(range(j1, j2))
                    else:
                        o_start = max(i1, v_start)
                        o_end = min(i2, v_end)
                        if o_start < o_end:
                            if tag == 'equal': val_bits.append("".join(tok0[o_start:o_end]))
                            else: 
                                unconsumed = [toki[j] for j in range(j1, j2) if j not in consumed_j_list[f_idx]]
                                val_bits.append("".join(unconsumed))
                                consumed_j_list[f_idx].update(range(j1, j2))
                val = "".join(val_bits)
            vals.append(val)
            
        cleaned_vals = [v.strip() for v in vals]
        if all(not v for v in cleaned_vals):
            template_parts.append("".join(tok0[v_start:v_end]))
        else:
            template_parts.append(f"{{{{VAR_PLACEHOLDER_{len(valid_var_ranges)}}}}}")
            valid_var_ranges.append((v_start, v_end))
            for f_idx in range(len(strings)):
                all_values_list[f_idx].append(cleaned_vals[f_idx])
                
        last_idx = v_end
        
    template_parts.append("".join(tok0[last_idx:]))
    
    all_values_dict = [{} for _ in range(len(strings))]
    for f_idx in range(len(strings)):
        for idx, val in enumerate(all_values_list[f_idx]):
            all_values_dict[f_idx][idx] = val
            
    return "".join(template_parts), valid_var_ranges, all_values_dict

def get_var_name(vals, diff_map, v_idx_list):
    norm_vals = tuple(get_norm_key(v) for v in vals)
    if norm_vals in diff_map: return diff_map[norm_vals]
    v_name = f"field_{v_idx_list[0]}"
    diff_map[norm_vals] = v_name
    v_idx_list[0] += 1
    return v_name

def align_blocks(base_items, comp_items, get_text_fn):
    def norm(txt): return re.sub(r'\s+', '', txt) if txt else ""
    base_sigs = [norm(get_text_fn(b)) for b in base_items]
    comp_sigs = [norm(get_text_fn(b)) for b in comp_items]
    sm = difflib.SequenceMatcher(None, base_sigs, comp_sigs)
    alignment = {}
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            for k in range(i2 - i1): alignment[i1 + k] = [j1 + k]
        elif tag == 'replace':
            limit = min(i2 - i1, j2 - j1)
            for k in range(limit):
                alignment[i1 + k] = [j1 + k]
    return alignment

def process_paragraph_list(blocks_list_per_file, diff_map, v_idx_list, all_data):
    num_files = len(blocks_list_per_file)
    if num_files == 0: return
    
    base_blocks = blocks_list_per_file[0]
    alignments = [{i: [i] for i in range(len(base_blocks))}]
    for f_idx in range(1, num_files):
        alignments.append(align_blocks(base_blocks, blocks_list_per_file[f_idx], lambda b: b.text if b else ""))

    def get_combined_chunks(p_items):
        if not p_items: return []
        chunks = []
        for i, p in enumerate(p_items):
            if i > 0:
                prev_props = chunks[-1][1] if chunks else None
                chunks.append(("\n", prev_props))
            chunks.extend(get_formatting_chunks(p))
        final_chunks = []
        if chunks:
            cur_text, cur_props = chunks[0]
            for t, p in chunks[1:]:
                if p == cur_props:
                    cur_text += t
                else:
                    final_chunks.append((cur_text, cur_props))
                    cur_text = t
                    cur_props = p
            final_chunks.append((cur_text, cur_props))
        return final_chunks

    for p_idx in range(len(base_blocks)):
        p_list = []
        for f_idx in range(num_files):
            a_map = alignments[f_idx]
            comp_indices = a_map.get(p_idx, [])
            comp_blocks = blocks_list_per_file[f_idx]
            
            valid_blocks = [comp_blocks[i] for i in comp_indices if i < len(comp_blocks)]
            p_list.append(valid_blocks if valid_blocks else None)
            
        all_texts = ["\n".join(x.text for x in p_items) if p_items else "" for p_items in p_list]
        if len(set(all_texts)) == 1: continue

        p0 = p_list[0][0]
        for extra_p in p_list[0][1:]: extra_p.clear()
        char_to_props = []
        for i, p in enumerate(p_list[0]):
            if i > 0:
                char_to_props.append(char_to_props[-1] if char_to_props else None)
            for r in p.runs:
                props = get_run_props(r)
                for _ in range(len(r.text)):
                    char_to_props.append(props)

        t_str, var_rs, vals_f = find_all_vars_in_slot(all_texts, diff_map, v_idx_list, char_to_props=char_to_props)
        if not var_rs: continue

        tok0 = tokenize(all_texts[0])
        tok_char_idx = []
        cur_len = 0
        for t in tok0:
            tok_char_idx.append(cur_len)
            cur_len += len(t)

        p0.clear()
        
        var_map = {}
        if var_rs:
            for idx, (vs, ve) in enumerate(var_rs):
                v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx_list)
                for f_idx in range(num_files):
                    all_data[f_idx][v_name] = vals_f[f_idx][idx]
                var_map[vs] = (ve, v_name)
                
        i = 0
        while i < len(tok0):
            if i in var_map:
                ve, v_name = var_map[i]
                run = p0.add_run(f"{{{{{v_name}}}}}")
                start_char = tok_char_idx[i]
                apply_run_props(run, char_to_props[start_char] if char_to_props else None)
                i = ve
            else:
                token_text = tok0[i]
                start_char = tok_char_idx[i]
                cur_text = ""
                cur_props = char_to_props[start_char] if start_char < len(char_to_props) else None
                
                for c_idx, char in enumerate(token_text):
                    props = char_to_props[start_char + c_idx] if (start_char + c_idx) < len(char_to_props) else None
                    if props == cur_props:
                        cur_text += char
                    else:
                        run = p0.add_run(cur_text)
                        apply_run_props(run, cur_props)
                        cur_text = char
                        cur_props = props
                
                if cur_text:
                    run = p0.add_run(cur_text)
                    apply_run_props(run, cur_props)
                    
                i += 1


def optimize_variables(all_data, diff_map, v_idx_list, template_doc=None, template_wb=None):
    if not all_data or not all_data[0]: return
    num_files = len(all_data)
    
    global_replacement_map = {}
    
    while True:
        var_names = [k for k in all_data[0].keys() if k != "__NAME_PATTERN__"]
        def max_len(v): return max(len(str(all_data[f].get(v, ""))) for f in range(num_files))
        var_names.sort(key=max_len, reverse=True)
        
        factorized_in_this_pass = False
        
        for i, v_big in enumerate(var_names):
            for j in range(i+1, len(var_names)):
                v_small = var_names[j]
                
                is_substring = True
                common_idx = None
                for f in range(num_files):
                    val_big = str(all_data[f].get(v_big, ""))
                    val_small = str(all_data[f].get(v_small, ""))
                    
                    if len(val_small.strip()) < 2 or not any(c.isalnum() for c in val_small):
                        is_substring = False; break
                        
                    idx = val_big.find(val_small)
                    if idx == -1:
                        is_substring = False; break
                    
                    end_idx = idx + len(val_small)
                    if idx > 0 and val_big[idx-1].isalnum():
                        is_substring = False; break
                    if end_idx < len(val_big) and val_big[end_idx].isalnum():
                        is_substring = False; break
                        
                    if common_idx is None: common_idx = idx
                    elif common_idx != idx:
                        is_substring = False; break
                        
                if is_substring:
                    prefixes = [str(all_data[f].get(v_big, ""))[:common_idx] for f in range(num_files)]
                    suffixes = [str(all_data[f].get(v_big, ""))[common_idx+len(str(all_data[f].get(v_small, ""))):] for f in range(num_files)]
                    
                    new_template_str = ""
                    if len(set(prefixes)) > 1:
                        t_str_p, var_rs_p, vals_f_p = find_all_vars_in_slot(prefixes, diff_map, v_idx_list)
                        if var_rs_p:
                            for idx_p in range(len(var_rs_p)):
                                v_name_p = get_var_name([v[idx_p] for v in vals_f_p], diff_map, v_idx_list)
                                t_str_p = t_str_p.replace(f"{{{{VAR_PLACEHOLDER_{idx_p}}}}}", f"{{{{{v_name_p}}}}}")
                                for f_idx in range(num_files): all_data[f_idx][v_name_p] = vals_f_p[f_idx][idx_p]
                        new_template_str += t_str_p
                    else:
                        new_template_str += prefixes[0]
                        
                    new_template_str += f"{{{{{v_small}}}}}"
                    
                    if len(set(suffixes)) > 1:
                        t_str_s, var_rs_s, vals_f_s = find_all_vars_in_slot(suffixes, diff_map, v_idx_list)
                        if var_rs_s:
                            for idx_s in range(len(var_rs_s)):
                                v_name_s = get_var_name([v[idx_s] for v in vals_f_s], diff_map, v_idx_list)
                                t_str_s = t_str_s.replace(f"{{{{VAR_PLACEHOLDER_{idx_s}}}}}", f"{{{{{v_name_s}}}}}")
                                for f_idx in range(num_files): all_data[f_idx][v_name_s] = vals_f_s[f_idx][idx_s]
                        new_template_str += t_str_s
                    else:
                        new_template_str += suffixes[0]
                    
                    global_replacement_map[v_big] = new_template_str
                    for f in range(num_files):
                        if v_big in all_data[f]:
                            del all_data[f][v_big]
                            
                    factorized_in_this_pass = True
                    break
                    
            if factorized_in_this_pass:
                break
                
        if not factorized_in_this_pass:
            break
            
    while True:
        var_names = list(all_data[0].keys())
        best_pair = None
        best_score = -1
        
        for i in range(len(var_names)):
            for j in range(i+1, len(var_names)):
                v1 = var_names[i]
                v2 = var_names[j]
                
                match_count = 0
                conflict_count = 0
                for f in range(num_files):
                    val1 = str(all_data[f].get(v1, "")).strip()
                    val2 = str(all_data[f].get(v2, "")).strip()
                    
                    if val1 == "None": val1 = ""
                    if val2 == "None": val2 = ""
                    
                    if val1 and val2 and val1 == val2:
                        match_count += 1
                    else:
                        if val1 == "" or val2 == "":
                            pass
                        elif val1 != "" and val2 != "":
                            # Strict check for numerics to avoid merging '5' into '15'
                            is_numeric = lambda s: bool(re.match(r'^\d+$', s.strip()))
                            if is_numeric(val1) or is_numeric(val2):
                                conflict_count += 1
                            elif val1 in val2 or val2 in val1:
                                pass
                            else:
                                conflict_count += 1
                            
                if match_count >= 2 and conflict_count == 0:
                    score = match_count
                    if score > best_score:
                        best_score = score
                        best_pair = (v1, v2)
                        
        if not best_pair:
            break
            
        v_keep, v_drop = best_pair
        global_replacement_map[v_drop] = f"{{{{{v_keep}}}}}"
        
        for f in range(num_files):
            val_keep = str(all_data[f].get(v_keep, "")).strip()
            val_drop = str(all_data[f].get(v_drop, "")).strip()
            
            if val_keep == "None": val_keep = ""
            if val_drop == "None": val_drop = ""
            
            if not val_keep and val_drop:
                all_data[f][v_keep] = val_drop
            elif val_keep and val_drop and len(val_drop) > len(val_keep):
                all_data[f][v_keep] = val_drop
                
            if v_drop in all_data[f]:
                del all_data[f][v_drop]
                
    if not global_replacement_map: return
    
    def resolve(text):
        changed = True
        while changed:
            changed = False
            for old_var, new_str in global_replacement_map.items():
                old_str = f"{{{{{old_var}}}}}"
                if old_str in text:
                    text = text.replace(old_str, new_str)
                    changed = True
        return text

    for v_big in global_replacement_map:
        global_replacement_map[v_big] = resolve(global_replacement_map[v_big])
        
    for v_big, new_str in global_replacement_map.items():
        old_str = f"{{{{{v_big}}}}}"
        if template_doc:
            for p in template_doc.paragraphs:
                for run in p.runs:
                    if old_str in run.text: run.text = run.text.replace(old_str, new_str)
            for t in template_doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            for run in p.runs:
                                if old_str in run.text: run.text = run.text.replace(old_str, new_str)
        if template_wb:
            for ws in template_wb.worksheets:
                if old_str in ws.title:
                    ws.title = ws.title.replace(old_str, new_str)
                for row in ws.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str) and old_str in cell.value:
                            cell.value = cell.value.replace(old_str, new_str)

def compare_word_group(files, template_out, relative_to_folder=None):
    docs = [Document(f) for f in files]
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    process_paragraph_list([d.paragraphs for d in docs], diff_map, v_idx, all_data)
    processed_cells = set()
    
    base_tables = docs[0].tables
    alignments_tables = [{i: [i] for i in range(len(base_tables))}]
    for f_idx in range(1, len(docs)):
        alignments_tables.append(align_blocks(base_tables, docs[f_idx].tables, lambda t: "".join(c.text for r in t.rows for c in r.cells) if t else ""))

    for t_idx in range(len(base_tables)):
        t_list = []
        for f_idx in range(len(docs)):
            comp_indices = alignments_tables[f_idx].get(t_idx, [])
            comp_idx = comp_indices[0] if comp_indices else None
            comp_tables = docs[f_idx].tables
            if comp_idx is not None and comp_idx < len(comp_tables):
                t_list.append(comp_tables[comp_idx])
            else:
                t_list.append(None)
                
        base_rows = t_list[0].rows if t_list[0] else []
        alignments_rows = [{i: [i] for i in range(len(base_rows))}]
        for f_idx in range(1, len(docs)):
            comp_rows = t_list[f_idx].rows if t_list[f_idx] else []
            alignments_rows.append(align_blocks(base_rows, comp_rows, lambda r: "".join(c.text for c in r.cells) if r else ""))

        for r_idx in range(len(base_rows)):
            r_list = []
            for f_idx in range(len(docs)):
                comp_indices = alignments_rows[f_idx].get(r_idx, [])
                comp_idx = comp_indices[0] if comp_indices else None
                comp_rows = t_list[f_idx].rows if t_list[f_idx] else []
                if comp_idx is not None and comp_idx < len(comp_rows):
                    r_list.append(comp_rows[comp_idx])
                else:
                    r_list.append(None)
                    
            base_cells = r_list[0].cells if r_list[0] else []
            c_len = len(base_cells)
            for c in range(c_len):
                c_list = [r.cells[c] if r and c < len(r.cells) else None for r in r_list]
                cell_key = tuple(cell._tc if cell is not None else None for cell in c_list)
                if any(x is None for x in cell_key) or cell_key in processed_cells:
                    continue
                processed_cells.add(cell_key)
                texts = [cell.text if cell is not None else "" for cell in c_list]
                if len(set(texts)) > 1:
                    process_paragraph_list([cell.paragraphs if cell else [] for cell in c_list], diff_map, v_idx, all_data)
                    
    optimize_variables(all_data, diff_map, v_idx, template_doc=docs[0])
    t_p = get_unique_path(template_out)
    docs[0].save(t_p)
    return t_p, all_data, v_idx[0]-1

def compare_excel_group(files, template_out, relative_to_folder=None):
    wbs_d = [openpyxl.load_workbook(f, data_only=True) for f in files]
    wb_t = openpyxl.load_workbook(files[0])
    for wb in wbs_d:
        for ws in wb.worksheets:
            normalize_sheet_numeric_cells(ws)
    for ws in wb_t.worksheets:
        normalize_sheet_numeric_cells(ws)
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    num_sheets = min(len(wb.sheetnames) for wb in wbs_d)
    for i in range(num_sheets):
        titles = [wb.sheetnames[i] for wb in wbs_d]
        if len(set(titles)) > 1:
            t_str, var_rs, vals_f = find_all_vars_in_slot(titles, diff_map, v_idx)
            if var_rs:
                for idx in range(len(var_rs)):
                    v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                    t_str = t_str.replace(f"{{{{VAR_PLACEHOLDER_{idx}}}}}", f"{{{{{v_name}}}}}")
                    for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                wb_t.worksheets[i].title = t_str
    for i in range(num_sheets):
        st = wb_t.worksheets[i]
        rows_per_doc = [list(wb.worksheets[i].iter_rows()) for wb in wbs_d]
        base_rows = rows_per_doc[0]
        alignments_rows = [{j: [j] for j in range(len(base_rows))}]
        for f_idx in range(1, len(files)):
            alignments_rows.append(align_blocks(base_rows, rows_per_doc[f_idx], lambda r: "".join(str(c.value or "") for c in r) if r else ""))

        for r_idx in range(len(base_rows)):
            r_list = []
            for f_idx in range(len(files)):
                comp_indices = alignments_rows[f_idx].get(r_idx, [])
                comp_idx = comp_indices[0] if comp_indices else None
                comp_rows = rows_per_doc[f_idx]
                if comp_idx is not None and comp_idx < len(comp_rows):
                    r_list.append(comp_rows[comp_idx])
                else:
                    r_list.append(None)
                    
            base_cells = r_list[0] if r_list[0] else []
            c_len = len(base_cells)
            for c in range(c_len):
                c_list = [r[c] if r and c < len(r) else None for r in r_list]
                cell_t = st.cell(row=r_idx+1, column=c+1)
                if type(cell_t).__name__ == 'MergedCell': continue
                if cell_t.data_type == 'f' or (isinstance(cell_t.value, str) and str(cell_t.value).startswith('=')): continue
                vals = [cell.value if cell else None for cell in c_list]
                if len(set(vals)) > 1:
                    if not all(isinstance(v, str) for v in vals if v is not None):
                        v_name = get_var_name(vals, diff_map, v_idx)
                        cell_t.value = f"{{{{{v_name}}}}}"
                        for f_idx in range(len(files)): all_data[f_idx][v_name] = vals[f_idx]
                    else:
                        t_str, var_rs, vals_f = find_all_vars_in_slot([str(v or "") for v in vals], diff_map, v_idx)
                        if var_rs:
                            for idx in range(len(var_rs)):
                                v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                                t_str = t_str.replace(f"{{{{VAR_PLACEHOLDER_{idx}}}}}", f"{{{{{v_name}}}}}")
                                for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                            cell_t.value = t_str
    

    optimize_variables(all_data, diff_map, v_idx, template_wb=wb_t)
    t_p = get_unique_path(template_out)
    wb_t.save(t_p)
    return t_p, all_data, v_idx[0]-1

def generate_name_template(filenames, all_data):
    import re
    base_names = filenames
    best_pattern = None
    max_placeholders = 0
    for i in range(len(base_names)):
        name = base_names[i]
        row_data = all_data[i]
        vars_sorted = sorted([(k, str(v).strip()) for k, v in row_data.items() if v and len(str(v).strip()) >= 2], key=lambda x: len(x[1]), reverse=True)
        pattern = name
        placeholders_count = 0
        for k, v in vars_sorted:
            if v in pattern:
                # Replace v only if it's not inside an existing {{...}} placeholder
                # We split the string by {{...}} and only replace in the literal parts
                parts = re.split(r'(\{\{.*?\}\})', pattern)
                new_parts = []
                replaced_in_this_step = False
                for part in parts:
                    if part.startswith('{{') and part.endswith('}}'):
                        new_parts.append(part)
                    else:
                        if v in part:
                            part = part.replace(v, f"{{{{{k}}}}}")
                            replaced_in_this_step = True
                        new_parts.append(part)
                if replaced_in_this_step:
                    pattern = "".join(new_parts)
                    placeholders_count += 1
        if placeholders_count > max_placeholders:
            max_placeholders = placeholders_count
            best_pattern = pattern
    if best_pattern: return best_pattern
    return f"{base_names[0]}_{{{{YYYY}}}}{{{{MM}}}}{{{{DD}}}}_{{{{hh}}}}{{{{mm}}}}"

def populate_config_sheet(ws, path, template_path, filenames, all_data, num_vars, relative_to_folder=None):
    try: rel_t = os.path.relpath(template_path, os.path.dirname(os.path.abspath(path)))
    except: rel_t = template_path
    
    rel_paths = []
    if relative_to_folder and filenames:
        folder_abs = os.path.abspath(relative_to_folder)
        for i, f in enumerate(filenames):
            try:
                f_abs = os.path.abspath(f)
                rel_file = os.path.relpath(f_abs, folder_abs)
                rel_file_no_ext = os.path.splitext(rel_file)[0]
                rel_paths.append(rel_file_no_ext.replace('\\', '/'))
            except Exception:
                rel_paths.append(os.path.splitext(os.path.basename(f))[0])
    else:
        rel_paths = [os.path.splitext(os.path.basename(f))[0] for f in filenames]
                    
    if all_data and "__NAME_PATTERN__" in all_data[0]:
        name_pattern = all_data[0]["__NAME_PATTERN__"]
        for d in all_data:
            d.pop("__NAME_PATTERN__", None)
    else:
        name_pattern = generate_name_template(rel_paths, all_data)
                
    # Check if the generated pattern matches all filenames
    pattern_matches_all = True
    if relative_to_folder and filenames:
        for i, f in enumerate(filenames):
            try:
                rel_file = os.path.relpath(os.path.abspath(f), folder_abs)
                orig_name = os.path.splitext(rel_file)[0].replace('\\', '/')
                rendered = render_string_template(name_pattern, all_data[i])
                if rendered != orig_name:
                    pattern_matches_all = False
                    break
            except Exception:
                pattern_matches_all = False
                break
            
    # If the pattern is a fallback pattern (contains YYYY) or doesn't match all original filenames, replace with file_name
    if "{{YYYY}}" in name_pattern or not pattern_matches_all:
        name_pattern = "{{rel_dir}}/{{file_name}}"
        
        # Clean up any variables that were created exclusively for the failed name pattern
        if all_data and "__PATH_ADDED_VARS__" in all_data[0]:
            path_added_vars = all_data[0]["__PATH_ADDED_VARS__"]
            for d in all_data:
                for pv in path_added_vars:
                    d.pop(pv, None)
                    
        if relative_to_folder and filenames:
            for i, f in enumerate(filenames):
                try:
                    rel_file = os.path.relpath(os.path.abspath(f), folder_abs)
                    rel_dir = os.path.dirname(rel_file).replace('\\', '/')
                    file_name = os.path.splitext(os.path.basename(f))[0]
                    if not rel_dir: rel_dir = "."
                    
                    # Template rel_dir and file_name
                    row_data = all_data[i]
                    vars_sorted = sorted([(k, str(v).strip()) for k, v in row_data.items() if v and len(str(v).strip()) >= 2 and k.startswith('field_')], key=lambda x: len(x[1]), reverse=True)
                    
                    def apply_template(text):
                        for k, v in vars_sorted:
                            if v in text:
                                parts = re.split(r'(\{\{.*?\}\})', text)
                                new_parts = []
                                for part in parts:
                                    if part.startswith('{{') and part.endswith('}}'):
                                        new_parts.append(part)
                                    else:
                                        new_parts.append(part.replace(v, f"{{{{{k}}}}}"))
                                text = "".join(new_parts)
                        return text
                        
                    all_data[i]["rel_dir"] = apply_template(rel_dir)
                    all_data[i]["file_name"] = apply_template(file_name)
                except Exception:
                    pass
        
    ws['A1'], ws['A2'] = rel_t, name_pattern
    
    headers = []
    style_props = ['bold', 'italic', 'font_size', 'font_name', 'font_color', 'fill', 'alignment', 'border', 'number_format']
    
    active_fields = set()
    for row_dict in all_data:
        for k in row_dict.keys():
            if k.startswith("field_") and k[6:].isdigit():
                active_fields.add(k)
    active_fields_sorted = sorted(list(active_fields), key=lambda x: int(x[6:]))
    
    for f_name in active_fields_sorted:
        headers.append(f_name)
        for prop in style_props:
            prop_key = f"{f_name}_{prop}"
            if any(prop_key in row_dict for row_dict in all_data):
                headers.append(prop_key)
    if all_data and "file_name" in all_data[0] and "{{file_name}}" in name_pattern:
        headers.insert(0, "file_name")
    if all_data and "rel_dir" in all_data[0] and "{{rel_dir}}" in name_pattern:
        headers.insert(0, "rel_dir")
        
    for i, h in enumerate(headers): ws.cell(row=4, column=i+1).value = h
    
    unique_rows = []
    seen = set()
    for row_dict in all_data:
        row_tuple = tuple(get_norm_key(row_dict.get(h, "")) for h in headers)
        if row_tuple not in seen:
            unique_rows.append(row_dict)
            seen.add(row_tuple)
            
    for r_idx, row_dict in enumerate(unique_rows):
        for c_idx, h in enumerate(headers):
            val = row_dict.get(h, "")
            cell = ws.cell(row=5 + r_idx, column=c_idx + 1)
            if isinstance(val, str):
                val_strip = val.strip()
                if val_strip.isdigit():
                    cell.value = int(val_strip)
                elif re.match(r'^\d+\.\d+$', val_strip):
                    cell.value = float(val_strip)
                else:
                    cell.value = val
            else:
                cell.value = val
            if isinstance(cell.value, str) and '\n' in cell.value: cell.alignment = Alignment(wrapText=True)

def save_single_config(path, template_path, all_data, num_vars, filenames, relative_to_folder=None):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Settings"
    populate_config_sheet(ws, path, template_path, filenames, all_data, num_vars, relative_to_folder=relative_to_folder)
    f_p = get_unique_path(path)
    wb.save(f_p)
    return f_p

def save_master_config(path, results, relative_to_folder=None):
    wb = openpyxl.Workbook()
    if wb.active: wb.remove(wb.active)
    for idx, (filenames, template_path, all_data, num_vars) in enumerate(results):
        base_name = os.path.splitext(os.path.basename(filenames[0]))[0]
        sheet_title = re.sub(r'[\\/*?:\[\]]', "", base_name)[:30]
        if not sheet_title: sheet_title = f"Group_{idx+1}"
        ws = wb.create_sheet(title=sheet_title)
        populate_config_sheet(ws, path, template_path, filenames, all_data, num_vars, relative_to_folder=relative_to_folder)
    f_p = get_unique_path(path)
    wb.save(f_p)
    return f_p

def create_bat_file(bat_path, config_path):
    s_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    m_py = os.path.join(s_dir, "_templates_machine_.py")
    v_py = os.path.join(s_dir, ".venv", "Scripts", "python.exe")
    if not os.path.exists(v_py):
        v_py = "python"
        
    bat_dir = os.path.dirname(os.path.abspath(bat_path))
    
    def get_rel(target):
        try:
            target_abs = os.path.abspath(target)
            if os.path.splitdrive(target_abs)[0].lower() == os.path.splitdrive(bat_dir)[0].lower():
                return os.path.relpath(target_abs, bat_dir)
        except Exception:
            pass
        return os.path.abspath(target)
        
    rel_m_py = get_rel(m_py)
    rel_config = get_rel(config_path)
    
    if v_py == "python":
        rel_v_py = "python"
    else:
        rel_v_py = get_rel(v_py)
        
    cnt = f'@echo off\ncd /d "%~dp0"\n"{rel_v_py}" "{rel_m_py}" "{rel_config}" all all\npause\n'
    with open(bat_path, "w", encoding="cp1251") as f:
        f.write(cnt)

def is_file_a_template(path, ext):
    try:
        if ext == '.docx':
            doc = Document(path)
            for p in doc.paragraphs:
                if re.search(r'\{\{.*?\}\}', p.text): return True
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        if re.search(r'\{\{.*?\}\}', c.text): return True
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if isinstance(cell, str) and re.search(r'\{\{.*?\}\}', cell): return True
    except: pass
    return False

def normalize_sheet_numeric_cells(ws):
    for row in ws.iter_rows():
        for cell in row:
            if not hasattr(cell, 'value') or cell.value is None:
                continue
            if cell.data_type == 'f' or (isinstance(cell.value, str) and str(cell.value).startswith('=')):
                continue
            val = cell.value
            if isinstance(val, str):
                val_stripped = val.strip()
                if not val_stripped:
                    continue
                # Try integer
                if re.match(r'^-?\d+$', val_stripped):
                    try:
                        cell.value = int(val_stripped)
                        if cell.number_format in [None, 'General', '@']:
                            cell.number_format = '0'
                        continue
                    except:
                        pass
                # Try float with dot decimal separator
                if re.match(r'^-?\d+\.\d+$', val_stripped):
                    try:
                        cell.value = float(val_stripped)
                        if cell.number_format in [None, 'General', '@']:
                            cell.number_format = '0.00'
                        continue
                    except:
                        pass
                # Try float with comma decimal separator (common in Ukraine)
                if re.match(r'^-?\d+,\d+$', val_stripped):
                    try:
                        normalized_val = val_stripped.replace(',', '.')
                        cell.value = float(normalized_val)
                        if cell.number_format in [None, 'General', '@']:
                            cell.number_format = '0.00'
                        continue
                    except:
                        pass
            elif isinstance(val, bool):
                continue
            elif isinstance(val, (int, float)):
                if isinstance(val, int):
                    if cell.number_format in [None, 'General', '@']:
                        cell.number_format = '0'
                elif isinstance(val, float):
                    if cell.number_format in [None, 'General', '@']:
                        cell.number_format = '0.00'

def get_structure_key(path, ext):
    try:
        if ext == '.docx':
            doc = Document(path)
            if doc.tables:
                table_shapes = tuple((len(t.rows), len(t.columns)) for t in doc.tables)
                return ("word", "with_tables", len(doc.tables), table_shapes)
            else:
                return ("word", "no_tables", len(doc.paragraphs) // 5)
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, read_only=True)
            sheet_info = []
            
            def get_cell_style_tuple(cell):
                f = cell.font
                font_tuple = (
                    bool(f.bold) if f else False,
                    bool(f.italic) if f else False,
                    float(f.size) if (f and f.size) else 11.0,
                    str(f.name) if (f and f.name) else "Calibri",
                    serialize_color(f.color) if (f and f.color) else ""
                )
                fill_str = serialize_fill(cell.fill) if cell.fill else ""
                border_str = serialize_border(cell.border) if cell.border else ""
                
                # Normalize number format in-memory for the signature
                val = cell.value
                num_format = str(cell.number_format or "General")
                if isinstance(val, str):
                    val_stripped = val.strip()
                    if val_stripped:
                        if re.match(r'^-?\d+$', val_stripped):
                            if num_format in ['General', None, '@']:
                                num_format = '0'
                        elif re.match(r'^-?\d+\.\d+$', val_stripped) or re.match(r'^-?\d+,\d+$', val_stripped):
                            if num_format in ['General', None, '@']:
                                num_format = '0.00'
                elif isinstance(val, bool):
                    pass
                elif isinstance(val, (int, float)):
                    if isinstance(val, int):
                        if num_format in ['General', None, '@']:
                            num_format = '0'
                    elif isinstance(val, float):
                        if num_format in ['General', None, '@']:
                            num_format = '0.00'
                            
                return (font_tuple, fill_str, border_str, num_format)
                
            DEFAULT_STYLE = ((False, False, 11.0, 'Calibri', ''), '', '', 'General')
            
            def normalize_header_text(val):
                if val is None: return ""
                val_str = str(val).strip()
                if re.match(r'^\d+([.,]\d+)?$', val_str):
                    return ""
                return "".join(c for c in val_str.lower() if c.isalnum())
                
            for idx, sheetname in enumerate(wb.sheetnames):
                ws = wb[sheetname]
                formulas_header = []
                formula_cols_data = set()
                header_rows_cells = {}
                sheet_styles = []
                
                for row in ws.iter_rows(max_row=100, max_col=100):
                    for cell in row:
                        if not hasattr(cell, 'row'):
                            continue
                        val = cell.value
                        r, c = cell.row, cell.column
                        
                        # Capture cell style/formatting (ignoring default styles to save space)
                        style_tup = get_cell_style_tuple(cell)
                        if style_tup != DEFAULT_STYLE:
                            sheet_styles.append((r, c, style_tup))
                            
                        if val is not None:
                            is_formula = cell.data_type == 'f' or (isinstance(val, str) and val.startswith('='))
                            if is_formula:
                                if r <= 15:
                                    formulas_header.append((r, c))
                                else:
                                    formula_cols_data.add(c)
                            else:
                                if r <= 15:
                                    if r not in header_rows_cells:
                                        header_rows_cells[r] = []
                                    header_rows_cells[r].append((c, val))
                                    
                header_row_idx = 1
                max_pop = 0
                for r in range(1, 16):
                    cells = header_rows_cells.get(r, [])
                    pop_count = sum(1 for col, v in cells if normalize_header_text(v))
                    if pop_count > max_pop:
                        max_pop = pop_count
                        header_row_idx = r
                        
                header_cells = header_rows_cells.get(header_row_idx, [])
                normalized_headers = tuple(col for col, v in header_cells if normalize_header_text(v))
                
                sheet_info.append((
                    idx,
                    tuple(formulas_header),
                    tuple(sorted(formula_cols_data)),
                    normalized_headers,
                    tuple(sheet_styles)
                ))
            return ("excel", len(wb.sheetnames), tuple(sheet_info))
    except: return None

def sort_files_by_complexity(files, ext):
    def get_complexity(f):
        try:
            if ext == '.docx':
                doc = Document(f)
                return sum(len(p.text) for p in doc.paragraphs) + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
            elif ext == '.xlsx':
                wb = openpyxl.load_workbook(f, data_only=True)
                return sum(len(str(cell.value or "")) for sheet in wb.worksheets for row in sheet.iter_rows() for cell in row)
        except:
            return 0
    sorted_files = sorted(files, key=get_complexity, reverse=True)
    if not sorted_files: return []
    median_idx = len(sorted_files) // 2
    median_file = sorted_files.pop(median_idx)
    return [median_file] + sorted_files

def run_compare_two(f1, f2, output_dir=None):
    ext = os.path.splitext(f1)[1].lower()
    if ext not in ['.docx', '.xlsx'] or os.path.splitext(f2)[1].lower() != ext:
        print("Помилка: Файли повинні мати однакове розширення (.docx або .xlsx).")
        return
    files = sort_files_by_complexity([f1, f2], ext)
    f1_a = os.path.abspath(files[0])
    f_dir = os.path.abspath(output_dir) if output_dir else os.path.dirname(f1_a)
    os.makedirs(f_dir, exist_ok=True)
    f_name = os.path.splitext(os.path.basename(f1_a))[0]
    t_o = os.path.join(f_dir, f"template_{f_name}{ext}")
    m_x = os.path.join(f_dir, f"config_{f_name}.xlsx")
    b_o = os.path.join(f_dir, f"{f_name}_run_all.bat")
    print(f"Аналіз 2 файлів (Режим: Порівняння)...")
    try:
        if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
        else: t_p, data, n_v = compare_excel_group(files, t_o)
        if n_v == 0:
            print(f"\nКонфіг та BAT-файл не створено, оскільки змінних не знайдено.")
            if os.path.exists(t_p):
                try: os.remove(t_p)
                except: pass
        else:
            m_p = save_single_config(m_x, t_p, data, n_v, files)
            create_bat_file(b_o, m_p)
            print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")
    except Exception as e: print(f"Помилка: {e}")

def run_package(sample, folder, output_dir=None):
    ext = os.path.splitext(sample)[1].lower()
    sample_key = get_structure_key(sample, ext)
    if not sample_key:
        print("Помилка: Не вдалося визначити структуру файлу-зразка.")
        return
    files = [sample]
    for root, _, fs in os.walk(folder):
        for f in fs:
            if f.lower().endswith(ext):
                if "template" in f.lower(): continue
                fp = os.path.join(root, f)
                if os.path.abspath(fp) == os.path.abspath(sample): continue
                if is_file_a_template(fp, ext): continue
                if get_structure_key(fp, ext) == sample_key:
                    files.append(fp)
    if len(files) < 2:
        print("Не знайдено подібних файлів для порівняння.")
        return
    files = sort_files_by_complexity(files, ext)
    sample_a = os.path.abspath(files[0])
    f_dir = os.path.abspath(output_dir) if output_dir else os.path.dirname(sample_a)
    os.makedirs(f_dir, exist_ok=True)
    f_name = os.path.splitext(os.path.basename(sample_a))[0]
    t_o = os.path.join(f_dir, f"template_{f_name}{ext}")
    m_x = os.path.join(f_dir, f"config_{f_name}.xlsx")
    b_o = os.path.join(f_dir, f"{f_name}_run_all.bat")
    print(f"Аналіз {len(files)} файлів (Режим: Пакетний)...")
    try:
        if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
        else: t_p, data, n_v = compare_excel_group(files, t_o)
        if n_v == 0:
            print(f"\nКонфіг та BAT-файл не створено, оскільки змінних не знайдено.")
            if os.path.exists(t_p):
                try: os.remove(t_p)
                except: pass
        else:
            m_p = save_single_config(m_x, t_p, data, n_v, files, relative_to_folder=folder)
            create_bat_file(b_o, m_p)
            print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")
    except Exception as e: print(f"Помилка: {e}")

def get_name_prefix(filename):
    name = os.path.splitext(filename)[0]
    name = re.sub(r'\d+$', '', name)
    parts = name.split('_')
    parts = [p for p in parts if p]
    if len(parts) > 1:
        last = parts[-1].lower()
        if last.isdigit() or len(last) == 1 or last in ['a', 'b', 'c', 'd', 'x', 'y', 'z']:
            parts = parts[:-1]
        if parts:
            last = parts[-1].lower()
            if last in ['pair', 'triplet', 'v2', 'v1', 'v3', 'copy']:
                parts = parts[:-1]
    res = "_".join(parts)
    res = re.sub(r'[\-_]+$', '', res)
    return res

def wb_sheetname_by_idx(filepath, idx):
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True)
        if idx < len(wb.sheetnames):
            return wb.sheetnames[idx]
    except:
        pass
    return f"Sheet {idx+1}"

def run_full_auto(folder, ignore_single=False, output_dir=None):
    print(f"Сканування папки: {folder} (Режим: Повний автопілот)...")
    groups = {}
    for root, _, fs in os.walk(folder):
        for f in fs:
            ext = os.path.splitext(f)[1].lower()
            if ext not in ['.docx', '.xlsx']: continue
            if "template" in f.lower(): continue
            fp = os.path.join(root, f)
            if is_file_a_template(fp, ext): continue
            key = get_structure_key(fp, ext)
            if key:
                if key not in groups: groups[key] = []
                groups[key].append(fp)
                
    if not groups:
        print("Не знайдено підходящих документів.")
        return
        
    print(f"\n--- Результати групування (Знайдено унікальних сигнатур: {len(groups)}) ---")
    for idx, (key, files) in enumerate(groups.items()):
        ftype = key[0]
        criteria_parts = []
        if ftype == "excel":
            num_sheets = key[1]
            criteria_parts.append(f"Тип: Excel, кількість аркушів: {num_sheets}")
            for sh_idx, (idx_sh, form_h, form_d, headers, styles) in enumerate(key[2]):
                sh_details = []
                if len(headers) > 0: sh_details.append(f"стовпців: {len(headers)}")
                if len(form_h) > 0: sh_details.append(f"формул у заголовку: {len(form_h)}")
                if len(form_d) > 0: sh_details.append(f"колонки з формулами: {list(form_d)}")
                if len(styles) > 0: sh_details.append(f"стилізованих комірок: {len(styles)}")
                if sh_details:
                    criteria_parts.append(f"  Аркуш {idx_sh+1} ({wb_sheetname_by_idx(files[0], idx_sh)}): {', '.join(sh_details)}")
        elif ftype == "word":
            if key[1] == "with_tables":
                criteria_parts.append(f"Тип: Word, з таблицями ({key[2]} шт.), розміри таблиць: {key[3]}")
            else:
                criteria_parts.append(f"Тип: Word, без таблиць (орієнтовно {key[2]*5} абзаців)")
                
        criteria_str = "\n".join(criteria_parts)
        print(f"\n[Сигнатура {idx+1}] ({len(files)} файлів):\n{criteria_str}")
        print("Склад групи:")
        for f in sorted(files):
            print(f"  - {os.path.basename(f)} (відносно: {os.path.relpath(f, folder)})")
    print("\n" + "="*60 + "\n")
    
    results = []
    out_dir = os.path.abspath(output_dir) if output_dir else os.path.abspath(folder)
    os.makedirs(out_dir, exist_ok=True)
    import shutil
    
    for key, files in groups.items():
        if len(files) < 2:
            if ignore_single:
                print(f"Група з файлом {os.path.basename(files[0])} має лише один файл, пропускаємо.")
                continue
            else:
                f = files[0]
                print(f"Група з файлом {os.path.basename(f)} має лише один файл. Копіюємо до _NODublicate_ із збереженням структури.")
                try:
                    rel_file = os.path.relpath(os.path.abspath(f), os.path.abspath(folder))
                    dest_file_path = os.path.join(out_dir, "_NODublicate_", rel_file)
                    os.makedirs(os.path.dirname(dest_file_path), exist_ok=True)
                    shutil.copy2(f, dest_file_path)
                except Exception as e:
                    print(f"Помилка при копіюванні {os.path.basename(f)}: {e}")
                continue
                
        ext = os.path.splitext(files[0])[1].lower()
        files = sort_files_by_complexity(files, ext)
        base_name = os.path.splitext(os.path.basename(files[0]))[0]
        t_o = os.path.join(out_dir, f"template_{base_name}{ext}")
        print(f"Обробка групи: {base_name} ({len(files)} файлів)...")
        try:
            if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o, folder)
            else: t_p, data, n_v = compare_excel_group(files, t_o, folder)
            print(f"  - Знайдено змінних: {n_v}")
            if n_v > 0:
                results.append((files, t_p, data, n_v))
            else:
                print(f"  - Аркуш пропущено: змінних не знайдено.")
                if os.path.exists(t_p):
                    try: os.remove(t_p)
                    except: pass
        except Exception as e:
            import traceback
            print(f"  - Помилка: {e}")
            traceback.print_exc()
            
    if not results:
        print("Не вдалося створити жодного шаблону в config_Auto.xlsx (або змінні відсутні).")
        return
        
    m_p = save_master_config(os.path.join(out_dir, "config_Auto.xlsx"), results, relative_to_folder=folder)
    create_bat_file(os.path.join(out_dir, "Auto_Run_All.bat"), m_p)
    print(f"\nГотово!\nСтворено конфіг: {m_p}\nОброблено груп: {len(results)}")
