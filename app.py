import streamlit as st
import os
import sys
import re
import openpyxl
import pandas as pd
import subprocess
import time
import logging
import warnings

# Silence noisy deprecation and user warnings
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=DeprecationWarning)

# Mute Streamlit's noisy local sources watcher tracebacks for missing optional third-party modules
logging.getLogger("streamlit.watcher.local_sources_watcher").setLevel(logging.ERROR)

# Configure page layout and style
import json
STATE_FILE = ".last_state.json"
_cached_file_state = None

def get_persisted_state_dict():
    global _cached_file_state
    if _cached_file_state is not None:
        return _cached_file_state
    if os.path.exists(STATE_FILE):
        try:
            with open(STATE_FILE, "r", encoding="utf-8") as f:
                _cached_file_state = json.load(f)
                return _cached_file_state
        except Exception:
            pass
    _cached_file_state = {}
    return _cached_file_state

def init_state_key(key, default_value):
    if key in st.session_state:
        return
    state = get_persisted_state_dict()
    if key in state and state[key] is not None:
        st.session_state[key] = state[key]
    else:
        st.session_state[key] = default_value

def load_persistent_state():
    global _cached_file_state
    _cached_file_state = None
    state = get_persisted_state_dict()
    for k, v in state.items():
        if v is not None:
            st.session_state[k] = v
    st.session_state["pm_cached_configs"] = {}
    if "sync_pm_editing_vars" in globals():
        sync_pm_editing_vars()

def make_json_serializable(obj):
    import math
    from datetime import datetime, date
    
    if obj is None:
        return None
    if isinstance(obj, (bool, str, int)):
        return obj
    if isinstance(obj, float):
        if math.isnan(obj) or math.isinf(obj):
            return str(obj)
        return obj
    if isinstance(obj, (datetime, date)):
        return obj.isoformat()
    if isinstance(obj, dict):
        return {str(k): make_json_serializable(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple, set)):
        return [make_json_serializable(v) for v in obj]
    
    # Try checking if pandas is installed and if obj is a dataframe/series or NaT/NaN
    try:
        import pandas as pd
        if isinstance(obj, (pd.DataFrame, pd.Series)):
            return make_json_serializable(obj.to_dict())
        if pd.isna(obj):
            return None
    except ImportError:
        pass

    # Try checking if numpy is installed
    try:
        import numpy as np
        if isinstance(obj, (np.integer, np.floating)):
            return make_json_serializable(obj.item())
        if isinstance(obj, np.ndarray):
            return make_json_serializable(obj.tolist())
    except ImportError:
        pass

    # Fallback to string representation if all else fails
    try:
        json.dumps(obj)
        return obj
    except TypeError:
        return str(obj)

def save_persistent_state():
    # Load the existing persisted state to avoid overwriting unrendered views' data with empty defaults
    state = dict(get_persisted_state_dict())
    
    tracked_keys = [
        "theme",
        "current_view",
        "last_opened_folder",
        "last_opened_config",
        "last_opened_template",
        "pm_folder_path",
        "editor_config_path",
        "editor_template_path",
        "editor_name_pattern",
        "pm_selected_doc",
        "analysis_mode",
        "analysis_output_dir",
        "txt_auto_folder",
        "txt_batch_sample",
        "txt_batch_folder",
        "txt_pair_file1",
        "txt_pair_file2",
        "editor_selected_sheet",
        "gen_config_path",
        "gen_output_dir",
        "pm_only_docs",
        "loaded_config_sheet",
        "current_sheet_headers",
        "current_sheet_data",
        "last_preview_row_idx",
        "loaded_gen_sheet",
        "gen_sheet_select",
        "gen_template_path",
        "gen_row_select",
        "pm_editing_vars",
        "pending_template_renames",
        "pm_loaded_doc_key",
        "gen_completion_status"
    ]
    
    for key in tracked_keys:
        if key in st.session_state:
            state[key] = make_json_serializable(st.session_state[key])
            
    try:
        with open(STATE_FILE, "w", encoding="utf-8") as f:
            json.dump(state, f, ensure_ascii=False, indent=4)
        global _cached_file_state
        _cached_file_state = state
    except Exception as e:
        try:
            import traceback
            with open("save_error.log", "w", encoding="utf-8") as f:
                traceback.print_exc(file=f)
        except Exception:
            pass

def load_excel_config(filepath):
    """Loads all worksheets from Excel config safely parsing metadata (A1 and A2), imported from core."""
    from _templates_machine_ import load_excel_config as core_load
    try:
        return core_load(filepath)
    except Exception as e:
        st.error(f"Помилка при завантаженні конфігу {filepath}: {e}")
        return None

def get_cached_config(config_path):
    """Retrieves config data from cache or loads it if stale/missing."""
    if "pm_cached_configs" not in st.session_state:
        st.session_state["pm_cached_configs"] = {}
        
    try:
        mtime = os.path.getmtime(config_path)
    except Exception:
        mtime = 0
        
    cache_entry = st.session_state["pm_cached_configs"].get(config_path)
    if cache_entry and cache_entry.get("mtime") == mtime:
        return cache_entry["data"]
        
    data = load_excel_config(config_path)
    if data:
        st.session_state["pm_cached_configs"][config_path] = {
            "mtime": mtime,
            "data": data
        }
    return data

def sync_pm_editing_vars():
    """Reloads the editing variables for the selected Project Manager document from the disk config file."""
    selected_doc = st.session_state.get("pm_selected_doc")
    if selected_doc:
        config_path = selected_doc.get("config_path")
        sheet_name = selected_doc.get("sheet_name")
        row_idx = selected_doc.get("row_idx")
        if config_path and sheet_name and row_idx is not None:
            current_doc_key = f"{config_path}_{sheet_name}_{row_idx}"
            loaded_doc_key = st.session_state.get("pm_loaded_doc_key")
            
            if current_doc_key != loaded_doc_key or st.session_state.get("pm_editing_vars") is None:
                if "pm_cached_configs" not in st.session_state:
                    st.session_state["pm_cached_configs"] = {}
                if config_path in st.session_state["pm_cached_configs"]:
                    del st.session_state["pm_cached_configs"][config_path]
                sheets_data = get_cached_config(config_path)
                if sheets_data and sheet_name in sheets_data:
                    rows = sheets_data[sheet_name].get("rows", [])
                    if row_idx < len(rows):
                        st.session_state["pm_editing_vars"] = dict(rows[row_idx])
                        st.session_state["pm_loaded_doc_key"] = current_doc_key
                        save_persistent_state()

def sync_data_editor_states():
    """Synchronizes data editor widget states (cells and headers) into session state dynamically on every rerun, ensuring unsaved edits are preserved."""
    if "current_sheet_data" in st.session_state and "current_sheet_headers" in st.session_state:
        cfg_path = st.session_state.get("editor_config_path", "")
        selected_sheet = st.session_state.get("editor_selected_sheet", "")
        if cfg_path and selected_sheet:
            clean_cfg_path = "".join([c if c.isalnum() else "_" for c in cfg_path])
            config_key = f"config_data_editor_{clean_cfg_path}_{selected_sheet}"
            headers_key = f"headers_data_editor_{clean_cfg_path}_{selected_sheet}"
        else:
            config_key = "config_data_editor"
            headers_key = "headers_data_editor"

        # 1. Sync config_data_editor (rows)
        if config_key in st.session_state and st.session_state[config_key] is not None:
            editor_state = st.session_state[config_key]
            
            # Sync edited cells
            edited_rows = editor_state.get("edited_rows", {})
            for row_idx_str, edits in edited_rows.items():
                row_idx = int(row_idx_str)
                if row_idx < len(st.session_state["current_sheet_data"]):
                    for col, val in edits.items():
                        if col == "Перегляд":
                            if val is True:
                                st.session_state["last_preview_row_idx"] = row_idx
                            elif val is False and st.session_state.get("last_preview_row_idx") == row_idx:
                                st.session_state["last_preview_row_idx"] = None
                        elif col in st.session_state["current_sheet_headers"]:
                            st.session_state["current_sheet_data"][row_idx][col] = str(val)
                            
            # Sync added rows
            added_rows = editor_state.get("added_rows", [])
            for row in added_rows:
                new_row = {h: str(row.get(h, "")) for h in st.session_state["current_sheet_headers"]}
                st.session_state["current_sheet_data"].append(new_row)
                
            # Sync deleted rows
            deleted_rows = editor_state.get("deleted_rows", [])
            for row_idx in sorted(deleted_rows, reverse=True):
                if row_idx < len(st.session_state["current_sheet_data"]):
                    st.session_state["current_sheet_data"].pop(row_idx)

            # Clear changes to prevent double-processing
            if isinstance(st.session_state.get(config_key), dict):
                st.session_state[config_key].setdefault("edited_rows", {}).clear()
                st.session_state[config_key].setdefault("added_rows", []).clear()
                st.session_state[config_key].setdefault("deleted_rows", []).clear()

        # 2. Sync headers_data_editor (headers)
        if headers_key in st.session_state and st.session_state[headers_key] is not None:
            h_editor_state = st.session_state[headers_key]
            edited_rows = h_editor_state.get("edited_rows", {})
            for row_idx_str, edits in edited_rows.items():
                row_idx = int(row_idx_str)
                if row_idx < len(st.session_state["current_sheet_headers"]):
                    orig = st.session_state["current_sheet_headers"][row_idx]
                    new_val = edits.get("Нова назва змінної", "").strip()
                    if new_val and new_val != orig:
                        if new_val not in st.session_state["current_sheet_headers"]:
                            # update header
                            st.session_state["current_sheet_headers"][row_idx] = new_val
                            # update row keys
                            for row in st.session_state["current_sheet_data"]:
                                if orig in row:
                                    row[new_val] = row.pop(orig)
                            # queue rename
                            if "pending_template_renames" not in st.session_state:
                                st.session_state["pending_template_renames"] = []
                            st.session_state["pending_template_renames"].append((orig, new_val))
                        else:
                            st.toast(f"⚠️ Змінна '{new_val}' вже існує!", icon="⚠️")

            # Clear changes to prevent double-processing
            if isinstance(st.session_state.get(headers_key), dict):
                st.session_state[headers_key].setdefault("edited_rows", {}).clear()

def sync_pm_inputs():
    """Synchronizes dynamic Project Manager variable inputs into st.session_state['pm_editing_vars']
    immediately at the start of execution, so changes are not lost during interrupted runs (like theme changes)."""
    selected_doc = st.session_state.get("pm_selected_doc")
    if selected_doc and isinstance(st.session_state.get("pm_editing_vars"), dict):
        config_path = selected_doc.get("config_path")
        sheet_name = selected_doc.get("sheet_name")
        row_idx = selected_doc.get("row_idx")
        if config_path and sheet_name and row_idx is not None:
            prefix = f"pm_input_{config_path}_{sheet_name}_{row_idx}_"
            edited_vars = st.session_state["pm_editing_vars"]
            changed = False
            for key in list(st.session_state.keys()):
                if key.startswith(prefix):
                    var_name = key[len(prefix):]
                    new_val = st.session_state[key]
                    if edited_vars.get(var_name) != new_val:
                        edited_vars[var_name] = new_val
                        changed = True
            if changed:
                save_persistent_state()

# Initialize session state variables
init_state_key("theme", "light")
if "state_loaded" not in st.session_state or st.session_state.get("theme_changed"):
    load_persistent_state()
    st.session_state["state_loaded"] = True
    st.session_state["theme_changed"] = False

theme = st.session_state["theme"]

if theme == "dark":
    bg_color = "#0f172a"
    text_color = "#f8fafc"
    card_bg = "rgba(30, 41, 59, 0.7)"
    card_border = "rgba(51, 65, 85, 0.8)"
    subtitle_color = "#94a3b8"
    input_border = "#334155"
    shadow_color = "rgba(0, 0, 0, 0.3)"
    popover_bg = "#1e293b"
    placeholder_color = "#94a3b8"
    alert_bg = "#1e293b"
    alert_border = "#334155"
else:
    bg_color = "#f8fafc"
    text_color = "#0f172a"
    card_bg = "rgba(255, 255, 255, 0.8)"
    card_border = "rgba(226, 232, 240, 0.8)"
    subtitle_color = "#64748b"
    input_border = "#e2e8f0"
    shadow_color = "rgba(31, 38, 135, 0.05)"
    popover_bg = "#ffffff"
    placeholder_color = "#64748b"
    alert_bg = "rgba(241, 245, 249, 0.8)"
    alert_border = "rgba(226, 232, 240, 0.8)"

# Configure page layout and style
st.set_page_config(
    page_title="TemplateMachine Control Center",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed"
)



# Custom premium styling
st.markdown(f"""
<style>
    /* Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;700&display=swap');
    
    :root {{
        --background-color: {bg_color};
        --text-color: {text_color};
        --secondary-background-color: {popover_bg};
    }}
    
    html, body, [data-testid="stAppViewContainer"] {{
        background-color: {bg_color} !important;
        color: {text_color} !important;
        transition: background-color 0.3s ease, color 0.3s ease;
    }}
    
    /* Hide Streamlit header, footer, and sidebar & remove top spacing */
    [data-testid="stHeader"], footer, [data-testid="stSidebar"] {{
        display: none !important;
    }}
    
    [data-testid="stMainBlockContainer"], .block-container {{
        padding-top: 0.5rem !important;
        margin-top: 0px !important;
    }}
    
    html, body, [class*="css"] {{
        font-family: 'Outfit', sans-serif;
    }}
    
    h1:not(.main-title), h2, h3, h4, h5, h6, p, span, label, li {{
        color: {text_color} !important;
    }}
    
    code, pre {{
        font-family: 'JetBrains Mono', monospace !important;
        background-color: {popover_bg} !important;
        color: {text_color} !important;
        border: 1px solid {card_border} !important;
        border-radius: 6px !important;
        padding: 0.2rem 0.4rem !important;
    }}
    
    pre {{
        padding: 1rem !important;
        overflow-x: auto !important;
        background-color: {popover_bg} !important;
    }}
    
    pre code {{
        background-color: transparent !important;
        border: none !important;
        padding: 0 !important;
    }}
    
    /* Elegant Title and Badges */
    .main-title {{
        background: linear-gradient(135deg, #4A90E2 0%, #50E3C2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3rem;
        font-weight: 700;
        margin-bottom: 0.2rem;
        letter-spacing: -0.5px;
    }}
    .subtitle {{
        color: {subtitle_color} !important;
        font-size: 1.15rem;
        margin-bottom: 2rem;
        font-weight: 300;
    }}
    
    /* Premium Styled Card Container */
    .card {{
        background: {card_bg};
        backdrop-filter: blur(10px);
        -webkit-backdrop-filter: blur(10px);
        border-radius: 16px;
        box-shadow: 0 8px 32px 0 {shadow_color};
        padding: 1.8rem;
        margin-bottom: 1.8rem;
        border: 1px solid {card_border};
        color: {text_color};
        transition: all 0.3s ease;
    }}
    
    /* Styled widgets & alerts */
    .stAlert, div[data-testid="stAlert"] {{
        border-radius: 12px !important;
        border: 1px solid {alert_border} !important;
        background-color: {alert_bg} !important;
        color: {text_color} !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.02) !important;
    }}
    
    /* Button custom hover effects */
    button[kind="primary"] {{
        background: linear-gradient(135deg, #4A90E2 0%, #357ABD 100%) !important;
        border: none !important;
        color: white !important;
        box-shadow: 0 4px 15px rgba(74, 144, 226, 0.3) !important;
        transition: all 0.25s ease-in-out !important;
    }}
    button[kind="primary"]:hover {{
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(74, 144, 226, 0.4) !important;
    }}
    
    button[kind="secondary"] {{
        border-radius: 8px !important;
        border: 1px solid {input_border} !important;
        background-color: {popover_bg} !important;
        color: {text_color} !important;
        transition: all 0.2s ease !important;
    }}
    button[kind="secondary"]:hover {{
        border-color: #4A90E2 !important;
        color: #4A90E2 !important;
        background-color: rgba(74, 144, 226, 0.03) !important;
    }}
    
    /* Code block copy toolbar buttons styled */
    button[kind="elementToolbar"], button[data-testid="stBaseButton-elementToolbar"] {{
        background-color: {popover_bg} !important;
        color: {text_color} !important;
        border: 1px solid {card_border} !important;
        border-radius: 6px !important;
    }}
    button[kind="elementToolbar"] svg, button[data-testid="stBaseButton-elementToolbar"] svg {{
        fill: {text_color} !important;
        color: {text_color} !important;
    }}
    button[kind="elementToolbar"]:hover, button[data-testid="stBaseButton-elementToolbar"]:hover {{
        background-color: #4A90E2 !important;
        color: white !important;
        border-color: #4A90E2 !important;
    }}
    button[kind="elementToolbar"]:hover svg, button[data-testid="stBaseButton-elementToolbar"]:hover svg {{
        fill: white !important;
        color: white !important;
    }}
    
    /* Tooltip popup styled */
    div[data-baseweb="tooltip"],
    div[data-baseweb="tooltip"] *,
    div[data-testid="stTooltipContent"],
    div[data-testid="stTooltipContent"] * {{
        background-color: {popover_bg} !important;
        color: {text_color} !important;
    }}
    div[data-baseweb="tooltip"] {{
        border: 1px solid {card_border} !important;
        border-radius: 8px !important;
        box-shadow: 0 4px 12px {shadow_color} !important;
    }}
    
    .badge-icon {{
        font-size: 1.5rem;
        margin-right: 0.5rem;
    }}
    
    /* Input fields and Selectboxes custom styling */
    .stTextInput input,
    .stSelectbox div[data-baseweb="select"],
    .stSelectbox div[data-baseweb="select"] > div,
    .stSelectbox div[role="combobox"],
    .stSelectbox div[role="combobox"] > div,
    .stNumberInput input,
    .stTextArea textarea,
    .stMultiSelect div[data-baseweb="select"],
    .stMultiSelect div[data-baseweb="select"] > div {{
        border-color: {input_border} !important;
        background-color: {popover_bg} !important;
        color: {text_color} !important;
    }}
    
    /* Input selected text color and control elements */
    .stSelectbox div[data-baseweb="select"] *, 
    .stSelectbox div[role="combobox"] *,
    .stMultiSelect div[data-baseweb="select"] * {{
        background-color: {popover_bg} !important;
        color: {text_color} !important;
    }}
    
    /* Input field placeholders styling */
    input::placeholder, textarea::placeholder {{
        color: {placeholder_color} !important;
        opacity: 0.75 !important;
    }}
    input::-webkit-input-placeholder, textarea::-webkit-input-placeholder {{
        color: {placeholder_color} !important;
        opacity: 0.75 !important;
    }}
    
    /* baseweb popover dropdown lists styled */
    div[data-baseweb="popover"], div[data-baseweb="menu"], ul[role="listbox"], li[role="option"] {{
        background-color: {popover_bg} !important;
        color: {text_color} !important;
    }}
    
    div[data-baseweb="popover"] ul, div[data-baseweb="popover"] li, div[data-baseweb="popover"] span {{
        background-color: {popover_bg} !important;
        color: {text_color} !important;
    }}
    
    /* baseweb hover highlights */
    div[data-baseweb="popover"] li:hover,
    div[data-baseweb="popover"] li[aria-selected="true"],
    div[data-baseweb="popover"] li:hover * {{
        background-color: #4A90E2 !important;
        color: #ffffff !important;
    }}
    
    /* Document Preview Container (Sheet/Word) Overrides */
    div.document-preview-container,
    div.document-preview-container *,
    div.word-preview-container,
    div.word-preview-container *,
    .document-preview-container p,
    .document-preview-container span,
    .document-preview-container td,
    .document-preview-container th,
    .document-preview-container tr,
    .document-preview-container table,
    .document-preview-container h1,
    .document-preview-container h2,
    .document-preview-container h3,
    .document-preview-container h4,
    .document-preview-container h5,
    .document-preview-container h6,
    .document-preview-container li,
    .document-preview-container ul,
    .word-preview-container p,
    .word-preview-container span,
    .word-preview-container td,
    .word-preview-container th,
    .word-preview-container tr,
    .word-preview-container table,
    .word-preview-container h1,
    .word-preview-container h2,
    .word-preview-container h3,
    .word-preview-container h4,
    .word-preview-container h5,
    .word-preview-container h6,
    .word-preview-container li,
    .word-preview-container ul {{
        color: #1a202c !important;
    }}
    div.document-preview-container,
    div.word-preview-container,
    div.document-preview-container div,
    div.word-preview-container div {{
        background-color: #ffffff !important;
    }}
    
    /* Details & Expanders */
    div[data-testid="stExpander"] {{
        background-color: {card_bg} !important;
        border: 1px solid {card_border} !important;
        border-radius: 8px !important;
        margin-bottom: 0.5rem !important;
    }}
    
    /* Force details and open details to be transparent to show the wrapper card background */
    div[data-testid="stExpander"] details,
    div[data-testid="stExpander"] details[open],
    div[data-testid="stExpander"] summary,
    div[data-testid="stExpander"] div[role="region"],
    div[data-testid="stExpander"] [data-testid="stExpanderDetails"],
    div[data-testid="stExpander"] .streamlit-expanderContent,
    div[data-testid="stExpander"] div[role="region"] [data-testid="stVerticalBlock"] {{
        background-color: transparent !important;
        background: transparent !important;
        border: none !important;
    }}
    
    /* Make sure all descendant container divs inside the expander are transparent,
       excluding actual widgets like input fields, selectboxes, buttons, alerts and nested expanders */
    div[data-testid="stExpander"] div[role="region"] div:not(.stTextInput):not(.stSelectbox):not(.stTextArea):not(.stButton):not(.stNumberInput):not(.stMultiSelect):not([data-testid="stExpander"]):not(.stAlert):not([data-testid="stAlert"]) {{
        background-color: transparent !important;
        background: transparent !important;
    }}
    
    /* Summary headers inside expanders */
    div[data-testid="stExpander"] > details > summary,
    div[data-testid="stExpander"] > details > summary * {{
        background-color: transparent !important;
        color: {text_color} !important;
    }}
    
    /* Make sure nested expanders restore their background and don't remain transparent */
    div[data-testid="stExpander"] div[role="region"] div[data-testid="stExpander"] {{
        background-color: {card_bg} !important;
        border: 1px solid {card_border} !important;
    }}
    
    /* Scoped styling for help page when the marker is present in DOM */
    div[data-testid="stMainBlockContainer"]:has(.help-page-marker),
    .block-container:has(.help-page-marker) {{
        background: {card_bg} !important;
        backdrop-filter: blur(10px) !important;
        -webkit-backdrop-filter: blur(10px) !important;
        border: 1px solid {card_border} !important;
        border-radius: 16px !important;
        box-shadow: 0 8px 32px 0 {shadow_color} !important;
        padding: 3rem !important;
        margin-top: 1.5rem !important;
        margin-bottom: 2.5rem !important;
    }}
    
    /* Responsive Theme Toggle */
</style>
""", unsafe_allow_html=True)

# ----------------------------------------------------
# HELPER FUNCTIONS FOR FILE & WINDOWS DIALOGS
# ----------------------------------------------------

def open_folder_picker(title="Оберіть папку", initialdir=None):
    """Opens a native Windows directory selection dialog."""
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.wm_attributes('-topmost', 1)
        if not initialdir:
            initialdir = st.session_state.get("last_opened_folder") or ""
        if not initialdir or not os.path.exists(initialdir) or not os.path.isdir(initialdir):
            initialdir = os.getcwd()
        folder = filedialog.askdirectory(parent=root, title=title, initialdir=initialdir)
        root.destroy()
        return folder
    except Exception as e:
        st.warning("Не вдалося відкрити діалогове вікно Windows. Будь ласка, введіть шлях вручну.")
        return ""

def open_file_picker(filetypes=None, initialdir=None):
    """Opens a native Windows file selection dialog."""
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk()
        root.withdraw()
        root.wm_attributes('-topmost', 1)
        if filetypes is None:
            filetypes = [("Усі підтримувані", "*.docx;*.xlsx"), ("Word файли", "*.docx"), ("Excel файли", "*.xlsx"), ("Усі файли", "*.*")]
        if not initialdir:
            initialdir = st.session_state.get("last_opened_folder") or ""
        if not initialdir or not os.path.exists(initialdir) or not os.path.isdir(initialdir):
            initialdir = os.getcwd()
        file = filedialog.askopenfilename(parent=root, title="Оберіть файл-зразок", filetypes=filetypes, initialdir=initialdir)
        root.destroy()
        return file
    except Exception as e:
        st.warning("Не вдалося відкрити діалогове вікно Windows. Будь ласка, введіть шлях вручну.")
        return ""

def create_new_sheet(filepath, new_sheet_name):
    """Creates a new worksheet with default metadata and template placeholders."""
    try:
        wb = openpyxl.load_workbook(filepath)
        if new_sheet_name in wb.sheetnames:
            st.error(f"Аркуш з назвою '{new_sheet_name}' вже існує!")
            return False
        ws = wb.create_sheet(title=new_sheet_name)
        # Initialize default values
        ws['A1'] = "template_placeholder.docx"
        ws['A2'] = "output_{{YYYY}}{{MM}}{{DD}}.docx"
        ws.cell(row=4, column=1).value = "field_1"
        ws.cell(row=5, column=1).value = "значення_1"
        wb.save(filepath)
        return True
    except Exception as e:
        st.error(f"Помилка при створенні аркуша: {e}")
        return False

def rename_sheet(filepath, old_sheet_name, new_sheet_name):
    """Safely renames a sheet in the Excel config."""
    try:
        wb = openpyxl.load_workbook(filepath)
        if old_sheet_name not in wb.sheetnames:
            st.error(f"Аркуш '{old_sheet_name}' не знайдено!")
            return False
        if new_sheet_name in wb.sheetnames:
            st.error(f"Аркуш з назвою '{new_sheet_name}' вже існує!")
            return False
        ws = wb[old_sheet_name]
        ws.title = new_sheet_name
        wb.save(filepath)
        return True
    except Exception as e:
        st.error(f"Помилка при перейменуванні аркуша: {e}")
        return False

def delete_sheet(filepath, sheet_name):
    """Deletes a sheet from the Excel config."""
    try:
        wb = openpyxl.load_workbook(filepath)
        if sheet_name not in wb.sheetnames:
            st.error(f"Аркуш '{sheet_name}' не знайдено!")
            return False
        if len(wb.sheetnames) <= 1:
            st.error("Неможливо видалити єдиний аркуш у книзі!")
            return False
        wb.remove(wb[sheet_name])
        wb.save(filepath)
        return True
    except Exception as e:
        st.error(f"Помилка при видаленні аркуша: {e}")
        return False

def update_config_template_path(filepath, sheet_name, new_template_path):
    """Safely updates only the template path (cell A1) of a specific sheet in Excel config."""
    try:
        # Convert template path to relative path if possible
        config_dir = os.path.dirname(os.path.abspath(filepath))
        try:
            target_abs = os.path.abspath(new_template_path)
            if os.path.splitdrive(target_abs)[0].lower() == os.path.splitdrive(config_dir)[0].lower():
                new_template_path = os.path.relpath(target_abs, config_dir).replace('\\', '/')
        except Exception:
            pass
            
        wb = openpyxl.load_workbook(filepath)
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            ws.cell(row=1, column=1).value = new_template_path
            wb.save(filepath)
            return True
        return False
    except Exception as e:
        st.error(f"Помилка при оновленні шляху шаблону в конфігу: {e}")
        return False

def get_formatted_documentation_markdown():
    """Returns the full technical guide content formatted as premium Markdown, loaded dynamically from _templates_machine_.txt."""
    try:
        filepath = "_templates_machine_.txt"
        if os.path.exists(filepath):
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
    except Exception as e:
        st.warning(f"Не вдалося завантажити файл довідки: {e}")
    
    # Fallback to absolute bare-bones in case the file gets deleted
    return "### 📖 Документація\nНе вдалося завантажити файл довідки `_templates_machine_.txt`. Будь ласка, переконайтеся, що файл існує у робочій папці."

# ----------------------------------------------------
# PROJECT MANAGER HELPERS & SCANNERS
# ----------------------------------------------------

def is_excel_config(file_path):
    """Robust algorithm to distinguish config Excel files from regular spreadsheets."""
    if not file_path.endswith('.xlsx') or os.path.basename(file_path).startswith('~$'):
        return False
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # Read cells A1 and A2
            a1 = ws.cell(row=1, column=1).value
            a2 = ws.cell(row=2, column=1).value
            if a1 and isinstance(a1, str):
                a1_lower = a1.lower()
                if any(ext in a1_lower for ext in ['.docx', '.xlsx']) or 'template' in a1_lower:
                    return True
            if a2 and isinstance(a2, str):
                a2_lower = a2.lower()
                if any(ext in a2_lower for ext in ['.docx', '.xlsx']) or '{{' in a2_lower:
                    return True
        return False
    except Exception:
        return False

def scan_recursive_configs(root_folder):
    """Walks the folder structure recursively and returns a list of config files."""
    config_files = []
    if not os.path.exists(root_folder) or not os.path.isdir(root_folder):
        return []
    for dirpath, _, filenames in os.walk(root_folder):
        for f in filenames:
            if f.endswith('.xlsx') and not f.startswith('~$'):
                full_path = os.path.abspath(os.path.join(dirpath, f))
                if is_excel_config(full_path):
                    config_files.append(full_path)
    return sorted(config_files)

def build_dir_tree(config_files, root_path):
    """Builds a nested directory tree dictionary from config paths."""
    tree = {}
    for path in config_files:
        rel_path = os.path.relpath(path, root_path)
        parts = rel_path.split(os.sep)
        current = tree
        for part in parts[:-1]:
            if part not in current:
                current[part] = {}
            current = current[part]
        current[parts[-1]] = path
    return tree

def build_docs_only_tree(config_files, root_path):
    """Builds a virtual tree where documents are grouped directly under folder structures,
    skipping config files and sheets nodes."""
    tree = {}
    abs_root_path = os.path.abspath(root_path)
    
    for path in config_files:
        abs_config_path = os.path.abspath(path)
        config_dir = os.path.dirname(abs_config_path)
            
        sheets_data = get_cached_config(abs_config_path)
        if not sheets_data:
            continue
            
        for sheet_name, info in sheets_data.items():
            rows = info.get("rows", [])
            template_path = info.get("template_path", "")
            name_pattern = info.get("name_pattern", "")
            
            for idx, row in enumerate(rows):
                is_selected = (
                    st.session_state.get("pm_selected_doc", {}).get("config_path") == abs_config_path and
                    st.session_state.get("pm_selected_doc", {}).get("sheet_name") == sheet_name and
                    st.session_state.get("pm_selected_doc", {}).get("row_idx") == idx
                )
                current_row_vars = row
                if is_selected and st.session_state.get("pm_editing_vars") is not None:
                    current_row_vars = st.session_state["pm_editing_vars"]
                
                doc_name = resolve_virtual_doc_name(name_pattern, current_row_vars, template_path)
                if not doc_name.strip():
                    doc_name = f"document_{idx + 1}"
                
                # Resolve absolute path of the document
                if os.path.isabs(doc_name):
                    abs_doc_path = os.path.abspath(doc_name)
                else:
                    abs_doc_path = os.path.abspath(os.path.join(config_dir, doc_name))
                    
                abs_doc_dir = os.path.dirname(abs_doc_path)
                
                # Determine relative folder directory to root_path
                try:
                    rel_doc_dir = os.path.relpath(abs_doc_dir, abs_root_path)
                    norm_rel = rel_doc_dir.replace("\\", "/")
                    rel_parts = [p for p in norm_rel.split("/") if p]
                    # If it goes outside the root (starts with '..') or is absolute, use absolute path
                    if (rel_parts and rel_parts[0] == "..") or os.path.isabs(rel_doc_dir):
                        rel_doc_dir = abs_doc_dir
                except ValueError:
                    rel_doc_dir = abs_doc_dir
                
                # Normalize separators for path splitting
                rel_doc_dir = rel_doc_dir.replace("\\", "/")
                
                if rel_doc_dir == "." or not rel_doc_dir:
                    parts = []
                else:
                    parts = rel_doc_dir.split("/")
                    parts = [p for p in parts if p and p != "."]
                    
                doc_info = {
                    "config_path": abs_config_path,
                    "sheet_name": sheet_name,
                    "row_idx": idx,
                    "doc_name": doc_name,
                    "template_path": template_path,
                    "name_pattern": name_pattern
                }
                
                current = tree
                for part in parts:
                    if part not in current:
                        current[part] = {}
                    current = current[part]
                
                if "__docs__" not in current:
                    current["__docs__"] = []
                current["__docs__"].append(doc_info)
                
    return tree

def render_docs_only_tree(tree, depth=0, path_prefix=""):
    selected_doc = st.session_state.get("pm_selected_doc")
    if selected_doc:
        config_path_clean = "".join([c if c.isalnum() else "_" for c in selected_doc.get('config_path', '')])
        selected_doc_id = f"{config_path_clean}_{selected_doc.get('sheet_name')}_{selected_doc.get('row_idx')}"
    else:
        selected_doc_id = "none"
    target_config = selected_doc.get("config_path") if selected_doc else None
    
    # 1. Render subfolders
    for key, value in tree.items():
        if key == "__docs__":
            continue
        def contains_selected(folder_dict):
            if "__docs__" in folder_dict:
                for doc in folder_dict["__docs__"]:
                    if target_config and os.path.abspath(doc["config_path"]).lower() == os.path.abspath(target_config).lower():
                        if selected_doc.get("sheet_name") == doc["sheet_name"] and selected_doc.get("row_idx") == doc["row_idx"]:
                            return True
            for k, v in folder_dict.items():
                if k != "__docs__" and isinstance(v, dict):
                    if contains_selected(v):
                        return True
            return False
            
        should_expand = (depth == 0) or (selected_doc and contains_selected(value))
        exp_key = f"pm_docs_only_exp_{path_prefix}_{key}_{selected_doc_id}"
        with st.expander("📁 " + key, expanded=should_expand, key=exp_key):
            render_docs_only_tree(value, depth + 1, path_prefix=f"{path_prefix}_{key}")
            
    # 2. Render files (documents) in the current folder
    if "__docs__" in tree:
        for doc in tree["__docs__"]:
            is_selected = (
                selected_doc and
                selected_doc.get("config_path") == doc["config_path"] and
                selected_doc.get("sheet_name") == doc["sheet_name"] and
                selected_doc.get("row_idx") == doc["row_idx"]
            )
            button_type = "primary" if is_selected else "secondary"
            if st.button(
                f"📄 {os.path.basename(doc['doc_name'])}",
                key=f"pm_doc_btn_{doc['config_path']}_{doc['sheet_name']}_{doc['row_idx']}",
                use_container_width=True,
                type=button_type
            ):
                st.session_state["pm_selected_doc"] = doc
                st.session_state["last_opened_config"] = doc["config_path"]
                st.session_state["last_opened_folder"] = os.path.dirname(os.path.abspath(doc["config_path"]))
                st.session_state["last_opened_template"] = doc["template_path"]
                st.session_state["pending_download"] = None
                clear_pm_input_keys()
                save_persistent_state()
                st.session_state["pm_editing_vars"] = None
                save_persistent_state()
                st.rerun()


def resolve_virtual_doc_name(pattern, row_data, template_path):
    """Resolves output document name using date/time variables and row data."""
    from datetime import datetime
    now = datetime.now()
    now_vars = {
        "YYYY": now.strftime("%Y"),
        "MM": now.strftime("%m"),
        "DD": now.strftime("%d"),
        "hh": now.strftime("%H"),
        "mm": now.strftime("%M"),
        "ss": now.strftime("%S")
    }
    variables = {**now_vars, **row_data}
    
    result = str(pattern)
    for key, val in variables.items():
        pattern_re = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
        result = re.sub(pattern_re, str(val), result)
        
    if template_path:
        ext = os.path.splitext(template_path)[1].lower()
        if not result.lower().endswith(ext):
            result += ext
    return result

def save_generated_document_dialog(template_path, variables, config_path, name_pattern=None):
    """Generates the document and opens a native Windows dialog to save it, with fallback."""
    import tkinter as tk
    from tkinter import filedialog
    from _templates_machine_ import process_word, process_excel
    import tempfile
    
    from _templates_machine_ import resolve_path
    cfg_dir = os.path.dirname(os.path.abspath(config_path))
    actual_t_path = resolve_path(cfg_dir, template_path)
        
    if not os.path.exists(actual_t_path):
        st.error(f"Шаблон не знайдено: {template_path}")
        return
        
    ext = os.path.splitext(actual_t_path)[1].lower()
    
    if not name_pattern:
        name_pattern = variables.get("name_pattern", "document")
    proposed_filename = os.path.basename(resolve_virtual_doc_name(name_pattern, variables, template_path))
    
    # Attempt native save dialog
    saved = False
    try:
        root = tk.Tk()
        root.withdraw()
        root.wm_attributes('-topmost', 1)
        
        filetypes = [("Word Document", "*.docx")] if ext == ".docx" else [("Excel Workbook", "*.xlsx")]
        default_ext = ext
        
        save_path = filedialog.asksaveasfilename(
            parent=root,
            title="Зберегти згенерований документ",
            filetypes=filetypes,
            defaultextension=default_ext,
            initialfile=proposed_filename
        )
        root.destroy()
        
        if save_path:
            if ext == ".docx":
                process_word(actual_t_path, save_path, variables)
            elif ext == ".xlsx":
                process_excel(actual_t_path, save_path, variables)
            st.success(f"🎉 Документ успішно збережено: {save_path}")
            saved = True
    except Exception as e:
        st.warning(f"Не вдалося відкрити діалогове вікно збереження Windows: {e}. Використовуємо завантаження через браузер.")
        
    if not saved:
        # Fallback to browser download button
        temp_dir = tempfile.gettempdir()
        temp_file = os.path.join(temp_dir, proposed_filename)
        try:
            if ext == ".docx":
                process_word(actual_t_path, temp_file, variables)
            elif ext == ".xlsx":
                process_excel(actual_t_path, temp_file, variables)
                
            with open(temp_file, "rb") as f:
                file_bytes = f.read()
                
            st.session_state["pending_download"] = {
                "bytes": file_bytes,
                "name": proposed_filename,
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == ".docx" else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            }
            save_persistent_state()
            st.rerun()
        except Exception as e_gen:
            st.error(f"Помилка генерації документа: {e_gen}")

# ----------------------------------------------------
# SYSTEM STATE & WORKSPACE SCANNING
# ----------------------------------------------------

def scan_workspace():
    """Scans the directory for configs and templates."""
    files = os.listdir('.')
    configs = [f for f in files if f.endswith('.xlsx') and not f.startswith('~$') and 'template' not in f.lower()]
    templates = [f for f in files if (f.endswith('.docx') or f.endswith('.xlsx')) and f.startswith('template_')]
    
    # Check inside example directory as well
    if os.path.exists('example'):
        for f in os.listdir('example'):
            if f.endswith('.xlsx') and not f.startswith('~$') and 'template' not in f.lower():
                configs.append(os.path.join('example', f))
            if (f.endswith('.docx') or f.endswith('.xlsx')) and f.startswith('template_'):
                templates.append(os.path.join('example', f))
                
    return sorted(configs), sorted(templates)

# ----------------------------------------------------
# EXCEL CONFIG READ/WRITE WITH OPENPYXL
# ----------------------------------------------------

def save_excel_config(filepath, sheet_name, template_path, name_pattern, headers, df_data):
    """Saves changes back to Excel config safely preserving other sheets, imported from core."""
    from _templates_machine_ import save_excel_config as core_save
    try:
        return core_save(filepath, sheet_name, template_path, name_pattern, headers, df_data)
    except Exception as e:
        st.error(f"Помилка при збереженні змін: {e}")
        return False

# ----------------------------------------------------
# LIVE SUBPROCESS LOG STREAMING
# ----------------------------------------------------

def run_subprocess_and_stream(args):
    """Runs the python automation script and streams console outputs to Streamlit."""
    python_exe = sys.executable
    script_path = "_templates_machine_.py"
    
    cmd = [python_exe, script_path] + args
    
    log_area = st.empty()
    progress_bar = st.progress(0.0)
    
    st.info(f"🚀 Запущено команду: `python _templates_machine_.py {' '.join(f'\"{a}\"' for a in args)}`")
    
    # Initialize last operation state
    st.session_state["last_operation_logs"] = []
    st.session_state["last_operation_status"] = "running"
    st.session_state["last_operation_cmd"] = f"python _templates_machine_.py {' '.join(f'\"{a}\"' for a in args)}"
    
    # Force Python stdout to UTF-8 using environment variables
    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"
    
    # Initialize process in binary mode to safely handle dynamic decoding fallbacks
    process = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=False,
        env=env
    )
    
    logs = []
    
    while True:
        line_bytes = process.stdout.readline()
        if not line_bytes:
            break
            
        # Bulletproof byte decoding fallback strategy (UTF-8 -> CP1251 -> UTF-8 with replacement)
        line = ""
        try:
            line = line_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                line = line_bytes.decode('cp1251')
            except Exception:
                line = line_bytes.decode('utf-8', errors='replace')
                
        cleaned_line = line.strip()
        if cleaned_line:
            logs.append(cleaned_line)
            st.session_state["last_operation_logs"] = logs
            # Display logs in real-time
            log_area.code("\n".join(logs[-15:])) # Show last 15 lines of progress
            
            # Simple progress heuristics
            if "Створено" in cleaned_line or "Обробка" in cleaned_line:
                progress_bar.progress(0.5)
            elif "Готово!" in cleaned_line or "Успішно!" in cleaned_line:
                progress_bar.progress(1.0)
                
    process.wait()
    
    if process.returncode == 0:
        progress_bar.progress(1.0)
        st.session_state["last_operation_status"] = "success"
        st.session_state["trigger_balloons"] = True
        st.session_state["gen_completion_status"] = "success"
    else:
        st.session_state["last_operation_status"] = "error"
        st.session_state["trigger_balloons"] = False
        st.session_state["gen_completion_status"] = "error"
        
    return process.returncode, logs

# ----------------------------------------------------
# HELPER FOR PERSISTENT CONSOLE LOGS
# ----------------------------------------------------

def show_last_operation_logs():
    """Displays the persistent logs of the last executed operation from session state."""
    if "last_operation_logs" in st.session_state and st.session_state["last_operation_logs"]:
        status = st.session_state.get("last_operation_status", "")
        cmd = st.session_state.get("last_operation_cmd", "")
        
        # Trigger balloons only once
        if st.session_state.get("trigger_balloons"):
            st.balloons()
            st.session_state["trigger_balloons"] = False
            
        st.markdown("---")
        if status == "success":
            st.success("🎉 Завдання успішно виконано!")
            st.markdown("### 🟢 Результат останньої операції: Успішно")
        elif status == "error":
            st.error(f"❌ Помилка виконання завдання!")
            st.markdown("### 🔴 Результат останньої операції: Помилка")
        else:
            st.markdown("### 🟡 Результат останньої операції: Виконується")
            
        if cmd:
            st.caption(f"Команда: `{cmd}`")
            
        with st.expander("📋 Показати повний лог консолі", expanded=True):
            st.code("\n".join(st.session_state["last_operation_logs"]))

def recreate_bat_file(bat_path, config_path):
    """Recreates the execution .bat file at the destination with correct relative paths."""
    try:
        s_dir = os.path.dirname(os.path.abspath(__file__))
        m_py = os.path.join(s_dir, "_templates_machine_.py")
        v_py = os.path.join(s_dir, ".venv", "Scripts", "python.exe")
        if not os.path.exists(v_py):
            v_py = "python"
            
        bat_dir = os.path.dirname(os.path.abspath(bat_path))
        
        def get_rel(target):
            try:
                target_abs = os.path.abspath(target)
                if os.path.splitdrive(target_abs)[0].lower() == os.path.splitdrive(bat_dir)[0].lower():
                    return os.path.relpath(target_abs, bat_dir)
            except Exception:
                pass
            return os.path.abspath(target)
            
        rel_m_py = get_rel(m_py)
        rel_config = get_rel(config_path)
        
        if v_py == "python":
            rel_v_py = "python"
        else:
            rel_v_py = get_rel(v_py)
            
        cnt = f'@echo off\ncd /d "%~dp0"\n"{rel_v_py}" "{rel_m_py}" "{rel_config}" all all\npause\n'
        with open(bat_path, "w", encoding="cp1251") as f:
            f.write(cnt)
    except Exception as e:
        st.error(f"Не вдалося оновити виконавчий .bat файл: {e}")

def move_autopilot_outputs(dest_dir):
    """Moves generated templates and configs of autopilot mode to a custom destination directory."""
    if not dest_dir:
        return
    try:
        import shutil
        dest_dir = os.path.abspath(dest_dir)
        os.makedirs(dest_dir, exist_ok=True)
        # Move Auto_Config.xlsx and Auto_Run_All.bat
        for filename in ["Auto_Config.xlsx", "Auto_Run_All.bat"]:
            src = os.path.join(os.getcwd(), filename)
            if os.path.exists(src):
                shutil.move(src, os.path.join(dest_dir, filename))
        # Move all template_* files from current directory
        moved_count = 0
        for filename in os.listdir(os.getcwd()):
            if filename.startswith("template_") and os.path.isfile(filename):
                shutil.move(filename, os.path.join(dest_dir, filename))
                moved_count += 1
                
        # Recreate the moved BAT file with relative paths inside the destination folder!
        dest_bat = os.path.join(dest_dir, "Auto_Run_All.bat")
        dest_cfg = os.path.join(dest_dir, "Auto_Config.xlsx")
        if os.path.exists(dest_bat) and os.path.exists(dest_cfg):
            recreate_bat_file(dest_bat, dest_cfg)
            
        st.toast(f"✅ Результати аналізу перенесено в: {dest_dir} (переміщено {moved_count} шаблонів)", icon="📁")
    except Exception as e:
        st.error(f"Помилка при перенесенні файлів результатів: {e}")

def move_batch_outputs(sample_path, dest_dir):
    """Moves generated templates and configs of batch mode to a custom destination directory."""
    if not dest_dir or not sample_path:
        return
    try:
        import shutil
        dest_dir = os.path.abspath(dest_dir)
        os.makedirs(dest_dir, exist_ok=True)
        f_dir = os.path.dirname(os.path.abspath(sample_path))
        base_name, ext = os.path.splitext(os.path.basename(sample_path))
        
        files_to_move = [
            f"template_{base_name}{ext}",
            f"{base_name}_config.xlsx",
            f"{base_name}_run_all.bat"
        ]
        moved_count = 0
        for filename in files_to_move:
            src = os.path.join(f_dir, filename)
            if os.path.exists(src):
                shutil.move(src, os.path.join(dest_dir, filename))
                moved_count += 1
                
        # Recreate the moved BAT file with relative paths inside the destination folder!
        dest_bat = os.path.join(dest_dir, f"{base_name}_run_all.bat")
        dest_cfg = os.path.join(dest_dir, f"{base_name}_config.xlsx")
        if os.path.exists(dest_bat) and os.path.exists(dest_cfg):
            recreate_bat_file(dest_bat, dest_cfg)
            
        st.toast(f"✅ Результати пакетного аналізу перенесено в: {dest_dir} (переміщено {moved_count} файлів)", icon="📁")
    except Exception as e:
        st.error(f"Помилка при перенесенні файлів результатів: {e}")

def move_pairwise_outputs(file1_path, dest_dir):
    """Moves generated templates and configs of pairwise mode to a custom destination directory."""
    if not dest_dir or not file1_path:
        return
    try:
        import shutil
        dest_dir = os.path.abspath(dest_dir)
        os.makedirs(dest_dir, exist_ok=True)
        f_dir = os.path.dirname(os.path.abspath(file1_path))
        base_name, ext = os.path.splitext(os.path.basename(file1_path))
        
        files_to_move = [
            f"template_{base_name}{ext}",
            f"{base_name}_config.xlsx",
            f"{base_name}_run_all.bat"
        ]
        moved_count = 0
        for filename in files_to_move:
            src = os.path.join(f_dir, filename)
            if os.path.exists(src):
                shutil.move(src, os.path.join(dest_dir, filename))
                moved_count += 1
                
        # Recreate the moved BAT file with relative paths inside the destination folder!
        dest_bat = os.path.join(dest_dir, f"{base_name}_run_all.bat")
        dest_cfg = os.path.join(dest_dir, f"{base_name}_config.xlsx")
        if os.path.exists(dest_bat) and os.path.exists(dest_cfg):
            recreate_bat_file(dest_bat, dest_cfg)
            
        st.toast(f"✅ Результати попарного порівняння перенесено в: {dest_dir} (переміщено {moved_count} файлів)", icon="📁")
    except Exception as e:
        st.error(f"Помилка при перенесенні файлів результатів: {e}")

# ----------------------------------------------------
# CORE DASHBOARD INTERFACE LAYOUT
# ----------------------------------------------------

# Initialize session state variables for file and folder picker widgets
for key in [
    "txt_auto_folder", "txt_batch_sample", "txt_batch_folder", 
    "txt_pair_file1", "txt_pair_file2", "editor_config_path", 
    "editor_template_path", "editor_name_pattern",
    "gen_config_path", "gen_output_dir", "analysis_output_dir",
    "last_opened_folder", "last_opened_config", "last_opened_template",
    "pm_folder_path", "analysis_mode", "editor_selected_sheet", "current_view",
    "loaded_config_sheet", "loaded_gen_sheet", "gen_template_path"
]:
    init_state_key(key, "")

init_state_key("gen_sheet_select", "all (Всі аркуші)")
init_state_key("gen_row_select", "all (Всі рядки з даними)")
init_state_key("pm_only_docs", False)
init_state_key("last_operation_logs", [])
init_state_key("last_operation_status", None)
init_state_key("last_operation_cmd", "")
init_state_key("pm_selected_doc", None)
init_state_key("current_sheet_headers", [])
init_state_key("current_sheet_data", [])
init_state_key("last_preview_row_idx", 0)
init_state_key("pm_editing_vars", None)
init_state_key("pending_template_renames", [])
init_state_key("pm_loaded_doc_key", "")
init_state_key("gen_completion_status", None)
init_state_key("last_gen_params", "")

def clear_pm_input_keys():
    """Removes all keys starting with pm_input_ from session state to force refreshing inputs."""
    keys_to_del = [k for k in st.session_state.keys() if k.startswith("pm_input_")]
    for k in keys_to_del:
        del st.session_state[k]

# Sync data editor states immediately on every run
sync_data_editor_states()
sync_pm_inputs()

def generate_docx_preview(template_path, variables, config_path=None):
    """Generates a temporary Word document and extracts its content for quick preview in high-fidelity HTML."""
    import tempfile
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from _templates_machine_ import process_word
    
    # Resolve relative template paths relative to the current Excel config directory if needed
    from _templates_machine_ import resolve_path
    c_path = config_path or st.session_state.get("editor_config_path")
    cfg_dir = os.path.dirname(os.path.abspath(c_path)) if c_path else os.getcwd()
    actual_path = resolve_path(cfg_dir, template_path)
            
    if not os.path.exists(actual_path):
        return f"<div style='color: red !important; font-weight: bold;'>Шаблон не знайдено за шляхом: {template_path}</div>"
        
    temp_dir = tempfile.gettempdir()
    temp_out = os.path.join(temp_dir, "temp_preview.docx")
    
    try:
        process_word(actual_path, temp_out, variables)
        doc = Document(temp_out)
        
        html = []
        html.append("<div class='word-preview-container' style='font-family: \"Times New Roman\", Times, serif; color: #1a202c !important; max-width: 800px; margin: 0 auto; line-height: 1.5; font-size: 14px; background-color: #ffffff !important;'>")
        
        # Pre-map paragraphs and tables to speed up lookup to O(1)
        p_map = {p._p: p for p in doc.paragraphs}
        t_map = {t._tbl: t for t in doc.tables}
        
        # Traverse paragraphs and tables in order
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph
                paragraph = p_map.get(element)
                if paragraph:
                    style_str = []
                    
                    # 1. Text Alignment
                    align = paragraph.alignment
                    if align == WD_ALIGN_PARAGRAPH.CENTER:
                        style_str.append("text-align: center;")
                    elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                        style_str.append("text-align: right;")
                    elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        style_str.append("text-align: justify;")
                    else:
                        style_str.append("text-align: left;")
                        
                    # 2. Spacing Before / After & Indents
                    space_before = paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else 0
                    space_after = paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 6
                    left_indent = paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0
                    right_indent = paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else 0
                    
                    style_str.append(f"margin: 0; margin-top: {space_before}pt; margin-bottom: {space_after}pt; margin-left: {left_indent}pt; margin-right: {right_indent}pt;")
                    
                    # 3. First Line Indent (paragraph tab spacing)
                    if paragraph.paragraph_format.first_line_indent:
                        fl_val = paragraph.paragraph_format.first_line_indent.pt
                        style_str.append(f"text-indent: {fl_val}pt;")
                        
                    # 4. Line Spacing
                    line_spacing = paragraph.paragraph_format.line_spacing
                    if line_spacing:
                        if isinstance(line_spacing, float):
                            style_str.append(f"line-height: {line_spacing};")
                        else:
                            style_str.append(f"line-height: {line_spacing.pt}pt;")
                    else:
                        style_str.append("line-height: 1.15;")
                        
                    # Ensure legibility in dark theme by adding explicit color
                    style_str.append("color: #1a202c !important;")
                    
                    # Write <p> tag
                    html.append(f"<p style='{' '.join(style_str)}'>")
                    for run in paragraph.runs:
                        r_style = []
                        if run.bold:
                            r_style.append("font-weight: bold;")
                        if run.italic:
                            r_style.append("font-style: italic;")
                        if run.underline:
                            r_style.append("text-decoration: underline;")
                        if run.font.size:
                            r_style.append(f"font-size: {run.font.size.pt}pt;")
                        if run.font.name:
                            r_style.append(f"font-family: '{run.font.name}', 'Times New Roman', serif;")
                        if run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            r_style.append(f"color: #{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x} !important;")
                        else:
                            r_style.append("color: #1a202c !important;")
                            
                        text_html = run.text.replace("\n", "<br>")
                        html.append(f"<span style='{' '.join(r_style)}'>{text_html}</span>")
                    
                    # Handle empty paragraphs (preserve spacing)
                    if not paragraph.runs:
                        html.append("&nbsp;")
                        
                    html.append("</p>")
            elif element.tag.endswith('tbl'):
                # It's a table
                table = t_map.get(element)
                if table:
                    html.append("<table style='border-collapse: collapse; width: 100%; margin: 15px 0; border: 1px solid #cbd5e1; color: #1a202c !important;'>")
                    for row in table.rows:
                        html.append("<tr>")
                        for cell in row.cells:
                            html.append("<td style='border: 1px solid #cbd5e1; padding: 8px; vertical-align: top; color: #1a202c !important;'>")
                            for cell_p in cell.paragraphs:
                                cell_p_style = []
                                align = cell_p.alignment
                                if align == WD_ALIGN_PARAGRAPH.CENTER:
                                    cell_p_style.append("text-align: center;")
                                elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                                    cell_p_style.append("text-align: right;")
                                elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                    cell_p_style.append("text-align: justify;")
                                else:
                                    cell_p_style.append("text-align: left;")
                                    
                                # Cell paragraph spacing should be compact
                                space_before = cell_p.paragraph_format.space_before.pt if cell_p.paragraph_format.space_before else 0
                                space_after = cell_p.paragraph_format.space_after.pt if cell_p.paragraph_format.space_after else 2
                                left_indent = cell_p.paragraph_format.left_indent.pt if cell_p.paragraph_format.left_indent else 0
                                right_indent = cell_p.paragraph_format.right_indent.pt if cell_p.paragraph_format.right_indent else 0
                                
                                cell_p_style.append(f"margin: 0; margin-top: {space_before}pt; margin-bottom: {space_after}pt; margin-left: {left_indent}pt; margin-right: {right_indent}pt;")
                                
                                if cell_p.paragraph_format.first_line_indent:
                                    fl_val = cell_p.paragraph_format.first_line_indent.pt
                                    cell_p_style.append(f"text-indent: {fl_val}pt;")
                                    
                                line_spacing = cell_p.paragraph_format.line_spacing
                                if line_spacing:
                                    if isinstance(line_spacing, float):
                                        cell_p_style.append(f"line-height: {line_spacing};")
                                    else:
                                        cell_p_style.append(f"line-height: {line_spacing.pt}pt;")
                                else:
                                    cell_p_style.append("line-height: 1.15;")
                                    
                                cell_p_style.append("color: #1a202c !important;")
                                html.append(f"<p style='{' '.join(cell_p_style)}'>")
                                for run in cell_p.runs:
                                    r_style = []
                                    if run.bold:
                                        r_style.append("font-weight: bold;")
                                    if run.italic:
                                        r_style.append("font-style: italic;")
                                    if run.underline:
                                        r_style.append("text-decoration: underline;")
                                    if run.font.size:
                                        r_style.append(f"font-size: {run.font.size.pt}pt;")
                                    if run.font.color and run.font.color.rgb:
                                        rgb = run.font.color.rgb
                                        r_style.append(f"color: #{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x} !important;")
                                    else:
                                        r_style.append("color: #1a202c !important;")
                                    text_html = run.text.replace("\n", "<br>")
                                    html.append(f"<span style='{' '.join(r_style)}'>{text_html}</span>")
                                if not cell_p.runs:
                                    html.append("&nbsp;")
                                html.append("</p>")
                            html.append("</td>")
                        html.append("</tr>")
                    html.append("</table>")
                    
        html.append("</div>")
        
        if os.path.exists(temp_out):
            try: os.remove(temp_out)
            except Exception: pass
            
        return "\n".join(html)
    except Exception as e:
        return f"<div style='color: red !important; font-weight: bold;'>Помилка попереднього перегляду Word: {e}</div>"

def generate_xlsx_preview(template_path, variables, config_path=None):
    """Generates a temporary Excel document and extracts its content for quick preview in high-fidelity HTML."""
    import tempfile
    import openpyxl
    from openpyxl.styles import PatternFill
    from _templates_machine_ import process_excel
    import html as py_html
    
    # Resolve relative template paths relative to the current Excel config directory if needed
    from _templates_machine_ import resolve_path
    c_path = config_path or st.session_state.get("editor_config_path")
    cfg_dir = os.path.dirname(os.path.abspath(c_path)) if c_path else os.getcwd()
    actual_path = resolve_path(cfg_dir, template_path)
            
    if not os.path.exists(actual_path):
        return f"<div style='color: red; font-weight: bold;'>Шаблон не знайдено за шляхом: {template_path}</div>"
        
    temp_dir = tempfile.gettempdir()
    temp_out = os.path.join(temp_dir, "temp_preview.xlsx")
    
    try:
        process_excel(actual_path, temp_out, variables)
        # Load with data_only=False to retrieve Excel formula strings instead of blank/None cells
        wb = openpyxl.load_workbook(temp_out, data_only=False)
        
        html = []
        html.append("<div style='font-family: \"Segoe UI\", Tahoma, Geneva, Verdana, sans-serif; color: #1a202c; overflow-x: auto;'>")
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            html.append(f"<h4 style='color: #2b6cb0 !important; margin-top: 15px; border-bottom: 2px solid #2b6cb0 !important; padding-bottom: 5px;'>Аркуш: {sheet_name}</h4>")
            # Get dimensions and restrict preview size
            total_rows = ws.max_row
            total_cols = ws.max_column
            
            max_preview_rows = min(total_rows or 0, 100)
            max_preview_cols = min(total_cols or 0, 25)
            
            if (total_rows and total_rows > 100) or (total_cols and total_cols > 25):
                html.append(f"<div style='background-color: #fffaf0; border: 1px solid #feebc8; color: #c05621; padding: 10px; border-radius: 6px; margin-bottom: 10px; font-size: 12px;'>⚠️ Аркуш занадто великий ({total_rows} рядків, {total_cols} колонок). Показано лише перші 100 рядків та 25 колонок.</div>")
            
            html.append("<table style='border-collapse: collapse; border: 1px solid #cbd5e1; font-size: 13px; min-width: 100%;'>")
            
            rows = list(ws.iter_rows(max_row=max_preview_rows, max_col=max_preview_cols))
            if not rows:
                html.append("<tr><td style='padding: 10px; color: #718096 !important;'>Аркуш порожній</td></tr>")
                html.append("</table>")
                continue
                
            for row in rows:
                html.append("<tr>")
                for cell in row:
                    val = cell.value
                    if val is None:
                        val = ""
                        
                    val_str = str(val).strip()
                    if val_str.startswith('='):
                        # Excel formula formatting
                        escaped_formula = py_html.escape(val_str)
                        val_display = f"<span style='color: #2b6cb0 !important; font-family: \"JetBrains Mono\", monospace; font-size: 11px; font-weight: 500; background-color: #ebf8ff !important; border: 1px solid #bee3f8; border-radius: 4px; padding: 2px 6px; display: inline-block;'>fx {escaped_formula}</span>"
                    else:
                        val_display = py_html.escape(str(val) if val is not None else "")
                        
                    styles = []
                    if cell.alignment:
                        h_align = cell.alignment.horizontal or "left"
                        v_align = cell.alignment.vertical or "center"
                        styles.append(f"text-align: {h_align};")
                        styles.append(f"vertical-align: {v_align};")
                    else:
                        styles.append("text-align: left; vertical-align: center;")
                        
                    if cell.font:
                        if cell.font.bold:
                            styles.append("font-weight: bold;")
                        if cell.font.italic:
                            styles.append("font-style: italic;")
                        if cell.font.size:
                            styles.append(f"font-size: {cell.font.size}pt;")
                        if cell.font.color and cell.font.color.rgb:
                            rgb = str(cell.font.color.rgb)
                            if len(rgb) == 8:
                                rgb = rgb[2:]
                            if rgb != "00000000":
                                styles.append(f"color: #{rgb} !important;")
                                
                    if cell.fill and isinstance(cell.fill, PatternFill) and cell.fill.fill_type == "solid":
                        if cell.fill.fgColor and cell.fill.fgColor.rgb:
                            rgb = str(cell.fill.fgColor.rgb)
                            if len(rgb) == 8:
                                rgb = rgb[2:]
                            if rgb != "00000000":
                                styles.append(f"background-color: #{rgb} !important;")
                                
                    styles_str = " ".join(styles)
                    html.append(f"<td style='border: 1px solid #cbd5e1; padding: 6px 12px; min-width: 80px; {styles_str}'>{val_display}</td>")
                html.append("</tr>")
            html.append("</table>")
            
        html.append("</div>")
        
        if os.path.exists(temp_out):
            try: os.remove(temp_out)
            except Exception: pass
            
        return "\n".join(html)
    except Exception as e:
        return f"<div style='color: red; font-weight: bold;'>Помилка попереднього перегляду Excel: {e}</div>"

def extract_placeholders_with_context(template_path):
    """Extracts jinja2 placeholders like {{var}} from a docx or xlsx template, aggregating all unique contexts."""
    import re
    import openpyxl
    from docx import Document
    
    placeholders_list = {}
    pattern = re.compile(r"\{\{\s*([^{}\s]+)\s*\}\}")
    
    if template_path.endswith(".docx"):
        try:
            doc = Document(template_path)
            full_text = []
            for p in doc.paragraphs:
                full_text.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            combined_text = "\n".join(full_text)
            matches = pattern.finditer(combined_text)
            for match in matches:
                var_name = match.group(1)
                start_idx = max(0, match.start() - 100)
                end_idx = min(len(combined_text), match.end() + 100)
                context = combined_text[start_idx:end_idx].strip().replace("\n", " ")
                
                if var_name not in placeholders_list:
                    placeholders_list[var_name] = []
                if context not in placeholders_list[var_name]:
                    placeholders_list[var_name].append(context)
        except Exception as e:
            st.error(f"Помилка зчитування Word: {e}")
            
    elif template_path.endswith(".xlsx"):
        try:
            wb = openpyxl.load_workbook(template_path, data_only=False)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for r in ws.iter_rows():
                    for cell in r:
                        val = cell.value
                        if isinstance(val, str):
                            matches = pattern.finditer(val)
                            for match in matches:
                                var_name = match.group(1)
                                start_idx = max(0, match.start() - 100)
                                end_idx = min(len(val), match.end() + 100)
                                context = f"Комірка {cell.coordinate}: {val[start_idx:end_idx].strip()}"
                                
                                if var_name not in placeholders_list:
                                    placeholders_list[var_name] = []
                                if context not in placeholders_list[var_name]:
                                    placeholders_list[var_name].append(context)
        except Exception as e:
            st.error(f"Помилка зчитування Excel: {e}")
            
    # Merge all unique contexts into a single descriptive string per variable
    placeholders = {}
    for var, contexts in placeholders_list.items():
        placeholders[var] = " | ".join(contexts)
        
    return placeholders

def rename_placeholder_in_template(template_path, old_name, new_name):
    """Automatically renames placeholders inside the Word (.docx) or Excel (.xlsx) template file."""
    import os
    if not os.path.exists(template_path):
        return False
    ext = os.path.splitext(template_path)[1].lower()
    if ext == ".docx":
        from docx import Document
        try:
            doc = Document(template_path)
            from _templates_machine_ import consolidate_jinja_tags
            consolidate_jinja_tags(doc)
            old_tag = f"{{{{{old_name}}}}}"
            new_tag = f"{{{{{new_name}}}}}"
            modified = False
            for p in doc.paragraphs:
                if old_tag in p.text:
                    for run in p.runs:
                        if old_tag in run.text:
                            run.text = run.text.replace(old_tag, new_tag)
                            modified = True
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if old_tag in p.text:
                                for run in p.runs:
                                    if old_tag in run.text:
                                        run.text = run.text.replace(old_tag, new_tag)
                                        modified = True
            if modified:
                doc.save(template_path)
                return True
        except Exception:
            pass
    elif ext == ".xlsx":
        import openpyxl
        try:
            wb = openpyxl.load_workbook(template_path)
            old_tag = f"{{{{{old_name}}}}}"
            new_tag = f"{{{{{new_name}}}}}"
            modified = False
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str) and old_tag in cell.value:
                            cell.value = cell.value.replace(old_tag, new_tag)
                            modified = True
            if modified:
                wb.save(template_path)
                return True
        except Exception:
            pass
    return False

configs, templates = scan_workspace()

# MAIN PAGE HEADER
col_title, col_theme = st.columns([11, 1])
with col_title:
    st.markdown('<div class="main-title">🚀 Панель керування TemplateMachine</div>', unsafe_allow_html=True)
    st.markdown('<div class="subtitle">Універсальний комбайн для автоматизації документів та аналізу архівів</div>', unsafe_allow_html=True)
with col_theme:
    st.write(" ") # spacer to push down button slightly
    theme_emoji = "🌙" if theme == "light" else "☀️"
    if st.button(theme_emoji, key="theme_toggle_btn", use_container_width=True):
        st.session_state["theme"] = "dark" if theme == "light" else "light"
        st.session_state["theme_changed"] = True
        save_persistent_state()
        st.rerun()

st.markdown("---")

# Initialize session state for active view retention
views_list = [
    "Менеджер Проектів",
    "✈️ Аналіз та Створення Шаблонів",
    "📝 Редактор Excel Конфігів",
    "⚡ Генерація Документів",
    "📖 Повна Довідка"
]

if "current_view" not in st.session_state:
    st.session_state["current_view"] = views_list[0]

if st.session_state["current_view"] not in views_list:
    st.session_state["current_view"] = views_list[0]

if "app_view_selector" not in st.session_state:
    st.session_state["app_view_selector"] = st.session_state["current_view"]

def on_view_change(new_view):
    if not new_view:
        return
    prev_view = st.session_state.get("current_view", views_list[0])
    if new_view != prev_view:
        # Save leaving view's state first
        save_persistent_state()
        
        # 1. Leaving views: Save metadata
        if prev_view == "📝 Редактор Excel Конфігів":
            if st.session_state.get("editor_config_path"):
                st.session_state["last_opened_config"] = st.session_state["editor_config_path"]
            if st.session_state.get("editor_template_path"):
                st.session_state["last_opened_template"] = st.session_state["editor_template_path"]
                
        # 2. Entering views: Restore cached paths
        if new_view == "📝 Редактор Excel Конфігів":
            st.session_state["editor_config_path"] = st.session_state.get("last_opened_config") or ""
            st.session_state["editor_template_path"] = st.session_state.get("last_opened_template") or ""
            st.session_state["loaded_config_sheet"] = ""
            
            # Reset any data editor session state widgets to avoid cached checkboxes when switching views
            cfg_path = st.session_state.get("last_opened_config") or ""
            if cfg_path:
                clean_cfg_path = "".join([c if c.isalnum() else "_" for c in cfg_path])
                prefix_config = f"config_data_editor_{clean_cfg_path}_"
                prefix_headers = f"headers_data_editor_{clean_cfg_path}_"
                keys_to_del = [k for k in st.session_state.keys() if k.startswith(prefix_config) or k.startswith(prefix_headers)]
                for k in keys_to_del:
                    del st.session_state[k]
            
        elif new_view == "Менеджер Проектів":
            if not st.session_state.get("pm_folder_path"):
                st.session_state["pm_folder_path"] = st.session_state.get("last_opened_folder") or ""
            st.session_state["pm_cached_configs"] = {}
            
            # Sync the selected doc index if we came back from Config Editor where preview selection might have changed
            selected_doc = st.session_state.get("pm_selected_doc")
            last_idx = st.session_state.get("last_preview_row_idx")
            if selected_doc and last_idx is not None and selected_doc.get("row_idx") != last_idx:
                selected_doc["row_idx"] = last_idx
                st.session_state["pm_selected_doc"] = selected_doc
                st.session_state["pm_editing_vars"] = None
                clear_pm_input_keys()
            sync_pm_editing_vars()
            
        elif new_view == "⚡ Генерація Документів":
            st.session_state["gen_config_path"] = st.session_state.get("last_opened_config") or ""
                
        st.session_state["current_view"] = new_view
        st.session_state["app_view_selector"] = new_view
        
        # Save entering view's state updates
        save_persistent_state()
        
        # Load persistent state now to populate session state with values for the entering view
        load_persistent_state()

def handle_view_selectbox_change():
    new_view = st.session_state.get("app_view_selector")
    if new_view:
        on_view_change(new_view)

selected_view = st.selectbox(
    "Оберіть розділ роботи:",
    views_list,
    key="app_view_selector",
    on_change=handle_view_selectbox_change
)

st.session_state["current_view"] = selected_view

st.markdown(" ")

# ----------------------------------------------------
# VIEW 0: PROJECT MANAGER & DOCUMENT TREE
# ----------------------------------------------------
if selected_view == "Менеджер Проектів":
    st.header("📁 Менеджер Проектів та Віртуальне Дерево Документів")
    st.write("Вкажіть шлях до папки, щоб автоматично просканувати всі конфіги та побудувати дерево віртуальних документів.")
    
    # 1. Folder picker / text input
    if "pm_folder_path" not in st.session_state:
        st.session_state["pm_folder_path"] = ""
        
    col_p1, col_p2 = st.columns([3, 1])
    with col_p1:
        folder_input = st.text_input(
            "📁 Шлях до папки з конфігами та шаблонами:",
            placeholder="Введіть або оберіть шлях до папки (наприклад, example)...",
            key="pm_folder_path",
            on_change=save_persistent_state
        )
    with col_p2:
        st.write(" ")
        st.write(" ")
        def select_pm_folder():
            selected = open_folder_picker("Оберіть папку з конфігами")
            if selected:
                st.session_state["pm_folder_path"] = selected
                st.session_state["last_opened_folder"] = selected
                save_persistent_state()
        st.button("📁 Обрати папку", key="btn_pm_folder", on_click=select_pm_folder)
        
    pm_path = st.session_state["pm_folder_path"]
    if pm_path and pm_path != st.session_state.get("last_opened_folder"):
        st.session_state["last_opened_folder"] = pm_path
        save_persistent_state()
    
    if not pm_path:
        st.info("Будь ласка, оберіть або введіть шлях до папки для сканування.")
    elif not os.path.exists(pm_path):
        st.error(f"Вказаний шлях '{pm_path}' не існує!")
    else:
        # Scan and build tree
        with st.spinner("⏳ Сканування папки..."):
            config_files = scan_recursive_configs(pm_path)
            
        if not config_files:
            st.warning("У вказаній папці не знайдено жодного Excel-конфігу.")
        else:
            # Build tree dictionary
            dir_tree = build_dir_tree(config_files, pm_path)
            
            # Setup split layout: Tree on the left, Preview & Edit on the right
            col_tree, col_content = st.columns([5, 7])
            
            with col_tree:
                only_docs = st.checkbox(
                    "Лише документи",
                    key="pm_only_docs",
                    on_change=save_persistent_state
                )
                
                st.subheader("🌳 Віртуальне дерево документів")
                
                # Recursive tree renderer with auto-expansion of current document path
                def contains_config_path(tree_dict, target_path):
                    if not target_path:
                        return False
                    try:
                        target_path = os.path.abspath(target_path).lower()
                    except Exception:
                        return False
                    
                    def check_node(val):
                        if isinstance(val, dict):
                            return any(check_node(v) for v in val.values())
                        elif isinstance(val, str):
                            try:
                                return os.path.abspath(val).lower() == target_path
                            except Exception:
                                return False
                        return False
                        
                    return check_node(tree_dict)

                def render_tree_node(node_name, node_value, depth=0, path_prefix=""):
                    selected_doc = st.session_state.get("pm_selected_doc")
                    if selected_doc:
                        config_path_clean = "".join([c if c.isalnum() else "_" for c in selected_doc.get('config_path', '')])
                        selected_doc_id = f"{config_path_clean}_{selected_doc.get('sheet_name')}_{selected_doc.get('row_idx')}"
                    else:
                        selected_doc_id = "none"
                    target_path = selected_doc.get("config_path") if selected_doc else None
                    
                    if isinstance(node_value, dict):
                        # Folder node: expand if root level or if it contains the selected config file path
                        should_expand = (depth == 0) or (target_path and contains_config_path(node_value, target_path))
                        exp_key = f"pm_folder_exp_{path_prefix}_{node_name}_{selected_doc_id}"
                        with st.expander("📁 " + node_name, expanded=should_expand, key=exp_key):
                            for name, val in node_value.items():
                                render_tree_node(name, val, depth + 1, path_prefix=f"{path_prefix}_{node_name}")
                    else:
                        # Config file node: expand if it is the current config of the selected document
                        config_path = node_value
                        config_name = node_name
                        is_current_config = False
                        if selected_doc and selected_doc.get("config_path"):
                            try:
                                is_current_config = os.path.abspath(selected_doc["config_path"]).lower() == os.path.abspath(config_path).lower()
                            except Exception:
                                pass
                        
                        clean_cfg = "".join([c if c.isalnum() else "_" for c in config_path])
                        exp_cfg_key = f"pm_cfg_exp_{path_prefix}_{clean_cfg}_{selected_doc_id}"
                        with st.expander("📊 " + config_name, expanded=is_current_config, key=exp_cfg_key):
                            sheets_data = get_cached_config(config_path)
                            if not sheets_data:
                                st.caption("Не вдалося завантажити або порожній конфіг")
                                return
                            
                            for sheet_name, info in sheets_data.items():
                                rows = info["rows"]
                                template_path = info["template_path"]
                                name_pattern = info["name_pattern"]
                                
                                # Sheet expander: expand if it is the current sheet of the selected document
                                is_current_sheet = is_current_config and selected_doc.get("sheet_name") == sheet_name
                                exp_sheet_key = f"pm_sheet_exp_{path_prefix}_{clean_cfg}_{sheet_name}_{selected_doc_id}"
                                with st.expander(f"📋 {sheet_name} ({len(rows)} док.)", expanded=bool(is_current_sheet), key=exp_sheet_key):
                                    if not rows:
                                        st.caption("Немає даних для генерації")
                                        continue
                                    
                                    for idx, row in enumerate(rows):
                                        is_selected = (
                                            st.session_state.get("pm_selected_doc", {}).get("config_path") == config_path and
                                            st.session_state.get("pm_selected_doc", {}).get("sheet_name") == sheet_name and
                                            st.session_state.get("pm_selected_doc", {}).get("row_idx") == idx
                                        )
                                        
                                        current_row_vars = row
                                        if is_selected and st.session_state.get("pm_editing_vars") is not None:
                                            current_row_vars = st.session_state["pm_editing_vars"]
                                            
                                        doc_name = resolve_virtual_doc_name(name_pattern, current_row_vars, template_path)
                                        if not doc_name.strip():
                                            doc_name = f"document_{idx + 5}"
                                            
                                        button_type = "primary" if is_selected else "secondary"
                                        
                                        if st.button(
                                            f"📄 {doc_name}",
                                            key=f"pm_btn_{config_path}_{sheet_name}_{idx}",
                                            use_container_width=True,
                                            type=button_type
                                        ):
                                            st.session_state["pm_selected_doc"] = {
                                                "config_path": config_path,
                                                "sheet_name": sheet_name,
                                                "row_idx": idx,
                                                "doc_name": doc_name,
                                                "template_path": template_path,
                                                "name_pattern": name_pattern
                                            }
                                            st.session_state["last_opened_config"] = config_path
                                            st.session_state["last_opened_folder"] = os.path.dirname(os.path.abspath(config_path))
                                            st.session_state["last_opened_template"] = template_path
                                            st.session_state["pending_download"] = None
                                            clear_pm_input_keys()
                                            save_persistent_state()
                                            st.session_state["pm_editing_vars"] = None
                                            save_persistent_state()
                                            st.rerun()

                if only_docs:
                    docs_tree = build_docs_only_tree(config_files, pm_path)
                    if not docs_tree:
                        st.info("Документів не знайдено.")
                    else:
                        render_docs_only_tree(docs_tree, 0)
                else:
                    # Render tree roots
                    for name, val in dir_tree.items():
                        render_tree_node(name, val, 0)
                    
            with col_content:
                selected_doc = st.session_state.get("pm_selected_doc")
                if not selected_doc:
                    st.info("👈 Оберіть віртуальний документ у дереві ліворуч, щоб почати перегляд та редагування.")
                else:
                    config_path = selected_doc["config_path"]
                    sheet_name = selected_doc["sheet_name"]
                    row_idx = selected_doc["row_idx"]
                    template_path = selected_doc["template_path"]
                    name_pattern = selected_doc["name_pattern"]
                    
                    st.subheader("🛠️ Перегляд та редагування")
                    
                    # Display metadata paths
                    st.caption(f"**Конфіг:** `{os.path.basename(config_path)}` | **Аркуш:** `{sheet_name}` | **Шаблон:** `{template_path}`")
                    
                    sheets_data = get_cached_config(config_path)
                    if not sheets_data or sheet_name not in sheets_data:
                        st.error("Помилка завантаження даних документа.")
                    else:
                        sheet_info = sheets_data[sheet_name]
                        headers = sheet_info["headers"]
                        row_vars = sheet_info["rows"][row_idx]
                        
                        # Initialize editing variables
                        if st.session_state.get("pm_editing_vars") is None:
                            st.session_state["pm_editing_vars"] = dict(row_vars)
                            
                        edited_vars = st.session_state["pm_editing_vars"]
                        
                        # High fidelity preview (rendered first)
                        st.markdown(" ")
                        st.markdown("##### 🔍 Попередній перегляд документа")
                        
                        # Resolve relative template path
                        from _templates_machine_ import resolve_path
                        cfg_dir = os.path.dirname(os.path.abspath(config_path))
                        actual_template_path = resolve_path(cfg_dir, template_path)
                            
                        if not os.path.exists(actual_template_path):
                            st.warning(f"Шаблон не знайдено за шляхом: {template_path} (враховуючи відносність до конфігу)")
                        else:
                            ext = os.path.splitext(actual_template_path)[1].lower()
                            with st.spinner("⏳ Генерація прев'ю..."):
                                if ext == ".docx":
                                    preview_html = generate_docx_preview(actual_template_path, edited_vars, config_path=config_path)
                                elif ext == ".xlsx":
                                    preview_html = generate_xlsx_preview(actual_template_path, edited_vars, config_path=config_path)
                                else:
                                    preview_html = f"<div style='color: #e53e3e;'>Непідтримуваний тип шаблону: {ext}</div>"
                                    
                            st.markdown(
                                f"""
                                <div class="document-preview-container" style="border: 2px solid #3182ce; border-radius: 8px; padding: 25px; background-color: #ffffff; max-height: 500px; overflow-y: auto; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); border-left: 8px solid #3182ce;">
                                    {preview_html}
                                </div>
                                """,
                                unsafe_allow_html=True
                            )
                            
                        # Show dynamic output name preview (rendered second)
                        current_resolved_name = resolve_virtual_doc_name(name_pattern, edited_vars, template_path)
                        st.markdown(" ")
                        st.markdown(f"##### 📄 Вихідне ім'я: `{current_resolved_name}`")
                        
                        # Render editing inputs
                        st.markdown("##### ✏️ Змінні цього документа:")
                        var_cols = st.columns(3)
                        for idx, h in enumerate(headers):
                            if not h:
                                continue
                            with var_cols[idx % 3]:
                                val = edited_vars.get(h, "")
                                new_val = st.text_input(f"{h}", value=val, key=f"pm_input_{config_path}_{sheet_name}_{row_idx}_{h}")
                                if new_val != val:
                                    edited_vars[h] = new_val
                                    save_persistent_state()
                                    st.rerun()
                                    
                        # Buttons row
                        st.markdown(" ")
                        btn_col1, btn_col2, btn_col3 = st.columns(3)
                        
                        with btn_col1:
                            if st.button("💾 Зберегти зміни в конфіг", type="primary", use_container_width=True):
                                # Load fresh config data to write
                                full_data = load_excel_config(config_path)
                                if full_data and sheet_name in full_data:
                                    # Update target row
                                    full_data[sheet_name]["rows"][row_idx] = edited_vars
                                    
                                    # Save to file
                                    success = save_excel_config(
                                        config_path,
                                        sheet_name,
                                        full_data[sheet_name]["template_path"],
                                        full_data[sheet_name]["name_pattern"],
                                        full_data[sheet_name]["headers"],
                                        full_data[sheet_name]["rows"]
                                    )
                                    if success:
                                        # Clear cache for this config
                                        if config_path in st.session_state["pm_cached_configs"]:
                                            del st.session_state["pm_cached_configs"][config_path]
                                        st.toast("🎉 Зміни успішно збережено в Excel!", icon="💾")
                                        time.sleep(0.5)
                                        st.rerun()
                                        
                        with btn_col2:
                            if st.button("📄 Зберегти документ", use_container_width=True):
                                # Resolve absolute template path
                                from _templates_machine_ import resolve_path
                                cfg_dir = os.path.dirname(os.path.abspath(config_path))
                                actual_t_path = resolve_path(cfg_dir, template_path)
                                    
                                save_generated_document_dialog(actual_t_path, edited_vars, config_path, name_pattern=name_pattern)
                                
                        with btn_col3:
                            def go_to_config_editor():
                                clean_cfg_path = "".join([c if c.isalnum() else "_" for c in config_path])
                                config_key = f"config_data_editor_{clean_cfg_path}_{sheet_name}"
                                headers_key = f"headers_data_editor_{clean_cfg_path}_{sheet_name}"
                                if config_key in st.session_state:
                                    del st.session_state[config_key]
                                if headers_key in st.session_state:
                                    del st.session_state[headers_key]
                                st.session_state["app_view_selector"] = "📝 Редактор Excel Конфігів"
                                st.session_state["current_view"] = "📝 Редактор Excel Конфігів"
                                st.session_state["editor_config_path"] = config_path
                                st.session_state["editor_selected_sheet"] = sheet_name
                                st.session_state["loaded_config_sheet"] = ""  # Force reload headers and data
                                st.session_state["last_preview_row_idx"] = row_idx
                                # Update last opened values immediately on transition
                                st.session_state["last_opened_config"] = config_path
                                st.session_state["last_opened_folder"] = os.path.dirname(os.path.abspath(config_path))
                                if template_path:
                                    st.session_state["last_opened_template"] = template_path
                                save_persistent_state()
                            st.button("✏️ Перейти до редактора конфігів", use_container_width=True, on_click=go_to_config_editor)
                            
                        pending_download = st.session_state.get("pending_download")
                        if pending_download and isinstance(pending_download, dict):
                            st.write(" ")
                            st.download_button(
                                label="⬇️ Завантажити згенерований документ через браузер",
                                data=pending_download["bytes"],
                                file_name=pending_download["name"],
                                mime=pending_download["mime"],
                                key="pm_download_fallback_btn",
                                use_container_width=True
                            )

# ----------------------------------------------------
# VIEW 1: ARCHIVE ANALYSIS & TEMPLATE CREATION
# ----------------------------------------------------
if selected_view == "✈️ Аналіз та Створення Шаблонів":
    st.header("🔍 Режими аналізу та розпізнавання шаблонів")
    st.write("Скрипт проведе інтелектуальне попарне або групове порівняння документів, виділить змінні і створить конфігураційний файл.")
    
    # Native Streamlit click callbacks to update widget states safely before instantiation
    def select_auto_folder():
        selected = open_folder_picker("Оберіть папку з архівом для аналізу")
        if selected:
            st.session_state["txt_auto_folder"] = selected
            save_persistent_state()
            
    def select_batch_sample():
        selected = open_file_picker()
        if selected:
            st.session_state["txt_batch_sample"] = selected
            save_persistent_state()
            
    def select_batch_folder():
        selected = open_folder_picker("Оберіть папку для пакетного аналізу")
        if selected:
            st.session_state["txt_batch_folder"] = selected
            save_persistent_state()
            
    def select_pair_file1():
        selected = open_file_picker()
        if selected:
            st.session_state["txt_pair_file1"] = selected
            save_persistent_state()
            
    def select_pair_file2():
        selected = open_file_picker()
        if selected:
            st.session_state["txt_pair_file2"] = selected
            save_persistent_state()
    
    analysis_modes = [
        "✈️ Повний автопілот (Сканування папки та автоматичне розбиття на групи)",
        "🔍 Пакетний аналіз (Один файл-зразок + папка з іншими файлами)",
        "⚖️ Попарне порівняння (Точне порівняння двох конкретних файлів)"
    ]
    if "analysis_mode" not in st.session_state:
        st.session_state["analysis_mode"] = analysis_modes[0]
    if st.session_state["analysis_mode"] not in analysis_modes:
        matched = False
        for m in analysis_modes:
            if st.session_state["analysis_mode"] and (st.session_state["analysis_mode"] in m or m in st.session_state["analysis_mode"]):
                st.session_state["analysis_mode"] = m
                matched = True
                break
        if not matched:
            st.session_state["analysis_mode"] = analysis_modes[0]

    mode = st.radio(
        "Оберіть режим аналізу:",
        analysis_modes,
        key="analysis_mode",
        on_change=save_persistent_state
    )
    
    # --- ALWAYS-VISIBLE ANALYSIS OUTPUT DIRECTORY SELECTION ---
    st.markdown("##### 📁 Папка для збереження результатів (Необов'язково)")
    st.caption("Якщо вказано, створені шаблони (`template_*`), Excel-конфіги та `.bat` файли будуть перенесені в цю папку. У безвіконному (headless) режимі введіть шлях вручну.")
    
    col_ao1, col_ao2 = st.columns([3, 1])
    with col_ao1:
        a_out_dir = st.text_input(
            "📁 Шлях до папки результатів:",
            placeholder="Наприклад: C:/MyTemplates (залиште порожнім для збереження за замовчуванням)",
            key="analysis_output_dir",
            on_change=save_persistent_state
        )
    with col_ao2:
        st.write(" ")
        st.write(" ")
        def select_analysis_output_dir():
            selected = open_folder_picker("Оберіть папку для збереження результатів")
            if selected:
                st.session_state["analysis_output_dir"] = selected
                save_persistent_state()
        st.button("📁 Обрати папку", key="btn_analysis_output_dir", on_click=select_analysis_output_dir)
        
    st.markdown("---")
    
    if "Повний автопілот" in mode:
        st.subheader("✈️ Режим 1: Повний автопілот")
        st.write("Аналізує вказану папку, групує схожі за структурою документи, створює для кожної групи окремий шаблон та зводить все в єдиний мульти-конфіг `Auto_Config.xlsx`.")
        
        col1, col2 = st.columns([3, 1])
        with col1:
            folder_path = st.text_input(
                "Шлях до папки з архівом документів:",
                placeholder="Наприклад: example/docs",
                key="txt_auto_folder",
                on_change=save_persistent_state
            )
        with col2:
            st.write(" ")
            st.write(" ")
            st.button("📁 Провідник Windows", key="btn_auto_folder", on_click=select_auto_folder)
                    
        if st.button("🚀 Запустити повний автопілот", type="primary"):
            if not folder_path:
                st.error("Будь ласка, вкажіть шлях до папки!")
            elif not os.path.exists(folder_path):
                st.error(f"Вказана папка '{folder_path}' не існує!")
            else:
                ret_code, _ = run_subprocess_and_stream([folder_path])
                if ret_code == 0 and st.session_state.get("analysis_output_dir"):
                    move_autopilot_outputs(st.session_state["analysis_output_dir"])
                configs, templates = scan_workspace()
                st.rerun()
                
    elif "Пакетний аналіз" in mode:
        st.subheader("🔍 Режим 2: Пакетний аналіз")
        st.write("Порівнює один обраний файл-зразок зі схожими файлами у папці та формує шаблон і конфігурацію.")
        
        # Row for Sample file
        col1, col2 = st.columns([3, 1])
        with col1:
            sample_file = st.text_input(
                "Шлях до файлу-зразка (.docx або .xlsx):",
                placeholder="Наприклад: example/w.docx",
                key="txt_batch_sample",
                on_change=save_persistent_state
            )
        with col2:
            st.write(" ")
            st.write(" ")
            st.button("📄 Обрати зразок", key="btn_batch_sample", on_click=select_batch_sample)
                    
        # Row for Folder
        col1, col2 = st.columns([3, 1])
        with col1:
            folder_path = st.text_input(
                "Шлях до папки порівняння:",
                placeholder="Наприклад: example",
                key="txt_batch_folder",
                on_change=save_persistent_state
            )
        with col2:
            st.write(" ")
            st.write(" ")
            st.button("📁 Обрати папку", key="btn_batch_folder", on_click=select_batch_folder)
                    
        if st.button("🚀 Запустити пакетний аналіз", type="primary"):
            if not sample_file or not folder_path:
                st.error("Будь ласка, вкажіть і файл-зразок, і папку порівняння!")
            elif not os.path.exists(sample_file):
                st.error(f"Файл-зразок '{sample_file}' не знайдено!")
            elif not os.path.exists(folder_path):
                st.error(f"Папку порівняння '{folder_path}' не знайдено!")
            else:
                ret_code, _ = run_subprocess_and_stream([sample_file, folder_path])
                if ret_code == 0 and st.session_state.get("analysis_output_dir"):
                    move_batch_outputs(sample_file, st.session_state["analysis_output_dir"])
                configs, templates = scan_workspace()
                st.rerun()
                
    elif "Попарне порівняння" in mode:
        st.subheader("⚖️ Режим 3: Попарне порівняння")
        st.write("Знаходить розбіжності між двома файлами, виділяє змінні та створює шаблон та індивідуальний конфіг.")
        
        # File 1
        col1, col2 = st.columns([3, 1])
        with col1:
            file_1 = st.text_input(
                "Шлях до Першого файлу:",
                placeholder="Наприклад: example/w1.docx",
                key="txt_pair_file1",
                on_change=save_persistent_state
            )
        with col2:
            st.write(" ")
            st.write(" ")
            st.button("📄 Обрати файл 1", key="btn_pair_file1", on_click=select_pair_file1)
                    
        # File 2
        col1, col2 = st.columns([3, 1])
        with col1:
            file_2 = st.text_input(
                "Шлях до Другого файлу:",
                placeholder="Наприклад: example/w2.docx",
                key="txt_pair_file2",
                on_change=save_persistent_state
            )
        with col2:
            st.write(" ")
            st.write(" ")
            st.button("📄 Обрати файл 2", key="btn_pair_file2", on_click=select_pair_file2)
                    
        if st.button("🚀 Запустити попарне порівняння", type="primary"):
            if not file_1 or not file_2:
                st.error("Будь ласка, вкажіть обидва файли для порівняння!")
            elif not os.path.exists(file_1) or not os.path.exists(file_2):
                st.error("Один або обидва вказані файли не знайдені!")
            else:
                ret_code, _ = run_subprocess_and_stream([file_1, file_2])
                if ret_code == 0 and st.session_state.get("analysis_output_dir"):
                    move_pairwise_outputs(file_1, st.session_state["analysis_output_dir"])
                configs, templates = scan_workspace()
                st.rerun()

    # Persistent log viewer at the bottom of analysis tab
    show_last_operation_logs()

# ----------------------------------------------------
# VIEW 2: INTERACTIVE EXCEL CONFIG EDITOR
# ----------------------------------------------------
elif selected_view == "📝 Редактор Excel Конфігів":
    st.header("📝 Інтерактивний редактор файлів конфігурації")
    st.write("Оберіть будь-який створений конфіг-файл, щоб відредагувати параметри шаблону, імені або змінити дані змінних та аркушів.")
    
    col_c1, col_c2 = st.columns([3, 1])
    with col_c1:
        selected_config = st.text_input(
            "Шлях до файлу конфігурації для редагування:",
            placeholder="Оберіть Excel-файл конфігурації...",
            key="editor_config_path",
            on_change=save_persistent_state
        )
    with col_c2:
        st.write(" ")
        st.write(" ")
        def select_editor_config():
            selected = open_file_picker(filetypes=[("Excel конфігурації", "*.xlsx"), ("Усі файли", "*.*")])
            if selected:
                st.session_state["editor_config_path"] = selected
                st.session_state["last_opened_config"] = selected
                st.session_state["last_opened_folder"] = os.path.dirname(os.path.abspath(selected))
                save_persistent_state()
        st.button("📁 Обрати конфіг", key="btn_editor_config", on_click=select_editor_config)
        
    cfg_path = st.session_state["editor_config_path"]
    if cfg_path and cfg_path != st.session_state.get("last_opened_config"):
        st.session_state["last_opened_config"] = cfg_path
        st.session_state["last_opened_folder"] = os.path.dirname(os.path.abspath(cfg_path))
        save_persistent_state()
    if not cfg_path:
        st.info("Будь ласка, оберіть файл конфігурації Excel (`*.xlsx`) для початку роботи!")
    elif not os.path.exists(cfg_path):
        st.error(f"Вказаний файл конфігурації '{cfg_path}' не знайдено!")
    else:
        sheets_data = load_excel_config(cfg_path)
        
        if sheets_data:
            sheet_names = list(sheets_data.keys())
            
            if "editor_selected_sheet" not in st.session_state or st.session_state["editor_selected_sheet"] not in sheet_names:
                st.session_state["editor_selected_sheet"] = sheet_names[0] if sheet_names else ""
                
            selected_sheet = st.selectbox(
                "Оберіть аркуш конфігурації:",
                sheet_names,
                key="editor_selected_sheet",
                on_change=save_persistent_state
            )
                
            # --- SHEET CRUD MANAGEMENT ---
            with st.expander("📂 Керування аркушами (Sheets Operations)", expanded=False):
                col_s1, col_s2, col_s3 = st.columns(3)
                
                with col_s1:
                    st.markdown("##### ➕ Створити новий аркуш")
                    new_s_name = st.text_input("Назва нового аркуша:", key="txt_new_sheet_name")
                    if st.button("➕ Створити аркуш", key="btn_create_sheet"):
                        if not new_s_name:
                            st.error("Введіть назву аркуша!")
                        else:
                            clean_s_name = re.sub(r'[\\/*?:\[\]]', "", new_s_name)[:31].strip()
                            if create_new_sheet(cfg_path, clean_s_name):
                                st.session_state["editor_selected_sheet"] = clean_s_name
                                save_persistent_state()
                                st.success(f"Аркуш '{clean_s_name}' успішно створено!")
                                time.sleep(1)
                                st.rerun()
                                
                with col_s2:
                    st.markdown("##### ✏️ Перейменувати поточний аркуш")
                    new_rename_name = st.text_input("Нова назва аркуша:", key="txt_rename_sheet_name")
                    if st.button("✏️ Перейменувати аркуш", key="btn_rename_sheet"):
                        if not new_rename_name:
                            st.error("Введіть нову назву!")
                        else:
                            clean_rename_name = re.sub(r'[\\/*?:\[\]]', "", new_rename_name)[:31].strip()
                            if rename_sheet(cfg_path, selected_sheet, clean_rename_name):
                                st.session_state["editor_selected_sheet"] = clean_rename_name
                                save_persistent_state()
                                st.success(f"Аркуш перейменовано на '{clean_rename_name}'!")
                                time.sleep(1)
                                st.rerun()
                                
                with col_s3:
                    st.markdown("##### ❌ Видалити поточний аркуш")
                    st.warning("Ця дія є незворотною!")
                    confirm_delete = st.checkbox("Підтверджую видалення аркуша", key="chk_confirm_delete_sheet")
                    if st.button("❌ Видалити аркуш", key="btn_delete_sheet"):
                        if not confirm_delete:
                            st.error("Будь ласка, підтвердіть видалення чекбоксом!")
                        else:
                            if delete_sheet(cfg_path, selected_sheet):
                                st.success(f"Аркуш '{selected_sheet}' видалено!")
                                remaining_sheets = [s for s in sheet_names if s != selected_sheet]
                                st.session_state["editor_selected_sheet"] = remaining_sheets[0] if remaining_sheets else ""
                                save_persistent_state()
                                time.sleep(1)
                                st.rerun()
                                
            # --- LOAD SHEET DATA ---
            sheet_info = sheets_data[selected_sheet]
            config_sheet_key = f"{cfg_path}_{selected_sheet}"
            
            # State synchronization on config/sheet change
            if st.session_state.get("loaded_config_sheet") != config_sheet_key:
                st.session_state["loaded_config_sheet"] = config_sheet_key
                st.session_state["current_sheet_headers"] = list(sheet_info["headers"])
                st.session_state["current_sheet_data"] = list(sheet_info["rows"])
                st.session_state["editor_template_path"] = sheet_info["template_path"]
                st.session_state["editor_name_pattern"] = sheet_info["name_pattern"]
                if sheet_info["template_path"]:
                    st.session_state["last_opened_template"] = sheet_info["template_path"]
                save_persistent_state()
                
            # --- METADATA (A1 & A2) ---
            st.markdown("### ⚙️ Налаштування генерації для обраного аркуша")
            col1, col2 = st.columns(2)
            
            with col1:
                col_t1, col_t2 = st.columns([3, 1])
                with col_t1:
                    t_path = st.text_input(
                        "📄 Шлях до шаблону (комірка A1):",
                        placeholder="Оберіть файл шаблону...",
                        key="editor_template_path",
                        on_change=save_persistent_state
                    )
                with col_t2:
                    st.write(" ")
                    st.write(" ")
                    def select_editor_template():
                        selected = open_file_picker(filetypes=[
                            ("Усі шаблони", "*.docx;*.xlsx"),
                            ("Word шаблони", "*.docx"),
                            ("Excel шаблони", "*.xlsx"),
                            ("Усі файли", "*.*")
                        ])
                        if selected:
                            st.session_state["editor_template_path"] = selected
                            save_persistent_state()
                    st.button("📁 Шаблон", key="btn_editor_template", on_click=select_editor_template)
                    
            with col2:
                n_pattern = st.text_input(
                    "✍️ Шаблон імені вихідних файлів (комірка A2):",
                    placeholder="Наприклад: output_{{YYYY}}.docx",
                    key="editor_name_pattern",
                    on_change=save_persistent_state
                )
                
            # Sync template path to last_opened_template immediately if changed
            curr_editor_t_path = st.session_state.get("editor_template_path", "")
            if curr_editor_t_path and curr_editor_t_path != st.session_state.get("last_opened_template", ""):
                st.session_state["last_opened_template"] = curr_editor_t_path
                save_persistent_state()
                
            # --- SLEEK ROW & COLUMN CONTROLS ---
            st.markdown("### 🛠️ Швидкі дії з рядками та стовпчиками (змінними)")
            col_ctrl1, col_ctrl2, col_ctrl3, col_ctrl4 = st.columns(4)
            
            with col_ctrl1:
                with st.popover("➕ Додати стовпчик (змінну)", width="stretch"):
                    new_col = st.text_input("Ім'я нової змінної (напр., client_name):", key="compact_new_col_input")
                    if st.button("➕ Додати стовпчик", key="btn_compact_add_col", width="stretch"):
                        if not new_col:
                            st.error("Введіть ім'я змінної!")
                        elif new_col in st.session_state["current_sheet_headers"]:
                            st.error("Такий стовпчик вже існує!")
                        else:
                            st.session_state["current_sheet_headers"].append(new_col)
                            for r in st.session_state["current_sheet_data"]:
                                r[new_col] = ""
                            save_persistent_state()
                            st.success(f"Стовпчик '{new_col}' додано!")
                            st.rerun()
                            
            with col_ctrl2:
                with st.popover("❌ Видалити стовпчик", width="stretch"):
                    if st.session_state["current_sheet_headers"]:
                        col_to_del = st.selectbox("Оберіть стовпчик для видалення:", st.session_state["current_sheet_headers"], key="compact_del_col_select")
                        confirm_col = st.checkbox("Підтверджую видалення стовпчика", key="compact_confirm_del_col")
                        if st.button("❌ Видалити стовпчик", key="btn_compact_del_col", type="primary", width="stretch"):
                            if not confirm_col:
                                st.error("Будь ласка, підтвердіть видалення!")
                            else:
                                st.session_state["current_sheet_headers"].remove(col_to_del)
                                for r in st.session_state["current_sheet_data"]:
                                    if col_to_del in r:
                                        r.pop(col_to_del)
                                save_persistent_state()
                                st.success(f"Стовпчик '{col_to_del}' видалено!")
                                rerun_needed = True
                                # We let it rerun at the end of the block if rerun_needed is set, but st.rerun() is immediate anyway.
                                st.rerun()
                    else:
                        st.caption("Немає активних стовпчиків.")
                        
            with col_ctrl3:
                if st.button("➕ Додати рядок", key="btn_compact_add_row", width="stretch"):
                    new_row = {h: "" for h in st.session_state["current_sheet_headers"]}
                    st.session_state["current_sheet_data"].append(new_row)
                    save_persistent_state()
                    st.success("Рядок додано!")
                    st.rerun()
                    
            with col_ctrl4:
                with st.popover("❌ Видалити рядок", width="stretch"):
                    if st.session_state["current_sheet_data"]:
                        row_to_del = st.number_input("Номер рядка для видалення (з 1):", min_value=1, max_value=len(st.session_state["current_sheet_data"]), step=1, key="compact_del_row_num")
                        confirm_row = st.checkbox("Підтверджую видалення рядка", key="compact_confirm_del_row")
                        if st.button("❌ Видалити рядок", key="btn_compact_del_row", type="primary", width="stretch"):
                            if not confirm_row:
                                st.error("Будь ласка, підтвердіть видалення!")
                            else:
                                idx = int(row_to_del) - 1
                                st.session_state["current_sheet_data"].pop(idx)
                                save_persistent_state()
                                st.success(f"Рядок {row_to_del} видалено!")
                                st.rerun()
                    else:
                        st.caption("Немає доступних рядків.")
                        
            # --- DATA EDITOR AND SAVE ---
            st.markdown("### 📊 Дані рядків для генерації документів (починаючи з рядка 5)")
            st.caption("Оберіть рядок для перегляду за допомогою прапорця в колонці **«Перегляд»** ліворуч. Клікніть двічі на будь-завгодно клітинку для редагування.")
            
            # --- INTERACTIVE COLUMN HEADERS EDITOR (EDITABLE BY DOUBLE-CLICK) ---
            st.markdown("##### ✏️ Редагувати назви змінних (подвійний клік на осередки правої колонки):")
            headers_list = st.session_state["current_sheet_headers"]
            headers_df = pd.DataFrame({
                "Поточне ім'я змінної": headers_list,
                "Нова назва змінної": headers_list
            })
            
            clean_cfg_path = "".join([c if c.isalnum() else "_" for c in cfg_path])
            edited_headers_df = st.data_editor(
                headers_df,
                column_config={
                    "Поточне ім'я змінної": st.column_config.TextColumn(disabled=True),
                    "Нова назва змінної": st.column_config.TextColumn(disabled=False)
                },
                hide_index=True,
                width="stretch",
                key=f"headers_data_editor_{clean_cfg_path}_{selected_sheet}"
            )

            headers = st.session_state["current_sheet_headers"]
            rows = st.session_state["current_sheet_data"]
            
            if rows:
                df = pd.DataFrame(rows, columns=headers)
            else:
                df = pd.DataFrame(columns=headers)
                
            # Insert the "Перегляд" selection column as boolean at the very beginning (index 0)
            df.insert(0, "Перегляд", False)
            
            # Setup session state tracking for preview index
            if "last_preview_row_idx" not in st.session_state:
                st.session_state["last_preview_row_idx"] = 0
                
            last_idx = st.session_state["last_preview_row_idx"]
            if last_idx is not None and last_idx < len(df):
                df.at[last_idx, "Перегляд"] = True
                
            edited_df = st.data_editor(
                df,
                num_rows="dynamic",
                width="stretch",
                key=f"config_data_editor_{clean_cfg_path}_{selected_sheet}"
            )
            
            # Determine selected index, default to the last tracked one
            selected_row_idx = st.session_state["last_preview_row_idx"]
            
            st.write(" ")
            
            if st.button("💾 Зберегти зміни в Excel", type="primary", key="btn_save_config", use_container_width=True):
                clean_rows = []
                for _, r in edited_df.iterrows():
                    row_dict = {}
                    for h in headers:
                        val = r.get(h, "")
                        row_dict[h] = str(val) if pd.notna(val) else ""
                    clean_rows.append(row_dict)
                    
                st.session_state["current_sheet_data"] = clean_rows
                
                success = save_excel_config(
                    cfg_path,
                    selected_sheet,
                    st.session_state["editor_template_path"],
                    st.session_state["editor_name_pattern"],
                    headers,
                    clean_rows
                )
                if success:
                    # Apply any pending template variable renames when saving config
                    pending_renames = st.session_state.get("pending_template_renames", [])
                    if pending_renames:
                        t_path = st.session_state.get("editor_template_path", "")
                        if t_path:
                            from _templates_machine_ import resolve_path
                            cfg_dir = os.path.dirname(os.path.abspath(cfg_path))
                            actual_t_path = resolve_path(cfg_dir, t_path)
                            
                            if os.path.exists(actual_t_path):
                                renamed_count = 0
                                for old_n, new_n in pending_renames:
                                    if rename_placeholder_in_template(actual_t_path, old_n, new_n):
                                        renamed_count += 1
                                if renamed_count > 0:
                                    st.toast(f"✅ Шаблон також оновлено (перейменовано {renamed_count} змінних)", icon="📄")
                        
                        # Clear pending renames list after successfully applying them
                        st.session_state["pending_template_renames"] = []

                    # Clear pm inputs and reset PM editing vars to force reload updated cells from Excel on navigation back
                    clear_pm_input_keys()
                    st.session_state["pm_editing_vars"] = None
                    st.success("🎉 Всі зміни успішно записані в Excel файл!")
                    st.balloons()
                    time.sleep(1)
                    st.rerun()
                    
            st.write(" ")  # spacer
            # --- HIGH-FIDELITY LIVE DOCUMENT PREVIEW WINDOW ---
            if selected_row_idx is not None and selected_row_idx < len(edited_df):
                st.write(" ")
                st.write(" ")
                st.markdown(f"### 🔍 Інтелектуальний миттєвий перегляд документа (Рядок {selected_row_idx + 5})")
                
                # Extract variables for this row
                row_vars = {}
                row_data = edited_df.iloc[selected_row_idx]
                for h in headers:
                    val = row_data.get(h, "")
                    row_vars[h] = str(val) if pd.notna(val) else ""
                    
                # Map pending renames to support both old and new names in live preview before saving
                pending_renames = st.session_state.get("pending_template_renames", [])
                for old_n, new_n in pending_renames:
                    if new_n in row_vars:
                        row_vars[old_n] = row_vars[new_n]
                    
                t_path = st.session_state["editor_template_path"]
                if t_path:
                    from _templates_machine_ import resolve_path
                    cfg_dir = os.path.dirname(os.path.abspath(cfg_path))
                    actual_template_path = resolve_path(cfg_dir, t_path)
                            
                    ext = os.path.splitext(t_path)[1].lower()
                    
                    with st.spinner("⏳ Генерація прев'ю..."):
                        if ext == ".docx":
                            preview_html = generate_docx_preview(actual_template_path, row_vars)
                        elif ext == ".xlsx":
                            preview_html = generate_xlsx_preview(actual_template_path, row_vars)
                        else:
                            preview_html = f"<div style='color: #e53e3e;'>Непідтримуваний тип шаблону: {ext}</div>"
                            
                    st.markdown(
                        f"""
                        <div class="document-preview-container" style="border: 2px solid #3182ce; border-radius: 8px; padding: 25px; background-color: #ffffff; max-height: 600px; overflow-y: auto; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1), 0 2px 4px -1px rgba(0,0,0,0.06); border-left: 8px solid #3182ce;">
                            {preview_html}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.warning("Будь ласка, вкажіть шлях до файлу шаблону в полі A1, щоб активувати попередній перегляд!")

# ----------------------------------------------------
# VIEW 3: DOCUMENT GENERATION & PROGRESS TRACKING
# ----------------------------------------------------
elif selected_view == "⚡ Генерація Документів":
    st.header("⚡ Масова генерація документів")
    st.write("Оберіть файл конфігурації, перевірте шлях до шаблону і запустіть автоматичне заповнення.")
    
    col_g1, col_g2 = st.columns([3, 1])
    with col_g1:
        g_config_input = st.text_input(
            "Шлях до файлу конфігурації для генерації:",
            placeholder="Оберіть Excel-файл конфігурації...",
            key="gen_config_path",
            on_change=save_persistent_state
        )
    with col_g2:
        st.write(" ")
        st.write(" ")
        def select_gen_config():
            selected = open_file_picker(filetypes=[("Excel конфігурації", "*.xlsx"), ("Усі файли", "*.*")])
            if selected:
                st.session_state["gen_config_path"] = selected
                st.session_state["last_opened_config"] = selected
                st.session_state["last_opened_folder"] = os.path.dirname(os.path.abspath(selected))
                save_persistent_state()
        st.button("📁 Обрати конфіг", key="btn_gen_config", on_click=select_gen_config)
        
    # Sync document generation path to last_opened state if it changes
    g_path = st.session_state["gen_config_path"]
    if g_path and g_path != st.session_state.get("last_opened_config"):
        try:
            abs_g_path = os.path.abspath(g_path)
            g_folder = os.path.dirname(abs_g_path)
            st.session_state["last_opened_config"] = g_path
            st.session_state["last_opened_folder"] = g_folder
            save_persistent_state()
        except Exception:
            pass
    
    # --- ALWAYS-VISIBLE OUTPUT DIRECTORY SELECTION ---
    st.markdown("##### 📁 Папка для збереження готових документів (Необов'язково)")
    st.caption("Якщо не вказано, файли зберігатимуться відносно папки Excel-конфігурації. У безвіконному (headless) режимі введіть шлях вручну.")
    
    col_go1, col_go2 = st.columns([3, 1])
    with col_go1:
        g_out_dir = st.text_input(
            "📁 Шлях до папки збереження:",
            placeholder="Наприклад: C:/ProcessedDocs (залиште порожнім за замовчуванням)",
            key="gen_output_dir",
            on_change=save_persistent_state
        )
    with col_go2:
        st.write(" ")
        st.write(" ")
        def select_gen_output_dir():
            selected = open_folder_picker("Оберіть папку для збереження готових документів")
            if selected:
                st.session_state["gen_output_dir"] = selected
                save_persistent_state()
        st.button("📁 Обрати папку", key="btn_gen_output_dir", on_click=select_gen_output_dir)
    if not g_path:
        st.info("Будь ласка, оберіть Excel-файл конфігурації для генерації!")
    elif not os.path.exists(g_path):
        st.error(f"Вказаний файл конфігурації '{g_path}' не знайдено!")
    else:
        g_sheets_data = load_excel_config(g_path)
        
        if g_sheets_data:
            g_sheet_names = ["all (Всі аркуші)"] + list(g_sheets_data.keys())
            selected_g_sheet = st.selectbox("Оберіть аркуш для обробки:", g_sheet_names, key="gen_sheet_select", on_change=save_persistent_state)
            
            actual_sheet_name = ""
            if selected_g_sheet != "all (Всі аркуші)":
                actual_sheet_name = selected_g_sheet
                
            # State synchronization on config/sheet change for generator
            gen_sheet_key = f"gen_{g_path}_{actual_sheet_name}"
            if st.session_state.get("loaded_gen_sheet") != gen_sheet_key:
                st.session_state["loaded_gen_sheet"] = gen_sheet_key
                if actual_sheet_name:
                    st.session_state["gen_template_path"] = g_sheets_data[actual_sheet_name]["template_path"]
                else:
                    st.session_state["gen_template_path"] = ""
                save_persistent_state()
                
            # --- VIEW/CHANGE TEMPLATE FILE VIA DIALOG ---
            if actual_sheet_name:
                sheet_data = g_sheets_data[actual_sheet_name]
                current_template = sheet_data["template_path"]
                
                st.markdown("##### 📄 Перевірка та зміна файлу шаблону (комірка A1)")
                
                col_gt1, col_gt2 = st.columns([3, 1])
                with col_gt1:
                    new_gt_path = st.text_input(
                        "📄 Поточний шлях до шаблону:",
                        placeholder="Оберіть файл шаблону...",
                        key="gen_template_path",
                        on_change=save_persistent_state
                    )
                with col_gt2:
                    st.write(" ")
                    st.write(" ")
                    def select_gen_template():
                        selected = open_file_picker(filetypes=[
                            ("Усі шаблони", "*.docx;*.xlsx"),
                            ("Word шаблони", "*.docx"),
                            ("Excel шаблони", "*.xlsx"),
                            ("Усі файли", "*.*")
                        ])
                        if selected:
                            st.session_state["gen_template_path"] = selected
                            save_persistent_state()
                            if update_config_template_path(g_path, actual_sheet_name, selected):
                                st.toast(f"✅ Шаблон оновлено в Excel: {os.path.basename(selected)}", icon="💾")
                    st.button("📁 Змінити шаблон", key="btn_gen_template", on_click=select_gen_template)
                    
                # If path was changed manually
                if st.session_state.get("gen_template_path", "") != current_template:
                    if st.button("💾 Зберегти новий шлях шаблону в Excel", key="btn_save_gen_template"):
                        manually_entered = st.session_state.get("gen_template_path", "")
                        if update_config_template_path(g_path, actual_sheet_name, manually_entered):
                            st.success("Шлях шаблону успішно оновлено в конфігурації!")
                            time.sleep(1)
                            st.rerun()
                            
            # Row selector filtering options
            selected_g_row = "all"
            if selected_g_sheet != "all (Всі аркуші)":
                rows_count = len(sheet_data["rows"])
                
                row_options = ["all (Всі рядки з даними)"]
                for i in range(rows_count):
                    preview_fields = []
                    row_dict = sheet_data["rows"][i]
                    headers_to_show = sheet_data["headers"][:3]
                    for h in headers_to_show:
                        val = row_dict.get(h, "")
                        if val:
                            preview_fields.append(f"{h}: {val}")
                    preview_str = ", ".join(preview_fields) if preview_fields else "Порожній рядок"
                    row_options.append(f"{5 + i} — ({preview_str})")
                    
                selected_g_row_str = st.selectbox("Оберіть рядок для обробки (за номером рядка в Excel):", row_options, key="gen_row_select", on_change=save_persistent_state)
                
                if "all" in selected_g_row_str:
                    selected_g_row = "all"
                else:
                    selected_g_row = selected_g_row_str.split(" — ")[0].strip()
                    
            st.markdown("---")
            st.subheader("🎯 Параметри запуску")
            
            st.markdown(f"""
            - **Файл конфігурації:** `{g_path}`
            - **Аркуш:** `{selected_g_sheet.split(' (')[0]}`
            - **Рядок:** `{selected_g_row}`
            - **Папка збереження:** `{st.session_state["gen_output_dir"] if st.session_state["gen_output_dir"] else "Папка конфігурації (за замовчуванням)"}`
            """)
            
            if st.button("⚡ Запустити генерацію документів", type="primary", key="btn_run_generation"):
                args = [g_path]
                
                sheet_arg = selected_g_sheet.split(' (')[0]
                row_arg = selected_g_row
                
                if sheet_arg != "all":
                    args.append(sheet_arg)
                    if row_arg != "all":
                        args.append(row_arg)
                elif row_arg != "all":
                    args.append("all")
                    args.append(row_arg)
                    
                # Append custom output folder if provided
                if st.session_state["gen_output_dir"]:
                    while len(args) < 3:
                        args.append("all")
                    args.append(st.session_state["gen_output_dir"])
                    
                run_subprocess_and_stream(args)
                st.rerun() # Refresh to display persistent logs immediately

            # Param change check & stable completion status rendering
            current_gen_params = f"{g_path}_{selected_g_sheet}_{selected_g_row}_{st.session_state.get('gen_output_dir', '')}"
            if st.session_state.get("last_gen_params") != current_gen_params:
                st.session_state["last_gen_params"] = current_gen_params
                st.session_state["gen_completion_status"] = None
                
            if st.session_state.get("gen_completion_status"):
                status = st.session_state["gen_completion_status"]
                st.write(" ")
                if status == "success":
                    st.success("🎉 Документи успішно згенеровано!")
                elif status == "error":
                    st.error("❌ Сталася помилка під час генерації документів. Перегляньте лог консолі нижче.")

    # Persistent log viewer at the bottom of generator tab
    show_last_operation_logs()

# ----------------------------------------------------
# VIEW 4: HELP & DOCUMENTATION
# ----------------------------------------------------
elif selected_view == "📖 Повна Довідка":
    # Render a marker element to trigger CSS scoped styling for the documentation page
    st.markdown('<div class="help-page-marker"></div>', unsafe_allow_html=True)

    st.header("📖 Повний посібник користувача")
    st.write("Детальний опис можливостей та технічний посібник роботи комбайна (завантажено з _templates_machine_.txt).")
    
    st.markdown("---")
    
    doc_markdown = get_formatted_documentation_markdown()
    st.markdown(doc_markdown)

# End of file