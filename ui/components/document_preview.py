import streamlit as st
import os
import tempfile
import html as py_html
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import PatternFill

from core.file_handlers.docx_handler import process_word
from core.file_handlers.xlsx_handler import process_excel
from core.text_processor import resolve_path

def generate_docx_preview(template_path, variables, config_path=None):
    """Generates a temporary Word document and extracts its content for quick preview in high-fidelity HTML."""
    c_path = config_path or st.session_state.get("editor_config_path")
    cfg_dir = os.path.dirname(os.path.abspath(c_path)) if c_path else os.getcwd()
    actual_path = resolve_path(cfg_dir, template_path)
            
    if not os.path.exists(actual_path):
        return f"<div style='color: red !important; font-weight: bold;'>Шаблон не знайдено за шляхом: {template_path}</div>"
        
    temp_dir = tempfile.gettempdir()
    temp_out = os.path.join(temp_dir, "temp_preview.docx")
    
    try:
        if not variables:
            doc = Document(actual_path)
        else:
            process_word(actual_path, temp_out, variables)
            doc = Document(temp_out)
        
        html = []
        html.append("<div class='word-preview-container' style='font-family: \"Times New Roman\", Times, serif; color: #1a202c !important; max-width: 800px; margin: 0 auto; line-height: 1.5; font-size: 14px; background-color: #ffffff !important;'>")
        
        # Pre-map paragraphs and tables to speed up lookup to O(1)
        p_map = {p._p: p for p in doc.paragraphs}
        t_map = {t._tbl: t for t in doc.tables}
        
        # Traverse paragraphs and tables in order
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph
                paragraph = p_map.get(element)
                if paragraph:
                    style_str = []
                    
                    # 1. Alignment
                    align = paragraph.alignment
                    if align == WD_ALIGN_PARAGRAPH.CENTER:
                        style_str.append("text-align: center;")
                    elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                        style_str.append("text-align: right;")
                    elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        style_str.append("text-align: justify;")
                    else:
                        style_str.append("text-align: left;")
                        
                    # 2. Spacing & Indents
                    space_before = paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else 0
                    space_after = paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 6
                    left_indent = paragraph.paragraph_format.left_indent.pt if paragraph.paragraph_format.left_indent else 0
                    right_indent = paragraph.paragraph_format.right_indent.pt if paragraph.paragraph_format.right_indent else 0
                    
                    style_str.append(f"margin: 0; margin-top: {space_before}pt; margin-bottom: {space_after}pt; margin-left: {left_indent}pt; margin-right: {right_indent}pt;")
                    
                    # 3. First Line Indent
                    if paragraph.paragraph_format.first_line_indent:
                        fl_val = paragraph.paragraph_format.first_line_indent.pt
                        style_str.append(f"text-indent: {fl_val}pt;")
                        
                    # 4. Line Spacing
                    line_spacing = paragraph.paragraph_format.line_spacing
                    if line_spacing:
                        if isinstance(line_spacing, float):
                            style_str.append(f"line-height: {line_spacing};")
                        else:
                            style_str.append(f"line-height: {line_spacing.pt}pt;")
                    else:
                        style_str.append("line-height: 1.15;")
                        
                    style_str.append("color: #1a202c !important;")
                    
                    html.append(f"<p style='{' '.join(style_str)}'>")
                    for run in paragraph.runs:
                        r_style = []
                        if run.bold:
                            r_style.append("font-weight: bold;")
                        if run.italic:
                            r_style.append("font-style: italic;")
                        if run.underline:
                            r_style.append("text-decoration: underline;")
                        if run.font.size:
                            r_style.append(f"font-size: {run.font.size.pt}pt;")
                        if run.font.name:
                            r_style.append(f"font-family: '{run.font.name}', 'Times New Roman', serif;")
                        if run.font.color and run.font.color.rgb:
                            rgb = run.font.color.rgb
                            r_style.append(f"color: #{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x} !important;")
                        else:
                            r_style.append("color: #1a202c !important;")
                            
                        text_html = run.text.replace("\n", "<br>")
                        html.append(f"<span style='{' '.join(r_style)}'>{text_html}</span>")
                    
                    if not paragraph.runs:
                        html.append("&nbsp;")
                        
                    html.append("</p>")
            elif element.tag.endswith('tbl'):
                # It's a table
                table = t_map.get(element)
                if table:
                    html.append("<table style='border-collapse: collapse; width: 100%; margin: 15px 0; border: 1px solid #cbd5e1; color: #1a202c !important;'>")
                    for row in table.rows:
                        html.append("<tr>")
                        for cell in row.cells:
                            html.append("<td style='border: 1px solid #cbd5e1; padding: 8px; vertical-align: top; color: #1a202c !important;'>")
                            for cell_p in cell.paragraphs:
                                cell_p_style = []
                                align = cell_p.alignment
                                if align == WD_ALIGN_PARAGRAPH.CENTER:
                                    cell_p_style.append("text-align: center;")
                                elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                                    cell_p_style.append("text-align: right;")
                                elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                    cell_p_style.append("text-align: justify;")
                                else:
                                    cell_p_style.append("text-align: left;")
                                    
                                space_before = cell_p.paragraph_format.space_before.pt if cell_p.paragraph_format.space_before else 0
                                space_after = cell_p.paragraph_format.space_after.pt if cell_p.paragraph_format.space_after else 2
                                left_indent = cell_p.paragraph_format.left_indent.pt if cell_p.paragraph_format.left_indent else 0
                                right_indent = cell_p.paragraph_format.right_indent.pt if cell_p.paragraph_format.right_indent else 0
                                
                                cell_p_style.append(f"margin: 0; margin-top: {space_before}pt; margin-bottom: {space_after}pt; margin-left: {left_indent}pt; margin-right: {right_indent}pt;")
                                
                                if cell_p.paragraph_format.first_line_indent:
                                    fl_val = cell_p.paragraph_format.first_line_indent.pt
                                    cell_p_style.append(f"text-indent: {fl_val}pt;")
                                    
                                line_spacing = cell_p.paragraph_format.line_spacing
                                if line_spacing:
                                    if isinstance(line_spacing, float):
                                        cell_p_style.append(f"line-height: {line_spacing};")
                                    else:
                                        cell_p_style.append(f"line-height: {line_spacing.pt}pt;")
                                else:
                                    cell_p_style.append("line-height: 1.15;")
                                    
                                cell_p_style.append("color: #1a202c !important;")
                                html.append(f"<p style='{' '.join(cell_p_style)}'>")
                                for run in cell_p.runs:
                                    r_style = []
                                    if run.bold:
                                        r_style.append("font-weight: bold;")
                                    if run.italic:
                                        r_style.append("font-style: italic;")
                                    if run.underline:
                                        r_style.append("text-decoration: underline;")
                                    if run.font.size:
                                        r_style.append(f"font-size: {run.font.size.pt}pt;")
                                    if run.font.color and run.font.color.rgb:
                                        rgb = run.font.color.rgb
                                        r_style.append(f"color: #{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x} !important;")
                                    else:
                                        r_style.append("color: #1a202c !important;")
                                    text_html = run.text.replace("\n", "<br>")
                                    html.append(f"<span style='{' '.join(r_style)}'>{text_html}</span>")
                                if not cell_p.runs:
                                    html.append("&nbsp;")
                                html.append("</p>")
                            html.append("</td>")
                        html.append("</tr>")
                    html.append("</table>")
                    
        html.append("</div>")
        
        if os.path.exists(temp_out):
            try: os.remove(temp_out)
            except Exception: pass
            
        return "\n".join(html)
    except Exception as e:
        return f"<div style='color: red !important; font-weight: bold;'>Помилка попереднього перегляду Word: {e}</div>"

def generate_xlsx_preview(template_path, variables, config_path=None):
    """Generates a temporary Excel document and extracts its content for quick preview in high-fidelity HTML."""
    c_path = config_path or st.session_state.get("editor_config_path")
    cfg_dir = os.path.dirname(os.path.abspath(c_path)) if c_path else os.getcwd()
    actual_path = resolve_path(cfg_dir, template_path)
            
    if not os.path.exists(actual_path):
        return f"<div style='color: red; font-weight: bold;'>Шаблон не знайдено за шляхом: {template_path}</div>"
        
    temp_dir = tempfile.gettempdir()
    temp_out = os.path.join(temp_dir, "temp_preview.xlsx")
    
    try:
        process_excel(actual_path, temp_out, variables)
        wb = openpyxl.load_workbook(temp_out, data_only=False)
        
        html = []
        html.append("<div style='font-family: \"Segoe UI\", Tahoma, Geneva, Verdana, sans-serif; color: #1a202c; overflow-x: auto;'>")
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            html.append(f"<h4 style='color: #2b6cb0 !important; margin-top: 15px; border-bottom: 2px solid #2b6cb0 !important; padding-bottom: 5px;'>Аркуш: {sheet_name}</h4>")
            
            total_rows = ws.max_row
            total_cols = ws.max_column
            
            max_preview_rows = min(total_rows or 0, 100)
            max_preview_cols = min(total_cols or 0, 25)
            
            if (total_rows and total_rows > 100) or (total_cols and total_cols > 25):
                html.append(f"<div style='background-color: #fffaf0; border: 1px solid #feebc8; color: #c05621; padding: 10px; border-radius: 6px; margin-bottom: 10px; font-size: 12px;'>⚠️ Аркуш занадто великий ({total_rows} рядків, {total_cols} колонок). Показано лише перші 100 рядків та 25 колонок.</div>")
            
            html.append("<table style='border-collapse: collapse; border: 1px solid #cbd5e1; font-size: 13px; min-width: 100%;'>")
            
            rows = list(ws.iter_rows(max_row=max_preview_rows, max_col=max_preview_cols))
            if not rows:
                html.append("<tr><td style='padding: 10px; color: #718096 !important;'>Аркуш порожній</td></tr>")
                html.append("</table>")
                continue
                
            for row in rows:
                html.append("<tr>")
                for cell in row:
                    val = cell.value
                    if val is None:
                        val = ""
                        
                    val_str = str(val).strip()
                    if val_str.startswith('='):
                        escaped_formula = py_html.escape(val_str)
                        val_display = f"<span style='color: #2b6cb0 !important; font-family: \"JetBrains Mono\", monospace; font-size: 11px; font-weight: 500; background-color: #ebf8ff !important; border: 1px solid #bee3f8; border-radius: 4px; padding: 2px 6px; display: inline-block;'>fx {escaped_formula}</span>"
                    else:
                        val_display = py_html.escape(str(val) if val is not None else "")
                        
                    styles = []
                    if cell.alignment:
                        h_align = cell.alignment.horizontal or "left"
                        v_align = cell.alignment.vertical or "center"
                        styles.append(f"text-align: {h_align};")
                        styles.append(f"vertical-align: {v_align};")
                    else:
                        styles.append("text-align: left; vertical-align: center;")
                        
                    if cell.font:
                        if cell.font.bold:
                            styles.append("font-weight: bold;")
                        if cell.font.italic:
                            styles.append("font-style: italic;")
                        if cell.font.size:
                            styles.append(f"font-size: {cell.font.size}pt;")
                        if cell.font.color:
                            try:
                                rgb = cell.font.color.rgb
                                if rgb and isinstance(rgb, str) and len(rgb) >= 6:
                                    if len(rgb) == 8:
                                        rgb = rgb[2:]
                                    if rgb != "000000" and all(c in "0123456789abcdefABCDEF" for c in rgb):
                                        styles.append(f"color: #{rgb} !important;")
                            except Exception:
                                pass
                                
                    if cell.fill and isinstance(cell.fill, PatternFill) and cell.fill.fill_type == "solid":
                        if cell.fill.fgColor:
                            try:
                                rgb = cell.fill.fgColor.rgb
                                if rgb and isinstance(rgb, str) and len(rgb) >= 6:
                                    if len(rgb) == 8:
                                        rgb = rgb[2:]
                                    if rgb != "000000" and all(c in "0123456789abcdefABCDEF" for c in rgb):
                                        styles.append(f"background-color: #{rgb} !important;")
                            except Exception:
                                pass
                                
                    styles_str = " ".join(styles)
                    html.append(f"<td style='border: 1px solid #cbd5e1; padding: 6px 12px; min-width: 80px; {styles_str}'>{val_display}</td>")
                html.append("</tr>")
            html.append("</table>")
            
        html.append("</div>")
        
        if os.path.exists(temp_out):
            try: os.remove(temp_out)
            except Exception: pass
            
        return "\n".join(html)
    except Exception as e:
        return f"<div style='color: red; font-weight: bold;'>Помилка попереднього перегляду Excel: {e}</div>"
