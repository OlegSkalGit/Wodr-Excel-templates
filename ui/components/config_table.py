import streamlit as st
import pandas as pd
import re
import os
import time

from core.io_utils import (
    load_excel_config,
    save_excel_config,
    create_new_sheet,
    rename_sheet,
    delete_sheet,
    update_config_template_path
)
from core.text_processor import resolve_path
from core.file_handlers.docx_handler import consolidate_jinja_tags, is_hf_defined
from ui.state_manager import (
    save_persistent_state,
    clear_pm_input_keys,
    open_file_picker
)

def rename_placeholder_in_template(template_path, old_name, new_name):
    """Automatically renames placeholders inside the Word (.docx) or Excel (.xlsx) template file, supporting variable whitespaces."""
    if not os.path.exists(template_path):
        return False
    ext = os.path.splitext(template_path)[1].lower()
    
    old_pattern = re.compile(r'(\{\{\s*)' + re.escape(old_name) + r'(\s*\}\})')
    
    if ext == ".docx":
        from docx import Document
        try:
            doc = Document(template_path)
            consolidate_jinja_tags(doc)
            modified = False
            
            def replace_in_runs(runs):
                nonlocal modified
                for r in runs:
                    if r.text and old_pattern.search(r.text):
                        r.text = old_pattern.sub(lambda m: m.group(1) + new_name + m.group(2), r.text)
                        modified = True
                        
            for p in doc.paragraphs:
                replace_in_runs(p.runs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_runs(p.runs)
            
            for section in doc.sections:
                for attr in ['header', 'footer', 'first_page_header', 'first_page_footer', 'even_page_header', 'even_page_footer']:
                    if not is_hf_defined(section, attr): continue
                    h_f = getattr(section, attr, None)
                    if h_f is not None:
                        for p in h_f.paragraphs:
                            replace_in_runs(p.runs)
                        for table in h_f.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        replace_in_runs(p.runs)
            if modified:
                doc.save(template_path)
                return True
        except Exception as e:
            import traceback
            traceback.print_exc()
            st.error(f"Помилка при перейменуванні в Word шаблоні: {e}. Можливо, файл відкритий в іншій програмі.")
            
    elif ext == ".xlsx":
        import openpyxl
        try:
            wb = openpyxl.load_workbook(template_path)
            modified = False
            for sheet in wb.worksheets:
                if old_pattern.search(sheet.title):
                    new_title = old_pattern.sub(lambda m: m.group(1) + new_name + m.group(2), sheet.title)
                    clean_title = re.sub(r'[\\/*?:\[\]]', "", new_title)[:31].strip()
                    if not clean_title: clean_title = "Sheet"
                    sheet.title = clean_title
                    modified = True
                    
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            if old_pattern.search(cell.value):
                                cell.value = old_pattern.sub(lambda m: m.group(1) + new_name + m.group(2), cell.value)
                                modified = True
            if modified:
                wb.save(template_path)
                return True
        except Exception as e:
            import traceback
            traceback.print_exc()
            st.error(f"Помилка при перейменуванні в Excel шаблоні: {e}. Можливо, файл відкритий в іншій програмі.")
    return False

def extract_placeholders_with_context(template_path):
    """Extracts jinja2 placeholders like {{var}} from a docx or xlsx template, aggregating all unique contexts with exact locations."""
    import openpyxl
    from docx import Document
    
    placeholders_list = {}
    pattern = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
    
    ext = os.path.splitext(template_path)[1].lower()
    
    if ext == ".docx":
        try:
            doc = Document(template_path)
            
            def scan_and_add(text, location_name):
                if not text:
                    return
                for match in pattern.finditer(text):
                    var_name = match.group(1).strip()
                    clean_ctx = text.strip().replace("\n", " ")
                    
                    # If paragraph is long, crop it around the match
                    if len(clean_ctx) > 120:
                        m_start = match.start()
                        m_end = match.end()
                        start_idx = max(0, m_start - 50)
                        end_idx = min(len(clean_ctx), m_end + 50)
                        prefix = "..." if start_idx > 0 else ""
                        suffix = "..." if end_idx < len(clean_ctx) else ""
                        clean_ctx = f"{prefix}{clean_ctx[start_idx:end_idx]}{suffix}"
                        
                    context = f"[{location_name}] {clean_ctx}"
                    
                    if var_name not in placeholders_list:
                        placeholders_list[var_name] = []
                    if context not in placeholders_list[var_name]:
                        placeholders_list[var_name].append(context)
                        
            # 1. Scan body paragraphs
            for idx, p in enumerate(doc.paragraphs):
                scan_and_add(p.text, f"Абзац {idx + 1}")
                
            # 2. Scan tables
            for t_idx, t in enumerate(doc.tables):
                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, p in enumerate(cell.paragraphs):
                            scan_and_add(p.text, f"Табл. {t_idx + 1}, Рядок {r_idx + 1}, Ком. {c_idx + 1}")
                            
            # 3. Scan headers and footers
            for s_idx, section in enumerate(doc.sections):
                for attr in ['header', 'footer', 'first_page_header', 'first_page_footer', 'even_page_header', 'even_page_footer']:
                    if not is_hf_defined(section, attr): continue
                    h_f = getattr(section, attr, None)
                    if h_f is not None:
                        for p_idx, p in enumerate(h_f.paragraphs):
                            scan_and_add(p.text, f"Розділ {s_idx + 1} Колонтитул")
                        for t_idx, t in enumerate(h_f.tables):
                            for r_idx, row in enumerate(t.rows):
                                for c_idx, cell in enumerate(row.cells):
                                    for p in cell.paragraphs:
                                        scan_and_add(p.text, f"Розділ {s_idx + 1} Колонтитул Табл.")
        except Exception as e:
            st.error(f"Помилка зчитування Word: {e}")
            
    elif ext == ".xlsx":
        try:
            wb = openpyxl.load_workbook(template_path, data_only=False)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for r in ws.iter_rows():
                    for cell in r:
                        val = cell.value
                        if isinstance(val, str):
                            matches = pattern.finditer(val)
                            for match in matches:
                                var_name = match.group(1).strip()
                                start_idx = max(0, match.start() - 50)
                                end_idx = min(len(val), match.end() + 50)
                                prefix = "..." if start_idx > 0 else ""
                                suffix = "..." if end_idx < len(val) else ""
                                clean_val = val[start_idx:end_idx].strip().replace("\n", " ")
                                context = f"[Аркуш: {sheet_name}, Ком: {cell.coordinate}] {prefix}{clean_val}{suffix}"
                                
                                if var_name not in placeholders_list:
                                    placeholders_list[var_name] = []
                                if context not in placeholders_list[var_name]:
                                    placeholders_list[var_name].append(context)
        except Exception as e:
            st.error(f"Помилка зчитування Excel: {e}")
            
    placeholders = {}
    for var, contexts in placeholders_list.items():
        placeholders[var] = " | ".join(contexts)
        
    return placeholders

def render_config_editor(cfg_path):
    """Renders interactive configuration sheet manager, variable editors, and data tables."""
    if not cfg_path:
        st.info("Будь ласка, оберіть файл конфігурації Excel (`*.xlsx`) для початку роботи!")
        return
        
    if not os.path.exists(cfg_path):
        st.error(f"Вказаний файл конфігурації '{cfg_path}' не знайдено!")
        return

    try:
        sheets_data = load_excel_config(cfg_path)
    except Exception as e:
        st.error(f"Не вдалося завантажити конфігурацію: {e}")
        return
        
    if sheets_data:
        sheet_names = list(sheets_data.keys())
        
        if "editor_selected_sheet" not in st.session_state or st.session_state["editor_selected_sheet"] not in sheet_names:
            st.session_state["editor_selected_sheet"] = sheet_names[0] if sheet_names else ""
            
        selected_sheet = st.selectbox(
            "Оберіть аркуш конфігурації:",
            sheet_names,
            key="editor_selected_sheet",
            on_change=save_persistent_state
        )
            
        # --- SHEET CRUD MANAGEMENT ---
        with st.expander("📂 Керування аркушами (Sheets Operations)", expanded=False):
            col_s1, col_s2, col_s3 = st.columns(3)
            
            with col_s1:
                st.markdown("##### ➕ Створити новий аркуш")
                new_s_name = st.text_input("Назва нового аркуша:", key="txt_new_sheet_name")
                if st.button("➕ Створити аркуш", key="btn_create_sheet"):
                    if not new_s_name:
                        st.error("Введіть назву аркуша!")
                    else:
                        clean_s_name = re.sub(r'[\\/*?:\[\]]', "", new_s_name)[:31].strip()
                        success, err = create_new_sheet(cfg_path, clean_s_name)
                        if success:
                            st.session_state["editor_selected_sheet"] = clean_s_name
                            save_persistent_state()
                            st.success(f"Аркуш '{clean_s_name}' успішно створено!")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error(err)
                            
            with col_s2:
                st.markdown("##### ✏️ Перейменувати поточний аркуш")
                new_rename_name = st.text_input("Нова назва аркуша:", key="txt_rename_sheet_name")
                if st.button("✏️ Перейменувати аркуш", key="btn_rename_sheet"):
                    if not new_rename_name:
                        st.error("Введіть нову назву!")
                    else:
                        clean_rename_name = re.sub(r'[\\/*?:\[\]]', "", new_rename_name)[:31].strip()
                        success, err = rename_sheet(cfg_path, selected_sheet, clean_rename_name)
                        if success:
                            st.session_state["editor_selected_sheet"] = clean_rename_name
                            save_persistent_state()
                            st.success(f"Аркуш перейменовано на '{clean_rename_name}'!")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error(err)
                            
            with col_s3:
                st.markdown("##### ❌ Видалити поточний аркуш")
                st.warning("Ця дія є незворотною!")
                confirm_delete = st.checkbox("Підтверджую видалення аркуша", key="chk_confirm_delete_sheet")
                if st.button("❌ Видалити аркуш", key="btn_delete_sheet"):
                    if not confirm_delete:
                        st.error("Будь ласка, підтвердіть видалення чекбоксом!")
                    else:
                        success, err = delete_sheet(cfg_path, selected_sheet)
                        if success:
                            st.success(f"Аркуш '{selected_sheet}' видалено!")
                            remaining_sheets = [s for s in sheet_names if s != selected_sheet]
                            st.session_state["editor_selected_sheet"] = remaining_sheets[0] if remaining_sheets else ""
                            save_persistent_state()
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error(err)
                            
        # --- LOAD SHEET DATA ---
        sheet_info = sheets_data[selected_sheet]
        try:
            mtime = os.path.getmtime(cfg_path)
        except Exception:
            mtime = 0
        config_sheet_key = f"{cfg_path}_{selected_sheet}_{mtime}"
        
        # State synchronization on config/sheet change
        if st.session_state.get("loaded_config_sheet") != config_sheet_key:
            st.session_state["loaded_config_sheet"] = config_sheet_key
            st.session_state["current_sheet_headers"] = list(sheet_info["headers"])
            st.session_state["current_sheet_data"] = list(sheet_info["rows"])
            st.session_state["editor_template_path"] = sheet_info["template_path"]
            st.session_state["editor_name_pattern"] = sheet_info["name_pattern"]
            st.session_state["pending_template_renames"] = []
            if sheet_info["template_path"]:
                st.session_state["last_opened_template"] = sheet_info["template_path"]
            save_persistent_state()
            
        # Pre-populate session state values from config if they were unmounted/deleted or are empty
        if "editor_template_path" not in st.session_state or not st.session_state["editor_template_path"]:
            st.session_state["editor_template_path"] = sheet_info["template_path"]
        if "editor_name_pattern" not in st.session_state or not st.session_state["editor_name_pattern"]:
            st.session_state["editor_name_pattern"] = sheet_info["name_pattern"]
            
        # --- METADATA (A1 & A2) ---
        st.markdown("### ⚙️ Налаштування генерації для обраного аркуша")
        col_t1, col_t2 = st.columns([3, 1], vertical_alignment="bottom")
        with col_t1:
            t_path = st.text_input(
                "📄 Шлях до шаблону (комірка A1):",
                placeholder="Оберіть файл шаблону...",
                key="editor_template_path",
                on_change=save_persistent_state
            )
        with col_t2:
            def select_editor_template():
                selected = open_file_picker(filetypes=[
                    ("Усі шаблони", "*.docx;*.xlsx"),
                    ("Word шаблони", "*.docx"),
                    ("Excel шаблони", "*.xlsx"),
                    ("Усі файли", "*.*")
                ])
                if selected:
                    st.session_state["editor_template_path"] = selected
                    save_persistent_state()
            st.button("📁 Обрати", key="btn_editor_template", on_click=select_editor_template, use_container_width=True)
            
        # Show a warning/error if the template file is specified but cannot be found or is empty
        if not t_path:
            st.error("⚠️ Шлях до шаблону не вказано (комірка A1 порожня)!")
        else:
            cfg_dir = os.path.dirname(os.path.abspath(cfg_path))
            actual_t_path = resolve_path(cfg_dir, t_path)
            if not os.path.exists(actual_t_path):
                st.error("⚠️ Файл шаблону не знайдено за вказаним шляхом!")
                
        col_n1, col_n2 = st.columns([3, 1])
        with col_n1:
            n_pattern = st.text_input(
                "✍️ Шаблон імені вихідних файлів (комірка A2):",
                placeholder="Наприклад: output_{{YYYY}}.docx",
                key="editor_name_pattern",
                on_change=save_persistent_state
            )
            
        # Sync template path to last_opened_template immediately if changed
        curr_editor_t_path = st.session_state.get("editor_template_path", "")
        if curr_editor_t_path and curr_editor_t_path != st.session_state.get("last_opened_template", ""):
            st.session_state["last_opened_template"] = curr_editor_t_path
            save_persistent_state()
            

                    
        # --- DATA EDITOR AND SAVE ---
        st.markdown("### 📊 Дані рядків для генерації документів (починаючи з рядка 5)")
        st.caption("Клікніть двічі на будь-яку клітинку для редагування даних. Для перегляду конкретного документа виберіть його у лівій панелі навігації.")
        
        # --- INTERACTIVE COLUMN HEADERS EDITOR ---
        with st.expander("✏️ Редагувати назви змінних (конфіг та шаблон)", expanded=False):
            st.caption("Подвійний клік на осередки правої колонки для перейменування змінної. Після збереження конфігу зміни застосуються до таблиці, а також, за наявності, до підключеного шаблону.")
            headers_list = st.session_state["current_sheet_headers"]
            headers_df = pd.DataFrame({
                "Поточне ім'я змінної": headers_list,
                "Нова назва змінної": headers_list
            })
            
            clean_cfg_path = "".join([c if c.isalnum() else "_" for c in cfg_path])
            edited_headers_df = st.data_editor(
                headers_df,
                column_config={
                    "Поточне ім'я змінної": st.column_config.TextColumn(disabled=True),
                    "Нова назва змінної": st.column_config.TextColumn(disabled=False)
                },
                hide_index=True,
                width="stretch",
                key=f"headers_data_editor_{clean_cfg_path}_{selected_sheet}"
            )

        headers = st.session_state["current_sheet_headers"]
        rows = st.session_state["current_sheet_data"]
        
        if rows:
            df = pd.DataFrame(rows, columns=headers)
        else:
            df = pd.DataFrame(columns=headers)
            

            
        edited_df = st.data_editor(
            df,
            num_rows="fixed",
            width="stretch",
            hide_index=True,
            key=f"config_data_editor_{clean_cfg_path}_{selected_sheet}"
        )
        
        st.write(" ")
        
        # --- SLEEK ROW & COLUMN CONTROLS ---
        with st.expander("🛠️ Швидкі дії з рядками та стовпчиками (змінними)", expanded=False):
            col_ctrl1, col_ctrl2, col_ctrl3, col_ctrl4 = st.columns(4)
            
            with col_ctrl1:
                with st.popover("➕ Додати стовпчик (змінну)", width="stretch"):
                    new_col = st.text_input("Ім'я нової змінної (напр., client_name):", key="compact_new_col_input")
                    if st.button("➕ Додати стовпчик", key="btn_compact_add_col", width="stretch"):
                        new_col_clean = new_col.strip()
                        if not new_col_clean:
                            st.error("Введіть ім'я змінної!")
                        elif " " in new_col_clean:
                            st.error("Ім'я змінної не повинно містити пробілів!")
                        elif new_col_clean in st.session_state["current_sheet_headers"]:
                            st.error("Такий стовпчик вже існує!")
                        else:
                            st.session_state["current_sheet_headers"].append(new_col_clean)
                            for r in st.session_state["current_sheet_data"]:
                                r[new_col_clean] = ""
                            save_persistent_state()
                            st.success(f"Стовпчик '{new_col_clean}' додано!")
                            st.rerun()
                            
            with col_ctrl2:
                with st.popover("❌ Видалити стовпчик", width="stretch"):
                    if st.session_state["current_sheet_headers"]:
                        col_to_del = st.selectbox("Оберіть стовпчик для видалення:", st.session_state["current_sheet_headers"], key="compact_del_col_select")
                        confirm_col = st.checkbox("Підтверджую видалення стовпчика", key=f"compact_confirm_del_col_{len(st.session_state['current_sheet_headers'])}")
                        if st.button("❌ Видалити стовпчик", key="btn_compact_del_col", type="primary", width="stretch"):
                            if not confirm_col:
                                st.error("Будь ласка, підтвердіть видалення!")
                            else:
                                st.session_state["current_sheet_headers"].remove(col_to_del)
                                for r in st.session_state["current_sheet_data"]:
                                    if col_to_del in r:
                                        r.pop(col_to_del)
                                save_persistent_state()
                                st.success(f"Стовпчик '{col_to_del}' видалено!")
                                st.rerun()
                    else:
                        st.caption("Немає активних стовпчиків.")
                        
            with col_ctrl3:
                if st.button("➕ Додати рядок", key="btn_compact_add_row", width="stretch"):
                    new_row = {h: "" for h in st.session_state["current_sheet_headers"]}
                    st.session_state["current_sheet_data"].append(new_row)
                    save_persistent_state()
                    st.success("Рядок додано!")
                    st.rerun()
                    
            with col_ctrl4:
                with st.popover("❌ Видалити рядок", width="stretch"):
                    if st.session_state["current_sheet_data"]:
                        row_to_del = st.number_input("Номер рядка для видалення (з 1):", min_value=1, max_value=len(st.session_state["current_sheet_data"]), step=1, key="compact_del_row_num")
                        confirm_row = st.checkbox("Підтверджую видалення рядка", key=f"compact_confirm_del_row_{len(st.session_state['current_sheet_data'])}")
                        if st.button("❌ Видалити рядок", key="btn_compact_del_row", type="primary", width="stretch"):
                            if not confirm_row:
                                st.error("Будь ласка, підтвердіть видалення!")
                            else:
                                idx = int(row_to_del) - 1
                                st.session_state["current_sheet_data"].pop(idx)
                                save_persistent_state()
                                st.success(f"Рядок {row_to_del} видалено!")
                                st.rerun()
                    else:
                        st.caption("Немає доступних рядків.")
                        
        st.write(" ")
        
        if st.button("💾 Зберегти зміни в Excel", type="primary", key="btn_save_config", use_container_width=True):
            clean_rows = []
            for _, r in edited_df.iterrows():
                row_dict = {}
                for h in headers:
                    val = r.get(h, "")
                    row_dict[h] = str(val) if pd.notna(val) else ""
                clean_rows.append(row_dict)
                
            st.session_state["current_sheet_data"] = clean_rows
            
            success = save_excel_config(
                cfg_path,
                selected_sheet,
                st.session_state["editor_template_path"],
                st.session_state["editor_name_pattern"],
                headers,
                clean_rows
            )
            if success:
                if "pm_cached_configs" in st.session_state:
                    st.session_state["pm_cached_configs"].pop(cfg_path, None)
                st.session_state["loaded_config_sheet"] = None
                pending_renames = st.session_state.get("pending_template_renames", [])
                if pending_renames:
                    t_path = st.session_state.get("editor_template_path", "")
                    if t_path:
                        cfg_dir = os.path.dirname(os.path.abspath(cfg_path))
                        actual_t_path = resolve_path(cfg_dir, t_path)
                        
                        if os.path.exists(actual_t_path):
                            renamed_count = 0
                            for old_n, new_n in pending_renames:
                                if rename_placeholder_in_template(actual_t_path, old_n, new_n):
                                    renamed_count += 1
                            if renamed_count > 0:
                                st.toast(f"✅ Шаблон також оновлено (перейменовано {renamed_count} змінних)", icon="📄")
                    
                    st.session_state["pending_template_renames"] = []

                clear_pm_input_keys()
                st.session_state["pm_editing_vars"] = None
                st.success("🎉 Всі зміни успішно записані в Excel файл!")
                time.sleep(1)
                st.rerun()
