import sys
import os
import re
import openpyxl
from openpyxl.styles import Alignment
from docx import Document
import difflib

# --- Utilities ---

def get_unique_path(full_path):
    if not os.path.exists(full_path): return full_path
    directory, filename = os.path.dirname(full_path), os.path.basename(full_path)
    name, ext = os.path.splitext(filename)
    if not os.path.exists(directory): os.makedirs(directory, exist_ok=True)
    counter = 1
    unique_path = full_path
    while os.path.exists(unique_path):
        unique_path = os.path.join(directory, f"{name}_{counter}{ext}")
        counter += 1
    return unique_path

def get_norm_key(v):
    if v is None: return ""
    if isinstance(v, (int, float)):
        try:
            if float(v).is_integer(): return str(int(v))
        except: pass
        return str(v)
    return str(v).strip()

def tokenize(s):
    if not isinstance(s, str): return [str(s)]
    return re.findall(r'[a-zA-Z0-9а-яА-ЯёЁіІїЇєЄґҐ]+|[^\w\s]+|\s+', s)

def get_run_props(r):
    return (r.bold, r.italic, r.underline, r.font.name, r.font.size)

def group_runs(p):
    if not p.runs: return []
    grps, cur_r, cur_p = [], [p.runs[0]], get_run_props(p.runs[0])
    for r in p.runs[1:]:
        props = get_run_props(r)
        if props == cur_p: cur_r.append(r)
        else:
            grps.append((cur_p, cur_r))
            cur_r, cur_p = [r], props
    grps.append((cur_p, cur_r))
    return grps

# --- Variable Search Logic (Shared exactly as in previous scripts) ---

def find_all_vars_in_slot(strings, diff_map, v_idx):
    s0 = str(strings[0] or "")
    tok0 = tokenize(s0)
    is_var = [False] * len(tok0)
    
    for si in strings[1:]:
        if s0 == si: continue
        toki = tokenize(str(si or ""))
        sm = difflib.SequenceMatcher(None, tok0, toki)
        raw_ops = sm.get_opcodes()
        smooth_ops = []
        for i in range(len(raw_ops)):
            tag, i1, i2, j1, j2 = raw_ops[i]
            if tag == 'equal':
                text = "".join(toki[j1:j2])
                if 0 < i < len(raw_ops)-1 and len(text) < 3: tag = 'replace'
            smooth_ops.append((tag, i1, i2, j1, j2))
        
        for tag, i1, i2, j1, j2 in smooth_ops:
            if tag != 'equal':
                for k in range(i1, i2): is_var[k] = True
                
    var_ranges = []
    i = 0
    while i < len(tok0):
        if is_var[i]:
            start = i
            while i < len(tok0) and is_var[i]: i += 1
            var_ranges.append((start, i))
        else: i += 1
        
    all_values = [{} for _ in range(len(strings))]
    template_parts = []
    last_idx = 0
    for idx, (v_start, v_end) in enumerate(var_ranges):
        template_parts.append("".join(tok0[last_idx:v_start]))
        vals = []
        for f_idx, si in enumerate(strings):
            if f_idx == 0: val = "".join(tok0[v_start:v_end])
            else:
                toki = tokenize(str(si or ""))
                sm = difflib.SequenceMatcher(None, tok0, toki)
                val_bits = []
                for tag, i1, i2, j1, j2 in sm.get_opcodes():
                    o_start = max(i1, v_start)
                    o_end = min(i2, v_end)
                    if o_start < o_end:
                        if tag == 'equal': val_bits.append("".join(tok0[o_start:o_end]))
                        else: val_bits.append("".join(toki[j1:j2]))
                val = "".join(val_bits)
            vals.append(val)
        template_parts.append(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}")
        for f_idx in range(len(strings)): all_values[f_idx][idx] = vals[f_idx]
        last_idx = v_end
    template_parts.append("".join(tok0[last_idx:]))
    return "".join(template_parts), var_ranges, all_values

def get_var_name(vals, diff_map, v_idx_list):
    norm_vals = tuple(get_norm_key(v) for v in vals)
    if norm_vals in diff_map: return diff_map[norm_vals]
    v_name = f"field_{v_idx_list[0]}"
    diff_map[norm_vals] = v_name
    v_idx_list[0] += 1
    return v_name

# --- Group Processing ---

def process_paragraph_list(pars_matrix, diff_map, v_idx_list, all_data):
    num_files = len(pars_matrix)
    num_pars = min(len(p) for p in pars_matrix)
    for p_idx in range(num_pars):
        p_list = [pars_matrix[f_idx][p_idx] for f_idx in range(num_files)]
        all_texts = [p.text for p in p_list]
        if len(set(all_texts)) == 1: continue
        
        g_all = [group_runs(p) for p in p_list]
        if all(len(g) == len(g_all[0]) for g in g_all) and g_all[0]:
            for g_idx in range(len(g_all[0])):
                run_texts = ["".join(r.text for r in g[g_idx][1]) for g in g_all]
                if len(set(run_texts)) == 1: continue
                t_str, var_rs, vals_f = find_all_vars_in_slot(run_texts, diff_map, v_idx_list)
                if var_rs:
                    for idx in range(len(var_rs)):
                        v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx_list)
                        t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                        for f_idx in range(num_files): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                    tr = g_all[0][g_idx][1]
                    tr[0].text = t_str
                    for r in tr[1:]: r.text = ""
        else:
            t_str, var_rs, vals_f = find_all_vars_in_slot(all_texts, diff_map, v_idx_list)
            if var_rs:
                for idx in range(len(var_rs)):
                    v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx_list)
                    t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                    for f_idx in range(num_files): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                p_list[0].text = t_str

def compare_word_group(files, template_out):
    docs = [Document(f) for f in files]
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    process_paragraph_list([d.paragraphs for d in docs], diff_map, v_idx, all_data)
    for t_i in range(min(len(d.tables) for d in docs)):
        t_list = [d.tables[t_i] for d in docs]
        for r in range(min(len(t.rows) for t in t_list)):
            for c in range(min(len(t.columns) for t in t_list)):
                c_list = [t.cell(r, c) for t in t_list]
                if len(set(c.text for c in c_list)) > 1:
                    process_paragraph_list([c.paragraphs for c in c_list], diff_map, v_idx, all_data)
    t_p = get_unique_path(template_out)
    docs[0].save(t_p)
    return t_p, all_data, v_idx[0]-1

def compare_excel_group(files, template_out):
    wbs_d = [openpyxl.load_workbook(f, data_only=True) for f in files]
    wb_t = openpyxl.load_workbook(files[0])
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    num_sheets = min(len(wb.sheetnames) for wb in wbs_d)
    for i in range(num_sheets):
        titles = [wb.sheetnames[i] for wb in wbs_d]
        if len(set(titles)) > 1:
            t_str, var_rs, vals_f = find_all_vars_in_slot(titles, diff_map, v_idx)
            if var_rs:
                for idx in range(len(var_rs)):
                    v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                    t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                    for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                wb_t.worksheets[i].title = t_str
    for i in range(num_sheets):
        sheets_d = [wb.worksheets[i] for wb in wbs_d]
        st = wb_t.worksheets[i]
        max_r = max((s.max_row or 1) for s in sheets_d)
        max_c = max((s.max_column or 1) for s in sheets_d)
        for r in range(1, max_r + 1):
            for c in range(1, max_c + 1):
                cell_t = st.cell(row=r, column=c)
                if cell_t.data_type == 'f' or (isinstance(cell_t.value, str) and str(cell_t.value).startswith('=')): continue
                vals = [s.cell(row=r, column=c).value for s in sheets_d]
                if len(set(vals)) > 1:
                    if not all(isinstance(v, str) for v in vals if v is not None):
                        v_name = get_var_name(vals, diff_map, v_idx)
                        cell_t.value = f"{{{{ {v_name} }}}}"
                        for f_idx in range(len(files)): all_data[f_idx][v_name] = vals[f_idx]
                    else:
                        t_str, var_rs, vals_f = find_all_vars_in_slot([str(v or "") for v in vals], diff_map, v_idx)
                        if var_rs:
                            for idx in range(len(var_rs)):
                                v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                                t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                                for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                            cell_t.value = t_str
    t_p = get_unique_path(template_out)
    wb_t.save(t_p)
    return t_p, all_data, v_idx[0]-1

# --- Config and File Naming logic ---

def generate_name_template(filenames, all_data):
    base_names = [os.path.splitext(os.path.basename(f))[0] for f in filenames]
    best_pattern = None
    max_placeholders = 0
    for i in range(len(base_names)):
        name = base_names[i]
        row_data = all_data[i]
        vars_sorted = sorted([(k, str(v).strip()) for k, v in row_data.items() if v and len(str(v).strip()) >= 2], key=lambda x: len(x[1]), reverse=True)
        pattern = name
        placeholders_count = 0
        for k, v in vars_sorted:
            if v in pattern:
                pattern = pattern.replace(v, f"{{{{ {k} }}}}")
                placeholders_count += 1
        if placeholders_count > max_placeholders:
            max_placeholders = placeholders_count
            best_pattern = pattern
    if best_pattern: return best_pattern
    return f"{base_names[0]}_{{{{ YYYY }}}}{{{{ MM }}}}{{{{ DD }}}}_{{{{ hh }}}}{{{{ mm }}}}"

def save_single_config(path, template_path, all_data, num_vars, filenames):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Settings"
    try: rel_t = os.path.relpath(template_path, os.path.dirname(os.path.abspath(path)))
    except: rel_t = template_path
    
    name_pattern = generate_name_template(filenames, all_data)
    ws['A1'], ws['A2'] = rel_t, name_pattern
    
    headers = [f"field_{i}" for i in range(1, num_vars + 1)]
    for i, h in enumerate(headers): ws.cell(row=4, column=i+1).value = h
    
    unique_rows = []
    seen = set()
    for row_dict in all_data:
        row_tuple = tuple(get_norm_key(row_dict.get(h, "")) for h in headers)
        if row_tuple not in seen:
            unique_rows.append(row_dict)
            seen.add(row_tuple)
            
    for r_idx, row_dict in enumerate(unique_rows):
        for c_idx, h in enumerate(headers):
            val = row_dict.get(h, "")
            cell = ws.cell(row=5 + r_idx, column=c_idx + 1)
            cell.value = val
            if isinstance(val, str) and '\n' in val: cell.alignment = Alignment(wrapText=True)
            
    f_p = get_unique_path(path)
    wb.save(f_p)
    return f_p

def save_master_config(path, results):
    wb = openpyxl.Workbook()
    if wb.active: wb.remove(wb.active)
    for idx, (filenames, template_path, all_data, num_vars) in enumerate(results):
        base_name = os.path.splitext(os.path.basename(filenames[0]))[0]
        sheet_title = re.sub(r'[\\/*?:\[\]]', "", base_name)[:30]
        if not sheet_title: sheet_title = f"Group_{idx+1}"
        ws = wb.create_sheet(title=sheet_title)
        try: rel_t = os.path.relpath(template_path, os.path.dirname(os.path.abspath(path)))
        except: rel_t = template_path
        
        name_pattern = generate_name_template(filenames, all_data)
        ws['A1'], ws['A2'] = rel_t, name_pattern
        
        headers = [f"field_{i}" for i in range(1, num_vars + 1)]
        for i, h in enumerate(headers): ws.cell(row=4, column=i+1).value = h
        unique_rows = []
        seen = set()
        for row_dict in all_data:
            row_tuple = tuple(get_norm_key(row_dict.get(h, "")) for h in headers)
            if row_tuple not in seen:
                unique_rows.append(row_dict)
                seen.add(row_tuple)
        for r_idx, row_dict in enumerate(unique_rows):
            for c_idx, h in enumerate(headers):
                val = row_dict.get(h, "")
                cell = ws.cell(row=5 + r_idx, column=c_idx + 1)
                cell.value = val
                if isinstance(val, str) and '\n' in val: cell.alignment = Alignment(wrapText=True)
    f_p = get_unique_path(path)
    wb.save(f_p)
    return f_p

def create_bat_file(bat_path, config_path):
    s_dir = os.path.dirname(os.path.abspath(__file__))
    m_py, v_py = os.path.join(s_dir, "main.py"), os.path.join(s_dir, ".venv", "Scripts", "python.exe")
    if not os.path.exists(v_py): v_py = "python"
    cnt = f'@echo off\n"{v_py}" "{m_py}" "{os.path.abspath(config_path)}" all all\npause\n'
    with open(bat_path, "w", encoding="cp1251") as f: f.write(cnt)

# --- Auto Filtering and Grouping ---

def is_file_a_template(path, ext):
    try:
        if ext == '.docx':
            doc = Document(path)
            for p in doc.paragraphs:
                if re.search(r'\{\{.*?\}\}', p.text): return True
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        if re.search(r'\{\{.*?\}\}', c.text): return True
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if isinstance(cell, str) and re.search(r'\{\{.*?\}\}', cell): return True
    except: pass
    return False

def get_structure_key(path, ext):
    try:
        if ext == '.docx':
            doc = Document(path)
            return ("word", len(doc.paragraphs), len(doc.tables))
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, read_only=True)
            return ("excel", len(wb.sheetnames))
    except: return None

# --- Application Modes ---

def run_compare_two(f1, f2):
    ext = os.path.splitext(f1)[1].lower()
    if ext not in ['.docx', '.xlsx'] or os.path.splitext(f2)[1].lower() != ext:
        print("Помилка: Файли повинні мати однакове розширення (.docx або .xlsx).")
        return
    
    files = [f1, f2]
    f1_a = os.path.abspath(f1)
    f_dir, f_name = os.path.dirname(f1_a), os.path.splitext(os.path.basename(f1_a))[0]
    t_o = os.path.join(f_dir, f"{f_name}_template{ext}")
    m_x = os.path.join(f_dir, f"{f_name}_config.xlsx")
    b_o = os.path.join(f_dir, f"{f_name}_run_all.bat")
    
    print(f"Аналіз 2 файлів (Режим: Порівняння)...")
    try:
        if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
        else: t_p, data, n_v = compare_excel_group(files, t_o)
        m_p = save_single_config(m_x, t_p, data, n_v, files)
        create_bat_file(b_o, m_p)
        print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")
    except Exception as e:
        print(f"Помилка: {e}")

def run_package(sample, folder):
    ext = os.path.splitext(sample)[1].lower()
    files = [sample]
    for root, _, fs in os.walk(folder):
        for f in fs:
            if f.lower().endswith(ext):
                if "template" in f.lower(): continue
                fp = os.path.join(root, f)
                if os.path.abspath(fp) == os.path.abspath(sample): continue
                if is_file_a_template(fp, ext): continue
                files.append(fp)
    
    if len(files) < 2:
        print("Не знайдено подібних файлів для порівняння.")
        return
        
    sample_a = os.path.abspath(sample)
    f_dir, f_name = os.path.dirname(sample_a), os.path.splitext(os.path.basename(sample_a))[0]
    t_o = os.path.join(f_dir, f"{f_name}_template{ext}")
    m_x = os.path.join(f_dir, f"{f_name}_config.xlsx")
    b_o = os.path.join(f_dir, f"{f_name}_run_all.bat")
    
    print(f"Аналіз {len(files)} файлів (Режим: Пакетний)...")
    try:
        if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
        else: t_p, data, n_v = compare_excel_group(files, t_o)
        m_p = save_single_config(m_x, t_p, data, n_v, files)
        create_bat_file(b_o, m_p)
        print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")
    except Exception as e:
        print(f"Помилка: {e}")

def run_full_auto(folder):
    print(f"Сканування папки: {folder} (Режим: Повний автопілот)...")
    groups = {}
    for root, _, fs in os.walk(folder):
        for f in fs:
            ext = os.path.splitext(f)[1].lower()
            if ext not in ['.docx', '.xlsx']: continue
            if "template" in f.lower(): continue
            fp = os.path.join(root, f)
            if is_file_a_template(fp, ext): continue
            key = get_structure_key(fp, ext)
            if key:
                if key not in groups: groups[key] = []
                groups[key].append(fp)
    if not groups:
        print("Не знайдено підходящих документів.")
        return
    results = []
    out_dir = os.getcwd()
    for key, files in groups.items():
        if len(files) < 2:
            print(f"Група {key[0]} ({key[1]}) має лише один файл, пропускаємо.")
            continue
        ext = os.path.splitext(files[0])[1].lower()
        base_name = os.path.splitext(os.path.basename(files[0]))[0]
        t_o = os.path.join(out_dir, f"{base_name}_template{ext}")
        print(f"Обробка групи: {base_name} ({len(files)} файлів)...")
        try:
            if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
            else: t_p, data, n_v = compare_excel_group(files, t_o)
            print(f"  - Знайдено змінних: {n_v}")
            results.append((files, t_p, data, n_v))
        except Exception as e:
            print(f"  - Помилка: {e}")
    if not results:
        print("Не вдалося створити жодного шаблону.")
        return
    m_p = save_master_config(os.path.join(out_dir, "Auto_Config.xlsx"), results)
    create_bat_file(os.path.join(out_dir, "Auto_Run_All.bat"), m_p)
    print(f"\nГотово!\nСтворено конфіг: {m_p}\nОброблено груп: {len(results)}")

def main():
    args = sys.argv[1:]
    if len(args) == 1:
        folder = args[0]
        if not os.path.isdir(folder):
            print("Помилка: Папка не знайдена.")
            return
        run_full_auto(folder)
    elif len(args) == 2:
        param1, param2 = args
        if not os.path.isfile(param1):
            print("Помилка: Перший параметр має бути файлом.")
            return
        if os.path.isfile(param2):
            run_compare_two(param1, param2)
        elif os.path.isdir(param2):
            run_package(param1, param2)
        else:
            print("Помилка: Другий параметр має бути файлом або папкою.")
    else:
        print("Використання:")
        print(" 1. Повний автопілот: python _templates_.py <папка>")
        print(" 2. Пакетний режим:   python _templates_.py <файл_зразок> <папка>")
        print(" 3. Порівняння двох:  python _templates_.py <файл_1> <файл_2>")

if __name__ == "__main__":
    main()
