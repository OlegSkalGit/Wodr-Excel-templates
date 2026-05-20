import sys
import os
import re
import openpyxl
from openpyxl.styles import Alignment
from docx import Document
import difflib

# --- Utilities ---

def get_unique_path(full_path):
    if not os.path.exists(full_path): return full_path
    directory, filename = os.path.dirname(full_path), os.path.basename(full_path)
    name, ext = os.path.splitext(filename)
    if not os.path.exists(directory): os.makedirs(directory, exist_ok=True)
    counter = 1
    unique_path = full_path
    while os.path.exists(unique_path):
        unique_path = os.path.join(directory, f"{name}_{counter}{ext}")
        counter += 1
    return unique_path

def get_norm_key(v):
    if v is None: return ""
    if isinstance(v, (int, float)):
        try:
            if float(v).is_integer(): return str(int(v))
        except: pass
        return str(v)
    return str(v).strip()

def tokenize(s):
    if not isinstance(s, str): return [str(s)]
    return re.findall(r'[a-zA-Z0-9а-яА-ЯёЁіІїЇєЄґҐ]+|[^\w\s]+|\s+', s)

def get_run_props(r):
    return (r.bold, r.italic, r.underline, r.font.name, r.font.size)

def group_runs(p):
    if not p.runs: return []
    grps, cur_r, cur_p = [], [p.runs[0]], get_run_props(p.runs[0])
    for r in p.runs[1:]:
        props = get_run_props(r)
        if props == cur_p: cur_r.append(r)
        else:
            grps.append((cur_p, cur_r))
            cur_r, cur_p = [r], props
    grps.append((cur_p, cur_r))
    return grps

# --- Multi-file Logic ---

def find_variables_in_strings(strings):
    if not strings: return "", [], []
    s0 = str(strings[0] or "")
    tok0 = tokenize(s0)
    is_constant = [True] * len(tok0)
    
    for si in strings[1:]:
        toki = tokenize(str(si or ""))
        sm = difflib.SequenceMatcher(None, tok0, toki)
        current_equal = [False] * len(tok0)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                for k in range(i1, i2): current_equal[k] = True
        for k in range(len(tok0)):
            if not current_equal[k]: is_constant[k] = False
            
    var_ranges = []
    i = 0
    while i < len(tok0):
        if not is_constant[i]:
            start = i
            while i < len(tok0) and not is_constant[i]: i += 1
            var_ranges.append((start, i))
        else: i += 1
            
    if not var_ranges: return s0, [], []
    
    result_template = []
    last_idx = 0
    for idx, (v_start, v_end) in enumerate(var_ranges):
        result_template.append("".join(tok0[last_idx:v_start]))
        result_template.append(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}")
        last_idx = v_end
    result_template.append("".join(tok0[last_idx:]))
    template_str = "".join(result_template)
    
    all_values = []
    for si in strings:
        toki = tokenize(str(si or ""))
        sm = difflib.SequenceMatcher(None, tok0, toki)
        values = {}
        for idx, (v_start, v_end) in enumerate(var_ranges):
            val_parts = []
            for tag, i1, i2, j1, j2 in sm.get_opcodes():
                overlap_start = max(i1, v_start)
                overlap_end = min(i2, v_end)
                if overlap_start < overlap_end:
                    if tag == 'equal': val_parts.append("".join(tok0[overlap_start:overlap_end]))
                    else: val_parts.append("".join(toki[j1:j2]))
            values[idx] = "".join(val_parts)
        all_values.append(values)
        
    return template_str, var_ranges, all_values

def get_var_name(vals, diff_map, v_idx_list):
    norm_vals = tuple(get_norm_key(v) for v in vals)
    if norm_vals in diff_map: return diff_map[norm_vals]
    v_name = f"field_{v_idx_list[0]}"
    diff_map[norm_vals] = v_name
    v_idx_list[0] += 1
    return v_name

# --- Document Processing ---

def process_paragraph_list(pars_matrix, diff_map, v_idx_list, all_data):
    num_files = len(pars_matrix)
    num_pars = min(len(p) for p in pars_matrix)
    for p_idx in range(num_pars):
        p_list = [pars_matrix[f_idx][p_idx] for f_idx in range(num_files)]
        all_texts = [p.text for p in p_list]
        if len(set(all_texts)) == 1: continue
        
        g_all = [group_runs(p) for p in p_list]
        if all(len(g) == len(g_all[0]) for g in g_all) and g_all[0]:
            for g_idx in range(len(g_all[0])):
                run_texts = ["".join(r.text for r in g[g_idx][1]) for g in g_all]
                if len(set(run_texts)) == 1: continue
                t_str, var_rs, vals_f = find_variables_in_strings(run_texts)
                if var_rs:
                    for idx in range(len(var_rs)):
                        v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx_list)
                        t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                        for f_idx in range(num_files): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                    tr = g_all[0][g_idx][1]
                    tr[0].text = t_str
                    for r in tr[1:]: r.text = ""
        else:
            t_str, var_rs, vals_f = find_variables_in_strings(all_texts)
            if var_rs:
                for idx in range(len(var_rs)):
                    v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx_list)
                    t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                    for f_idx in range(num_files): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                p_list[0].text = t_str

def compare_all_word(files, template_out, config_out):
    print(f"Аналіз {len(files)} Word документів...")
    docs = [Document(f) for f in files]
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    process_paragraph_list([d.paragraphs for d in docs], diff_map, v_idx, all_data)
    for t_i in range(min(len(d.tables) for d in docs)):
        t_list = [d.tables[t_i] for d in docs]
        for r in range(min(len(t.rows) for t in t_list)):
            for c in range(min(len(t.columns) for t in t_list)):
                c_list = [t.cell(r, c) for t in t_list]
                if len(set(c.text for c in c_list)) > 1:
                    process_paragraph_list([c.paragraphs for c in c_list], diff_map, v_idx, all_data)
    
    t_p = get_unique_path(template_out)
    docs[0].save(t_p)
    m_p = save_config(config_out, t_p, all_data, v_idx[0]-1)
    return t_p, m_p

def compare_all_excel(files, template_out, config_out):
    print(f"Аналіз {len(files)} Excel файлів...")
    wbs_d = [openpyxl.load_workbook(f, data_only=True) for f in files]
    wb_t = openpyxl.load_workbook(files[0])
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    
    # Sheets titles
    num_sheets = min(len(wb.sheetnames) for wb in wbs_d)
    for i in range(num_sheets):
        titles = [wb.sheetnames[i] for wb in wbs_d]
        if len(set(titles)) > 1:
            t_str, var_rs, vals_f = find_variables_in_strings(titles)
            if var_rs:
                for idx in range(len(var_rs)):
                    v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                    t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                    for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                wb_t.worksheets[i].title = t_str
                
    # Cells
    for i in range(num_sheets):
        sheets_d = [wb.worksheets[i] for wb in wbs_d]
        st = wb_t.worksheets[i]
        max_r = max(s.max_row for s in sheets_d)
        max_c = max(s.max_column for s in sheets_d)
        for r in range(1, max_r + 1):
            for c in range(1, max_c + 1):
                cell_t = st.cell(row=r, column=c)
                if cell_t.data_type == 'f' or (isinstance(cell_t.value, str) and cell_t.value.startswith('=')): continue
                vals = [s.cell(row=r, column=c).value for s in sheets_d]
                if len(set(vals)) > 1:
                    # If any are not strings, use simple replacement
                    if not all(isinstance(v, str) for v in vals if v is not None):
                        v_name = get_var_name(vals, diff_map, v_idx)
                        cell_t.value = f"{{{{ {v_name} }}}}"
                        for f_idx in range(len(files)): all_data[f_idx][v_name] = vals[f_idx]
                    else:
                        str_vals = [str(v or "") for v in vals]
                        t_str, var_rs, vals_f = find_variables_in_strings(str_vals)
                        if var_rs:
                            for idx in range(len(var_rs)):
                                v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                                t_str = t_str.replace(f"{{{{ VAR_PLACEHOLDER_{idx} }}}}", f"{{{{ {v_name} }}}}")
                                for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                            cell_t.value = t_str
                            
    t_p = get_unique_path(template_out)
    wb_t.save(t_p)
    m_p = save_config(config_out, t_p, all_data, v_idx[0]-1)
    return t_p, m_p

def save_config(path, template_path, all_data, num_vars):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Settings"
    try: rel_t = os.path.relpath(template_path, os.path.dirname(os.path.abspath(path)))
    except: rel_t = template_path
    ws['A1'], ws['A2'] = rel_t, "Result_{{ YYYY }}{{ MM }}{{ DD }}_{{ hh }}{{ mm }}{{ ss }}"
    
    headers = [f"field_{i}" for i in range(1, num_vars + 1)]
    for i, h in enumerate(headers): ws.cell(row=4, column=i+1).value = h
    
    # Deduplicate rows
    unique_rows = []
    seen = set()
    for row_dict in all_data:
        row_tuple = tuple(get_norm_key(row_dict.get(h, "")) for h in headers)
        if row_tuple not in seen:
            unique_rows.append(row_dict)
            seen.add(row_tuple)
            
    for r_idx, row_dict in enumerate(unique_rows):
        for c_idx, h in enumerate(headers):
            val = row_dict.get(h, "")
            cell = ws.cell(row=5 + r_idx, column=c_idx + 1)
            cell.value = val
            if isinstance(val, str) and '\n' in val: cell.alignment = Alignment(wrapText=True)
            
    f_p = get_unique_path(path)
    wb.save(f_p)
    return f_p

# --- Main ---

def is_file_a_template(path, ext):
    try:
        if ext == '.docx':
            doc = Document(path)
            for p in doc.paragraphs:
                if "{{" in p.text and "}}" in p.text: return True
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        if "{{" in c.text and "}}" in c.text: return True
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if isinstance(cell, str) and "{{" in cell and "}}" in cell:
                            return True
    except: pass
    return False

def create_bat_file(bat_path, config_path):
    s_dir = os.path.dirname(os.path.abspath(__file__))
    m_py, v_py = os.path.join(s_dir, "main.py"), os.path.join(s_dir, ".venv", "Scripts", "python.exe")
    if not os.path.exists(v_py): v_py = "python"
    cnt = f'@echo off\n"{v_py}" "{m_py}" "{os.path.abspath(config_path)}" all all\npause\n'
    with open(bat_path, "w", encoding="cp1251") as f: f.write(cnt)

def main():
    if len(sys.argv) < 3:
        print("Использование: create_template_package.py <файл_зразок> <папка_з_документами>")
        return
    sample, folder = sys.argv[1:3]
    if not os.path.exists(sample) or not os.path.isdir(folder):
        print("Ошибка: Файл зразок або папка не знайдена.")
        return
        
    ext = os.path.splitext(sample)[1].lower()
    files = [sample]
    for root, _, fs in os.walk(folder):
        for f in fs:
            if f.lower().endswith(ext):
                # Filter 1: name contains "template"
                if "template" in f.lower(): continue
                
                fp = os.path.join(root, f)
                if os.path.abspath(fp) == os.path.abspath(sample): continue
                
                # Filter 2: content contains {{ ... }}
                if is_file_a_template(fp, ext): continue
                
                files.append(fp)
    
    if len(files) < 2:
        print("Не знайдено подібних файлів для порівняння.")
        return
        
    sample_a = os.path.abspath(sample)
    f_dir, f_name = os.path.dirname(sample_a), os.path.splitext(os.path.basename(sample_a))[0]
    m_x, t_o, b_o = os.path.join(f_dir, f"{f_name}_config.xlsx"), os.path.join(f_dir, f"{f_name}_template{ext}"), os.path.join(f_dir, f"{f_name}_run_all.bat")
    
    try:
        if ext == '.docx': tp, mp = compare_all_word(files, t_o, m_x)
        elif ext == '.xlsx': tp, mp = compare_all_excel(files, t_o, m_x)
        else:
            print(f"Формат {ext} не підтримується.")
            return
        create_bat_file(b_o, mp)
        print(f"\nУспішно!\nФайлів оброблено: {len(files)}\nШаблон: {tp}\nКонфіг: {mp}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Помилка: {e}")

if __name__ == "__main__":
    main()
