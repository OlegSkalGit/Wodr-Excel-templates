import sys
import os
import re
import openpyxl
from openpyxl.styles import Alignment
from docx import Document
import difflib

SUPPORTED_EXCEL = ('.xlsx',)
SUPPORTED_WORD = ('.docx',)

# --- ДОПОМІЖНІ ФУНКЦІЇ ---

def get_norm_key(v):
    if v is None: return ""
    if isinstance(v, (int, float)):
        try:
            if float(v).is_integer(): return str(int(v))
        except: pass
        return str(v)
    return str(v).strip().replace('\xa0', ' ')

def tokenize(s):
    if not isinstance(s, str): return [str(s)]
    return re.findall(r'[a-zA-Z0-9а-яА-ЯёЁіІїЇєЄґҐіІїЇєЄґҐіІїЇєЄґҐ]+|[^\w\s]+|\s+', s)

def smart_diff_replace(t1, t2, diff_map, vars_old, vars_new, var_idx_list):
    n1, n2 = get_norm_key(t1), get_norm_key(t2)
    if not isinstance(t1, str) or not isinstance(t2, str):
        if n1 == n2: return t2, False, []
        k = (n1, n2)
        v_n = diff_map.get(k)
        if not v_n:
            v_n = f"field_{var_idx_list[0]}"
            diff_map[k], vars_old[v_n], vars_new[v_n] = v_n, t1, t2
            var_idx_list[0] += 1
        return f"{{{{ {v_n} }}}}", True, [(v_n, t2)]
    if n1 == n2: return t2, False, []
    tok1, tok2 = tokenize(t1), tokenize(t2)
    s = difflib.SequenceMatcher(None, tok1, tok2)
    raw_ops = s.get_opcodes()
    smooth_ops = []
    for i in range(len(raw_ops)):
        tag, i1, i2, j1, j2 = raw_ops[i]
        if tag == 'equal':
            text = "".join(str(x) for x in tok2[j1:j2])
            if 0 < i < len(raw_ops)-1 and len(text) < 3: tag = 'replace'
        smooth_ops.append((tag, i1, i2, j1, j2))
    final_ops = []
    if smooth_ops:
        cur = list(smooth_ops[0])
        for nxt in smooth_ops[1:]:
            if cur[0] != 'equal' and nxt[0] != 'equal': cur[2], cur[4] = nxt[2], nxt[4]
            else:
                final_ops.append(tuple(cur))
                cur = list(nxt)
        final_ops.append(tuple(cur))
    result, modified, local_vars = [], False, []
    for tag, i1, i2, j1, j2 in final_ops:
        if tag == 'equal': result.append("".join(str(x) for x in tok2[j1:j2]))
        else:
            old_p, new_p = "".join(str(x) for x in tok1[i1:i2]), "".join(str(x) for x in tok2[j1:j2])
            if not old_p and not new_p: continue
            k = (get_norm_key(old_p), get_norm_key(new_p))
            v_n = diff_map.get(k)
            if not v_n:
                v_n = f"field_{var_idx_list[0]}"
                diff_map[k], vars_old[v_n], vars_new[v_n] = v_n, old_p, new_p
                var_idx_list[0] += 1
            result.append(f"{{{{ {v_n} }}}}")
            local_vars.append((v_n, new_p))
            modified = True
    return "".join(result), modified, local_vars

def group_runs(p):
    if not p.runs: return []
    grps, cur_r, cur_p = [], [p.runs[0]], (p.runs[0].bold, p.runs[0].italic, p.runs[0].underline, p.runs[0].font.name, p.runs[0].font.size)
    for r in p.runs[1:]:
        props = (r.bold, r.italic, r.underline, r.font.name, r.font.size)
        if props == cur_p: cur_r.append(r)
        else:
            grps.append((cur_p, cur_r))
            cur_r, cur_p = [r], props
    grps.append((cur_p, cur_r))
    return grps

def induce_name_pattern(base_name, vars_old):
    name, ext = os.path.splitext(base_name)
    pattern = name
    sorted_vars = sorted(vars_old.items(), key=lambda x: len(str(x[1])), reverse=True)
    for var_name, var_val in sorted_vars:
        if var_name.startswith('_'): continue
        val_str = get_norm_key(var_val)
        if val_str and val_str in pattern:
            pattern = pattern.replace(val_str, f"{{{{ {var_name} }}}}")
    if pattern == name: return f"Result_{{{{ YYYY }}}}{{{{ MM }}}}{{{{ DD }}}}"
    return pattern

def create_bat_file(bat_path, config_path):
    s_dir = os.path.dirname(os.path.abspath(__file__))
    m_py, v_py = os.path.join(s_dir, "main.py"), os.path.join(s_dir, ".venv", "Scripts", "python.exe")
    if not os.path.exists(v_py): v_py = "python"
    cnt = f'@echo off\n"{v_py}" "{m_py}" "{os.path.abspath(config_path)}" all all\npause\n'
    with open(bat_path, "w", encoding="cp1251") as f: f.write(cnt)

# --- КЕРУВАННЯ EXCEL КОНФІГОМ ---

def add_config_sheet(wb, sheet_name, template_path, vars_old, vars_new, batch_data=None):
    if sheet_name in wb.sheetnames: ws = wb[sheet_name]
    else: ws = wb.create_sheet(sheet_name)
    config_dir = os.path.dirname(os.path.abspath(sheet_name))
    try: rel_template = os.path.relpath(template_path, config_dir)
    except: rel_template = os.path.basename(template_path)
    base_file_name = os.path.basename(vars_new.get('_source_file', 'Result'))
    name_pattern = induce_name_pattern(base_file_name, vars_old)
    ws['A1'], ws['A2'] = rel_template, name_pattern
    headers = [h for h in vars_new.keys() if not h.startswith('_')]
    for i, h in enumerate(headers): ws.cell(row=4, column=i+1).value = h
    for i, h in enumerate(headers):
        v_o, v_n = vars_old.get(h, ""), vars_new.get(h, "")
        for r, v in [(5, v_o), (6, v_n)]:
            cell = ws.cell(row=r, column=i+1)
            cell.value = v
            if isinstance(v, str) and '\n' in v: cell.alignment = Alignment(wrapText=True)
    if batch_data:
        for r_idx, entry in enumerate(batch_data):
            for i, h in enumerate(headers):
                val = entry.get(h, "")
                cell = ws.cell(row=7+r_idx, column=i+1)
                cell.value = val
                if isinstance(val, str) and '\n' in val: cell.alignment = Alignment(wrapText=True)

# --- ЕКСТРАКЦІЯ ТА ОБРОБКА ГРУП ---

def extract_batch(f1, files_to_scan, ext, diff_info):
    results = []
    for f_path in files_to_scan:
        try:
            entry, has_diff = {}, False
            if ext == '.xlsx':
                wb, wb1 = openpyxl.load_workbook(f_path, data_only=True), openpyxl.load_workbook(f1, data_only=True)
                for (s_idx, r, c, var_name) in diff_info:
                    if s_idx < len(wb.worksheets):
                        ws, ws1 = wb.worksheets[s_idx], wb1.worksheets[s_idx]
                        v1, v_cur = (ws1.title, ws.title) if r == -1 else (ws1.cell(row=r, column=c).value, ws.cell(row=r, column=c).value)
                        if get_norm_key(v1) != get_norm_key(v_cur):
                            has_diff = True
                            _, _, local = smart_diff_replace(v1, v_cur, {}, {}, {}, [999])
                            for (_, val) in local: entry[var_name] = val
            elif ext == '.docx':
                doc, doc1 = Document(f_path), Document(f1)
                for (obj_type, idx, sub_idx, var_name) in diff_info:
                    if obj_type == 'p' and idx < len(doc.paragraphs):
                        t1, t_cur = doc1.paragraphs[idx].text, doc.paragraphs[idx].text
                        if get_norm_key(t1) != get_norm_key(t_cur):
                            has_diff = True
                            _, _, local = smart_diff_replace(t1, t_cur, {}, {}, {}, [999])
                            for (_, val) in local: entry[var_name] = val
                    elif obj_type == 't' and idx < len(doc.tables):
                        r, c = sub_idx
                        t1, t_cur = doc1.tables[idx].cell(r, c).text, doc.tables[idx].cell(r, c).text
                        if get_norm_key(t1) != get_norm_key(t_cur):
                            has_diff = True
                            _, _, local = smart_diff_replace(t1, t_cur, {}, {}, {}, [999])
                            for (_, val) in local: entry[var_name] = val
            if has_diff: results.append(entry)
            else: print(f"  [Пропуск] Файл ідентичний за змінними: {os.path.basename(f_path)}")
        except: pass
    return results

def process_excel_group(f_base, other_files, template_out):
    try:
        wb1_d, wb_t = openpyxl.load_workbook(f_base, data_only=True), openpyxl.load_workbook(f_base)
        vars_old, vars_new, diff_map, v_idx, diff_locs = {'_source_file': f_base}, {'_source_file': f_base}, {}, [1], []
        for i, f_other in enumerate(other_files, 1):
            try:
                print(f"  [{i}/{len(other_files)}] Аналіз: {os.path.basename(f_other)}")
                wb2_d = openpyxl.load_workbook(f_other, data_only=True)
                for s_i in range(min(len(wb1_d.sheetnames), len(wb2_d.sheetnames))):
                    s1, s2, st = wb1_d.worksheets[s_i], wb2_d.worksheets[s_i], wb_t.worksheets[s_i]
                    nv, mod, local = smart_diff_replace(s1.title, s2.title, diff_map, vars_old, vars_new, v_idx)
                    if mod:
                        st.title = nv
                        for (vn, _) in local: diff_locs.append((s_i, -1, -1, vn))
                    for r in range(1, max(s1.max_row, s2.max_row) + 1):
                        for c in range(1, max(s1.max_column, s2.max_column) + 1):
                            ct = st.cell(row=r, column=c)
                            if ct.data_type == 'f' or (isinstance(ct.value, str) and str(ct.value).startswith('=')): continue
                            v1, v2 = s1.cell(row=r, column=c).value, s2.cell(row=r, column=c).value
                            if get_norm_key(v1) != get_norm_key(v2):
                                nv, mod, local = smart_diff_replace(v1, v2, diff_map, vars_old, vars_new, v_idx)
                                if mod:
                                    ct.value = nv
                                    for (vn, _) in local: diff_locs.append((s_i, r, c, vn))
            except: continue
        wb_t.save(template_out)
        batch = extract_batch(f_base, other_files, '.xlsx', list(set(diff_locs)))
        return vars_old, vars_new, batch
    except: return {}, {}, []

def process_word_group(f_base, other_files, template_out):
    try:
        doc1, doc_t = Document(f_base), Document(f_base)
        vars_old, vars_new, diff_map, v_idx, diff_locs = {'_source_file': f_base}, {'_source_file': f_base}, {}, [1], []
        for i, f_other in enumerate(other_files, 1):
            try:
                print(f"  [{i}/{len(other_files)}] Аналіз: {os.path.basename(f_other)}")
                doc2 = Document(f_other)
                for p_i in range(min(len(doc1.paragraphs), len(doc2.paragraphs))):
                    p1, p2, pt = doc1.paragraphs[p_i], doc2.paragraphs[p_i], doc_t.paragraphs[p_i]
                    if get_norm_key(p1.text) != get_norm_key(p2.text):
                        g1, g2, gt = group_runs(p1), group_runs(p2), group_runs(pt)
                        for j in range(min(len(g1), len(g2))):
                            t1, t2 = "".join(str(r.text) for r in g1[j][1]), "".join(str(r.text) for r in g2[j][1])
                            if get_norm_key(t1) != get_norm_key(t2):
                                nv, mod, local = smart_diff_replace(t1, t2, diff_map, vars_old, vars_new, v_idx)
                                if mod:
                                    gt[j][1][0].text = nv
                                    for r in gt[j][1][1:]: r.text = ""
                                    for (vn, _) in local: diff_locs.append(('p', p_i, j, vn))
                for t_i in range(min(len(doc1.tables), len(doc2.tables))):
                    t1, t2, tt = doc1.tables[t_i], doc2.tables[t_i], doc_t.tables[t_i]
                    for r in range(min(len(t1.rows), len(t2.rows))):
                        for c in range(min(len(t1.columns), len(t2.columns))):
                            c1, c2, ct = t1.cell(r, c), t2.cell(r, c), tt.cell(r, c)
                            if get_norm_key(c1.text) != get_norm_key(c2.text):
                                nv, mod, local = smart_diff_replace(c1.text, c2.text, diff_map, vars_old, vars_new, v_idx)
                                if mod:
                                    ct.text = nv
                                    for (vn, _) in local: diff_locs.append(('t', t_i, (r, c), vn))
            except: continue
        doc_t.save(template_out)
        batch = extract_batch(f_base, other_files, '.docx', list(set(diff_locs)))
        return vars_old, vars_new, batch
    except: return {}, {}, []

def get_file_signature(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == '.docx':
            doc = Document(path)
            return (ext, len(doc.paragraphs), len(doc.tables))
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, read_only=True)
            return (ext, len(wb.sheetnames), tuple(wb.sheetnames))
    except: return None

# --- ГОЛОВНИЙ МОДУЛЬ ---

def main():
    if len(sys.argv) < 3:
        print("\nВикористання:\n1. Порівняння: create_template.py f1 f2\n2. Пакетний збір: create_template.py f1 f2 dir\n3. АВТО-АНАЛІЗ: create_template.py src_dir out_dir")
        return
    p1, p2 = sys.argv[1], sys.argv[2]
    if os.path.isdir(p1) and os.path.isdir(p2):
        abs_src, abs_out = os.path.abspath(p1), os.path.abspath(p2)
        if abs_src.startswith(abs_out) or abs_out.startswith(abs_src):
            print("Помилка: Папки не можуть бути вкладені!")
            return
        if os.listdir(abs_out):
            print("Помилка: Вихідна папка має бути порожньою!")
            return
        print(f"\n--- АВТО-АНАЛІЗ РОЗПОЧАТО ---")
        groups = {}
        for root, _, fs in os.walk(p1):
            for f in fs:
                if f.lower().endswith(SUPPORTED_EXCEL + SUPPORTED_WORD) and not f.startswith("~$"):
                    path = os.path.join(root, f)
                    sig = get_file_signature(path)
                    if sig:
                        if sig not in groups: groups[sig] = []
                        groups[sig].append(path)
        wb_cfg = openpyxl.Workbook()
        if 'Sheet' in wb_cfg.sheetnames: del wb_cfg['Sheet']
        g_idx = 1
        for sig, files in groups.items():
            ext, f_base = sig[0], files[0]
            t_name = f"template_{g_idx}{ext}"
            t_path = os.path.join(abs_out, t_name)
            print(f"\nОбробка Групи №{g_idx} ({len(files)} файлів):")
            v_o, v_n, b_d = process_excel_group(f_base, files[1:], t_path) if ext == '.xlsx' else process_word_group(f_base, files[1:], t_path)
            if v_o or v_n:
                add_config_sheet(wb_cfg, f"Group_{g_idx}", t_path, v_o, v_n, b_d)
                g_idx += 1
        cfg_path = os.path.join(abs_out, "auto_config.xlsx")
        wb_cfg.save(cfg_path)
        create_bat_file(os.path.join(abs_out, "run_auto.bat"), cfg_path)
        print(f"\n--- ЗАВЕРШЕНО ---\nЗгенеровано конфіг та BAT: {abs_out}")
    else:
        f1, f2 = p1, p2
        scan_dir = sys.argv[3] if len(sys.argv) > 3 else None
        f1_a = os.path.abspath(f1)
        ex = os.path.splitext(f1_a)[1].lower()
        f1_d, f1_n = os.path.dirname(f1_a), os.path.splitext(os.path.basename(f1_a))[0]
        m_x, t_o, b_o = os.path.join(f1_d, f"{f1_n}_config.xlsx"), os.path.join(f1_d, f"{f1_n}_template{ex}"), os.path.join(f1_d, f"{f1_n}_run_all.bat")
        v_o, v_n, b_d = process_excel_group(f1, [f2], t_o) if ex == '.xlsx' else process_word_group(f1, [f2], t_o)
        if scan_dir:
            files = [os.path.join(root, f) for root, _, fs in os.walk(scan_dir) for f in fs if f.lower().endswith(ex) and not f.startswith("~$")]
            v_o2, v_n2, b_d2 = process_excel_group(f1, files, t_o) if ex == '.xlsx' else process_word_group(f1, files, t_o)
            b_d.extend(b_d2)
        wb = openpyxl.Workbook()
        add_config_sheet(wb, "Settings", t_o, v_o, v_n, b_d)
        if 'Sheet' in wb.sheetnames: del wb['Sheet']
        wb.save(m_x)
        create_bat_file(b_o, m_x)
        print(f"Готово! Шаблон: {t_o}\nКонфіг та BAT створено.")

if __name__ == "__main__": main()
