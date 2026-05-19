import sys
import os
import re
import openpyxl
from openpyxl.styles import Alignment
from docx import Document
import difflib

def get_unique_path(full_path):
    if not os.path.exists(full_path): return full_path
    directory, filename = os.path.dirname(full_path), os.path.basename(full_path)
    name, ext = os.path.splitext(filename)
    if not os.path.exists(directory): os.makedirs(directory, exist_ok=True)
    counter = 1
    unique_path = full_path
    while os.path.exists(unique_path):
        unique_path = os.path.join(directory, f"{name}_{counter}{ext}")
        counter += 1
    return unique_path

def create_main_excel(path, template_path, vars_old, vars_new):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Settings"
    try:
        rel_template = os.path.relpath(template_path, os.path.dirname(os.path.abspath(path)))
    except: rel_template = template_path
    ws['A1'], ws['A2'] = rel_template, "Result_{{ YYYY }}{{ MM }}{{ DD }}_{{ hh }}{{ mm }}{{ ss }}"
    headers = list(vars_new.keys())
    for i, h in enumerate(headers): ws.cell(row=4, column=i+1).value = h
    for i, h in enumerate(headers):
        v_o, v_n = vars_old.get(h, ""), vars_new.get(h, "")
        for r, v in [(5, v_o), (6, v_n)]:
            c = ws.cell(row=r, column=i+1)
            c.value = v
            if isinstance(v, str) and '\n' in v: c.alignment = Alignment(wrapText=True)
    final_path = get_unique_path(path)
    wb.save(final_path)
    return final_path

def get_norm_key(v):
    if v is None: return ""
    if isinstance(v, (int, float)):
        try:
            if float(v).is_integer(): return str(int(v))
        except: pass
        return str(v)
    return str(v).strip()

def tokenize(s):
    if not isinstance(s, str): return [str(s)]
    return re.findall(r'[a-zA-Z0-9а-яА-ЯёЁіІїЇєЄґҐіІїЇєЄґҐіІїЇєЄґҐ]+|[^\w\s]+|\s+', s)

def smart_diff_replace(t1, t2, diff_map, vars_old, vars_new, var_idx_list):
    n1, n2 = get_norm_key(t1), get_norm_key(t2)
    if not isinstance(t1, str) or not isinstance(t2, str):
        if n1 == n2: return t2, False
        k = (n1, n2)
        if k in diff_map: v_n = diff_map[k]
        else:
            v_n = f"field_{var_idx_list[0]}"
            diff_map[k], vars_old[v_n], vars_new[v_n] = v_n, t1, t2
            var_idx_list[0] += 1
        return f"{{{{ {v_n} }}}}", True
    if n1 == n2: return t2, False

    tok1, tok2 = tokenize(t1), tokenize(t2)
    s = difflib.SequenceMatcher(None, tok1, tok2)
    
    # Спершу збираємо всі опікоди та "згладжуємо" дрібні однакові фрагменти всередині змін
    raw_ops = s.get_opcodes()
    smooth_ops = []
    for i in range(len(raw_ops)):
        tag, i1, i2, j1, j2 = raw_ops[i]
        # Правило згладжування: якщо 'equal' коротший за 3 символи і він НЕ на краях рядка
        if tag == 'equal':
            text = "".join(tok2[j1:j2])
            if 0 < i < len(raw_ops)-1 and len(text) < 3:
                # Перетворюємо цей 'equal' на 'replace' (зливаємо з сусідами)
                tag = 'replace'
        smooth_ops.append((tag, i1, i2, j1, j2))

    # Тепер об'єднуємо послідовні не-'equal' блоки в один великий 'replace'
    final_ops = []
    if smooth_ops:
        cur = list(smooth_ops[0])
        for nxt in smooth_ops[1:]:
            if cur[0] != 'equal' and nxt[0] != 'equal':
                cur[2], cur[4] = nxt[2], nxt[4] # Розширюємо межі
            else:
                final_ops.append(tuple(cur))
                cur = list(nxt)
        final_ops.append(tuple(cur))

    result, modified = [], False
    for tag, i1, i2, j1, j2 in final_ops:
        if tag == 'equal':
            result.append("".join(str(x) for x in tok2[j1:j2]))
        else:
            old_p, new_p = "".join(str(x) for x in tok1[i1:i2]), "".join(str(x) for x in tok2[j1:j2])
            if not old_p and not new_p: continue
            k = (get_norm_key(old_p), get_norm_key(new_p))
            if k in diff_map: v_n = diff_map[k]
            else:
                v_n = f"field_{var_idx_list[0]}"
                diff_map[k], vars_old[v_n], vars_new[v_n] = v_n, old_p, new_p
                var_idx_list[0] += 1
            result.append(f"{{{{ {v_n} }}}}")
            modified = True
            
    return "".join(result), modified

def get_run_props(r):
    return (r.bold, r.italic, r.underline, r.font.name, r.font.size)

def group_runs(p):
    if not p.runs: return []
    grps, cur_r, cur_p = [], [p.runs[0]], get_run_props(p.runs[0])
    for r in p.runs[1:]:
        props = get_run_props(r)
        if props == cur_p: cur_r.append(r)
        else:
            grps.append((cur_p, cur_r))
            cur_r, cur_p = [r], props
    grps.append((cur_p, cur_r))
    return grps

def compare_excel(f1, f2, main_xls, template_out):
    print("Аналіз Excel файлів...")
    wb1_d, wb2_d = openpyxl.load_workbook(f1, data_only=True), openpyxl.load_workbook(f2, data_only=True)
    wb_t = openpyxl.load_workbook(f1)
    vars_old, vars_new, diff_map, v_idx = {}, {}, {}, [1]
    for i in range(min(len(wb1_d.sheetnames), len(wb2_d.sheetnames))):
        s1_n, s2_n = wb1_d.sheetnames[i], wb2_d.sheetnames[i]
        nv, mod = smart_diff_replace(s1_n, s2_n, diff_map, vars_old, vars_new, v_idx)
        if mod: wb_t.worksheets[i].title = nv
    for i in range(min(len(wb1_d.worksheets), len(wb2_d.worksheets))):
        s1, s2, st = wb1_d.worksheets[i], wb2_d.worksheets[i], wb_t.worksheets[i]
        for r in range(1, max(s1.max_row, s2.max_row) + 1):
            for c in range(1, max(s1.max_column, s2.max_column) + 1):
                ct = st.cell(row=r, column=c)
                if ct.data_type == 'f' or (isinstance(ct.value, str) and ct.value.startswith('=')): continue
                v1, v2 = s1.cell(row=r, column=c).value, s2.cell(row=r, column=c).value
                if v1 != v2:
                    nv, mod = smart_diff_replace(v1, v2, diff_map, vars_old, vars_new, v_idx)
                    if mod: ct.value = nv
    t_p = get_unique_path(template_out)
    wb_t.save(t_p)
    m_p = create_main_excel(main_xls, t_p, vars_old, vars_new)
    print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")

def process_paragraphs(pars1, pars2, parst, diff_map, vars_old, vars_new, v_idx):
    for i in range(min(len(pars1), len(pars2))):
        p1, p2, pt = pars1[i], pars2[i], parst[i]
        if p1.text != p2.text:
            g1, g2, gt = group_runs(p1), group_runs(p2), group_runs(pt)
            if g1 and g2 and len(g1) == len(g2) == len(gt):
                for j in range(len(g1)):
                    t1, t2 = "".join(r.text for r in g1[j][1]), "".join(r.text for r in g2[j][1])
                    if t1 != t2:
                        nv, mod = smart_diff_replace(t1, t2, diff_map, vars_old, vars_new, v_idx)
                        if mod:
                            gt[j][1][0].text = nv
                            for r in gt[j][1][1:]: r.text = ""
            else:
                nv, mod = smart_diff_replace(p1.text, p2.text, diff_map, vars_old, vars_new, v_idx)
                if mod: pt.text = nv

def compare_word(f1, f2, main_xls, template_out):
    print("Аналіз Word документів...")
    doc1, doc2, doc_t = Document(f1), Document(f2), Document(f1)
    vars_old, vars_new, diff_map, v_idx = {}, {}, {}, [1]
    process_paragraphs(doc1.paragraphs, doc2.paragraphs, doc_t.paragraphs, diff_map, vars_old, vars_new, v_idx)
    for t_i in range(min(len(doc1.tables), len(doc2.tables))):
        t1, t2, tt = doc1.tables[t_i], doc2.tables[t_i], doc_t.tables[t_i]
        for r in range(min(len(t1.rows), len(t2.rows))):
            for c in range(min(len(t1.columns), len(t2.columns))):
                c1, c2, ct = t1.cell(r, c), t2.cell(r, c), tt.cell(r, c)
                if c1.text != c2.text:
                    process_paragraphs(c1.paragraphs, c2.paragraphs, ct.paragraphs, diff_map, vars_old, vars_new, v_idx)
    t_p = get_unique_path(template_out)
    doc_t.save(t_p)
    m_p = create_main_excel(main_xls, t_p, vars_old, vars_new)
    print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")

def create_bat_file(bat_path, config_path):
    s_dir = os.path.dirname(os.path.abspath(__file__))
    m_py, v_py = os.path.join(s_dir, "main.py"), os.path.join(s_dir, ".venv", "Scripts", "python.exe")
    if not os.path.exists(v_py): v_py = "python"
    cnt = f'@echo off\n"{v_py}" "{m_py}" "{os.path.abspath(config_path)}" all all\npause\n'
    with open(bat_path, "w", encoding="cp1251") as f: f.write(cnt)

def main():
    if len(sys.argv) < 3: return
    f1, f2 = sys.argv[1:3]
    if not os.path.exists(f1) or not os.path.exists(f2): return
    f1_a = os.path.abspath(f1)
    f1_d, f1_n, ex = os.path.dirname(f1_a), os.path.splitext(os.path.basename(f1_a))[0], os.path.splitext(f1_a)[1].lower()
    m_x, t_o, b_o = os.path.join(f1_d, f"{f1_n}_config.xlsx"), os.path.join(f1_d, f"{f1_n}_template{ex}"), os.path.join(f1_d, f"{f1_n}_run_all.bat")
    try:
        if ex == '.xlsx': compare_excel(f1, f2, m_x, t_o)
        elif ex == '.docx': compare_word(f1, f2, m_x, t_o)
        create_bat_file(b_o, m_x)
    except Exception as e: print(f"Помилка: {e}")

if __name__ == "__main__": main()
