import sys
import os
import re
from datetime import datetime
import difflib

import openpyxl
from openpyxl.styles import Alignment
from docx import Document
from docxtpl import DocxTemplate, Listing

# ==========================================
# СПІЛЬНІ УТИЛІТИ 
# ==========================================

def get_unique_path(full_path):
    directory = os.path.dirname(full_path)
    if directory and not os.path.exists(directory): 
        os.makedirs(directory, exist_ok=True)
    if not os.path.exists(full_path): return full_path
    filename = os.path.basename(full_path)
    name, ext = os.path.splitext(filename)
    counter = 1
    unique_path = full_path
    while os.path.exists(unique_path):
        unique_path = os.path.join(directory, f"{name}_{counter}{ext}")
        counter += 1
    return unique_path

def get_norm_key(v):
    if v is None: return ""
    if isinstance(v, (int, float)):
        try:
            if float(v).is_integer(): return str(int(v))
        except: pass
        return str(v)
    return str(v)

def tokenize(s):
    if not isinstance(s, str): return [str(s)]
    return re.findall(r'\w+|[^\w\s]+|\s+', s)

from lxml import etree
from docx.oxml import parse_xml

def get_run_props(r):
    if r._r.rPr is not None:
        return etree.tostring(r._r.rPr).decode('utf-8')
    return None

def get_formatting_chunks(p):
    char_to_props = []
    for r in p.runs:
        props = get_run_props(r)
        for _ in range(len(r.text)):
            char_to_props.append(props)
    chunks = []
    if not p.text: return chunks
    cur_text = p.text[0]
    cur_props = char_to_props[0]
    for i in range(1, len(p.text)):
        if char_to_props[i] == cur_props:
            cur_text += p.text[i]
        else:
            chunks.append((cur_text, cur_props))
            cur_text = p.text[i]
            cur_props = char_to_props[i]
    chunks.append((cur_text, cur_props))
    return chunks

def apply_run_props(r, props):
    if props is not None:
        try:
            rPr = parse_xml(props)
            r._r.insert(0, rPr)
        except Exception:
            pass

def get_default_props():
    return None

def resolve_path(base_dir, path):
    if os.path.isabs(path): return path
    return os.path.abspath(os.path.join(base_dir, path))

# ==========================================
# ЧАСТИНА 1: СТВОРЕННЯ ШАБЛОНІВ
# ==========================================

def find_all_vars_in_slot(strings, diff_map, v_idx, char_to_props=None):
    s0 = str(strings[0] or "")
    tok0 = tokenize(s0)
    
    tok_char_idx = []
    curr_idx = 0
    for tok in tok0:
        tok_char_idx.append(curr_idx)
        curr_idx += len(tok)
    tok_char_idx.append(curr_idx)

    is_var = [False] * len(tok0)
    for si in strings[1:]:
        if s0 == si: continue
        toki = tokenize(str(si or ""))
        sm = difflib.SequenceMatcher(None, tok0, toki)
        raw_ops = sm.get_opcodes()
        smooth_ops = []
        for i in range(len(raw_ops)):
            tag, i1, i2, j1, j2 = raw_ops[i]
            if tag == 'equal':
                text = "".join(toki[j1:j2])
                if 0 < i < len(raw_ops)-1 and len(text) < 3:
                    should_merge = True
                    if char_to_props is not None:
                        prev_op = raw_ops[i-1]
                        next_op = raw_ops[i+1]
                        start_char = tok_char_idx[prev_op[1]]
                        end_char = tok_char_idx[next_op[2]] if next_op[2] < len(tok_char_idx) else len(char_to_props)
                        props_slice = char_to_props[start_char:end_char]
                        if props_slice:
                            should_merge = (len(set(props_slice)) <= 1)
                    if should_merge:
                        tag = 'replace'
            smooth_ops.append((tag, i1, i2, j1, j2))
        for tag, i1, i2, j1, j2 in smooth_ops:
            if tag != 'equal':
                if tag == 'insert':
                    if i1 < len(tok0): is_var[i1] = True
                    elif i1 > 0: is_var[i1 - 1] = True
                for k in range(i1, i2): is_var[k] = True
    var_ranges = []
    i = 0
    while i < len(tok0):
        if is_var[i]:
            start = i
            i += 1
            if char_to_props is not None:
                start_char = tok_char_idx[start]
                while i < len(tok0) and is_var[i]:
                    curr_char = tok_char_idx[i]
                    next_char = tok_char_idx[i+1] if i+1 < len(tok_char_idx) else len(char_to_props)
                    slice_props = char_to_props[curr_char:next_char]
                    if slice_props and len(set([char_to_props[start_char]] + slice_props)) > 1:
                        break
                    i += 1
            else:
                while i < len(tok0) and is_var[i]: i += 1
            var_ranges.append((start, i))
        else: i += 1
    all_values = [{} for _ in range(len(strings))]
    template_parts = []
    last_idx = 0
    for idx, (v_start, v_end) in enumerate(var_ranges):
        template_parts.append("".join(tok0[last_idx:v_start]))
        vals = []
        for f_idx, si in enumerate(strings):
            if f_idx == 0: val = "".join(tok0[v_start:v_end])
            else:
                toki = tokenize(str(si or ""))
                sm = difflib.SequenceMatcher(None, tok0, toki)
                val_bits = []
                for tag, i1, i2, j1, j2 in sm.get_opcodes():
                    if tag == 'insert':
                        if v_start <= i1 <= v_end:
                            val_bits.append("".join(toki[j1:j2]))
                    else:
                        o_start = max(i1, v_start)
                        o_end = min(i2, v_end)
                        if o_start < o_end:
                            if tag == 'equal': val_bits.append("".join(tok0[o_start:o_end]))
                            else: val_bits.append("".join(toki[j1:j2]))
                val = "".join(val_bits)
            vals.append(val)
        template_parts.append(f"{{{{VAR_PLACEHOLDER_{idx}}}}}")
        for f_idx in range(len(strings)): all_values[f_idx][idx] = vals[f_idx]
        last_idx = v_end
    template_parts.append("".join(tok0[last_idx:]))
    return "".join(template_parts), var_ranges, all_values

def get_var_name(vals, diff_map, v_idx_list):
    norm_vals = tuple(get_norm_key(v) for v in vals)
    if norm_vals in diff_map: return diff_map[norm_vals]
    v_name = f"field_{v_idx_list[0]}"
    diff_map[norm_vals] = v_name
    v_idx_list[0] += 1
    return v_name

def process_paragraph_list(pars_matrix, diff_map, v_idx_list, all_data):
    num_files = len(pars_matrix)
    num_pars = min(len(p) for p in pars_matrix)
    for p_idx in range(num_pars):
        p_list = [pars_matrix[f_idx][p_idx] for f_idx in range(num_files)]
        
        chunk_lists = [get_formatting_chunks(p) for p in p_list]
        can_compare_chunks = all(len(cl) == len(chunk_lists[0]) for cl in chunk_lists) and chunk_lists[0]

        if can_compare_chunks:
            p0 = p_list[0]
            p0.clear()
            for c_idx in range(len(chunk_lists[0])):
                chunk_texts = [cl[c_idx][0] for cl in chunk_lists]
                chunk_props = chunk_lists[0][c_idx][1]

                if len(set(chunk_texts)) == 1:
                    run = p0.add_run(chunk_texts[0])
                    apply_run_props(run, chunk_props)
                    continue

                t_str, var_rs, vals_f = find_all_vars_in_slot(chunk_texts, diff_map, v_idx_list)
                if var_rs:
                    for idx in range(len(var_rs)):
                        v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx_list)
                        t_str = t_str.replace(f"{{{{VAR_PLACEHOLDER_{idx}}}}}", f"{{{{{v_name}}}}}")
                        for f_idx in range(num_files):
                            all_data[f_idx][v_name] = vals_f[f_idx][idx]
                
                run = p0.add_run(t_str)
                apply_run_props(run, chunk_props)
            continue

        all_texts = [p.text for p in p_list]
        if len(set(all_texts)) == 1: continue

        p0 = p_list[0]
        char_to_props = []
        for r in p0.runs:
            props = get_run_props(r)
            for _ in range(len(r.text)):
                char_to_props.append(props)

        t_str, var_rs, vals_f = find_all_vars_in_slot(all_texts, diff_map, v_idx_list, char_to_props=char_to_props)
        if not var_rs: continue

        tok0 = tokenize(all_texts[0])
        tok_char_idx = []
        curr_idx = 0
        for tok in tok0:
            tok_char_idx.append(curr_idx)
            curr_idx += len(tok)
        tok_char_idx.append(curr_idx)

        new_char_props = []
        last_idx = 0

        for idx, (v_start, v_end) in enumerate(var_rs):
            v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx_list)
            for f_idx in range(num_files):
                all_data[f_idx][v_name] = vals_f[f_idx][idx]

            start_char = tok_char_idx[last_idx]
            end_char = tok_char_idx[v_start]
            for i in range(start_char, end_char):
                if i < len(char_to_props):
                    new_char_props.append((all_texts[0][i], char_to_props[i]))

            placeholder = f"{{{{{v_name}}}}}"
            var_start_char = tok_char_idx[v_start]
            
            if var_start_char < len(char_to_props):
                p_props = char_to_props[var_start_char]
            else:
                p_props = char_to_props[-1] if char_to_props else None

            for ch in placeholder:
                new_char_props.append((ch, p_props))

            last_idx = v_end

        start_char = tok_char_idx[last_idx]
        for i in range(start_char, len(all_texts[0])):
            if i < len(char_to_props):
                new_char_props.append((all_texts[0][i], char_to_props[i]))

        new_fragments = []
        if new_char_props:
            cur_text = new_char_props[0][0]
            cur_props = new_char_props[0][1]
            for ch, props in new_char_props[1:]:
                if props == cur_props:
                    cur_text += ch
                else:
                    new_fragments.append((cur_text, cur_props))
                    cur_text = ch
                    cur_props = props
            new_fragments.append((cur_text, cur_props))

        p0.clear()
        for text, props in new_fragments:
            run = p0.add_run(text)
            apply_run_props(run, props)

def compare_word_group(files, template_out):
    docs = [Document(f) for f in files]
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    process_paragraph_list([d.paragraphs for d in docs], diff_map, v_idx, all_data)
    for t_i in range(min(len(d.tables) for d in docs)):
        t_list = [d.tables[t_i] for d in docs]
        for r in range(min(len(t.rows) for t in t_list)):
            c_len = min(len(t.rows[r].cells) for t in t_list)
            for c in range(c_len):
                c_list = [t.rows[r].cells[c] for t in t_list]
                if len(set(cell.text for cell in c_list)) > 1:
                    process_paragraph_list([cell.paragraphs for cell in c_list], diff_map, v_idx, all_data)
    t_p = get_unique_path(template_out)
    docs[0].save(t_p)
    return t_p, all_data, v_idx[0]-1

def compare_excel_group(files, template_out):
    wbs_d = [openpyxl.load_workbook(f, data_only=True) for f in files]
    wb_t = openpyxl.load_workbook(files[0])
    diff_map, v_idx, all_data = {}, [1], [{} for _ in range(len(files))]
    num_sheets = min(len(wb.sheetnames) for wb in wbs_d)
    for i in range(num_sheets):
        titles = [wb.sheetnames[i] for wb in wbs_d]
        if len(set(titles)) > 1:
            t_str, var_rs, vals_f = find_all_vars_in_slot(titles, diff_map, v_idx)
            if var_rs:
                for idx in range(len(var_rs)):
                    v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                    t_str = t_str.replace(f"{{{{VAR_PLACEHOLDER_{idx}}}}}", f"{{{{{v_name}}}}}")
                    for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                wb_t.worksheets[i].title = t_str
    for i in range(num_sheets):
        sheets_d = [wb.worksheets[i] for wb in wbs_d]
        st = wb_t.worksheets[i]
        max_r = max((s.max_row or 1) for s in sheets_d)
        max_c = max((s.max_column or 1) for s in sheets_d)
        for r in range(1, max_r + 1):
            for c in range(1, max_c + 1):
                cell_t = st.cell(row=r, column=c)
                if cell_t.data_type == 'f' or (isinstance(cell_t.value, str) and str(cell_t.value).startswith('=')): continue
                vals = [s.cell(row=r, column=c).value for s in sheets_d]
                if len(set(vals)) > 1:
                    if not all(isinstance(v, str) for v in vals if v is not None):
                        v_name = get_var_name(vals, diff_map, v_idx)
                        cell_t.value = f"{{{{{v_name}}}}}"
                        for f_idx in range(len(files)): all_data[f_idx][v_name] = vals[f_idx]
                    else:
                        t_str, var_rs, vals_f = find_all_vars_in_slot([str(v or "") for v in vals], diff_map, v_idx)
                        if var_rs:
                            for idx in range(len(var_rs)):
                                v_name = get_var_name([v[idx] for v in vals_f], diff_map, v_idx)
                                t_str = t_str.replace(f"{{{{VAR_PLACEHOLDER_{idx}}}}}", f"{{{{{v_name}}}}}")
                                for f_idx in range(len(files)): all_data[f_idx][v_name] = vals_f[f_idx][idx]
                            cell_t.value = t_str
    t_p = get_unique_path(template_out)
    wb_t.save(t_p)
    return t_p, all_data, v_idx[0]-1

def generate_name_template(filenames, all_data):
    base_names = [os.path.splitext(os.path.basename(f))[0] for f in filenames]
    best_pattern = None
    max_placeholders = 0
    for i in range(len(base_names)):
        name = base_names[i]
        row_data = all_data[i]
        vars_sorted = sorted([(k, str(v).strip()) for k, v in row_data.items() if v and len(str(v).strip()) >= 2], key=lambda x: len(x[1]), reverse=True)
        pattern = name
        placeholders_count = 0
        for k, v in vars_sorted:
            if v in pattern:
                pattern = pattern.replace(v, f"{{{{{k}}}}}")
                placeholders_count += 1
        if placeholders_count > max_placeholders:
            max_placeholders = placeholders_count
            best_pattern = pattern
    if best_pattern: return best_pattern
    return f"{base_names[0]}_{{{{YYYY}}}}{{{{MM}}}}{{{{DD}}}}_{{{{hh}}}}{{{{mm}}}}"

def populate_config_sheet(ws, path, template_path, filenames, all_data, num_vars):
    try: rel_t = os.path.relpath(template_path, os.path.dirname(os.path.abspath(path)))
    except: rel_t = template_path
    name_pattern = generate_name_template(filenames, all_data)
    ws['A1'], ws['A2'] = rel_t, name_pattern
    headers = [f"field_{i}" for i in range(1, num_vars + 1)]
    for i, h in enumerate(headers): ws.cell(row=4, column=i+1).value = h
    unique_rows = []
    seen = set()
    for row_dict in all_data:
        row_tuple = tuple(get_norm_key(row_dict.get(h, "")) for h in headers)
        if row_tuple not in seen:
            unique_rows.append(row_dict)
            seen.add(row_tuple)
    for r_idx, row_dict in enumerate(unique_rows):
        for c_idx, h in enumerate(headers):
            val = row_dict.get(h, "")
            cell = ws.cell(row=5 + r_idx, column=c_idx + 1)
            cell.value = val
            if isinstance(val, str) and '\n' in val: cell.alignment = Alignment(wrapText=True)

def save_single_config(path, template_path, all_data, num_vars, filenames):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Settings"
    populate_config_sheet(ws, path, template_path, filenames, all_data, num_vars)
    f_p = get_unique_path(path)
    wb.save(f_p)
    return f_p

def save_master_config(path, results):
    wb = openpyxl.Workbook()
    if wb.active: wb.remove(wb.active)
    for idx, (filenames, template_path, all_data, num_vars) in enumerate(results):
        base_name = os.path.splitext(os.path.basename(filenames[0]))[0]
        sheet_title = re.sub(r'[\\/*?:\[\]]', "", base_name)[:30]
        if not sheet_title: sheet_title = f"Group_{idx+1}"
        ws = wb.create_sheet(title=sheet_title)
        populate_config_sheet(ws, path, template_path, filenames, all_data, num_vars)
    f_p = get_unique_path(path)
    wb.save(f_p)
    return f_p

def create_bat_file(bat_path, config_path):
    s_dir = os.path.dirname(os.path.abspath(__file__))
    m_py, v_py = os.path.join(s_dir, "_templates_machine_.py"), os.path.join(s_dir, ".venv", "Scripts", "python.exe")
    if not os.path.exists(v_py): v_py = "python"
    cnt = f'@echo off\n"{v_py}" "{m_py}" "{os.path.abspath(config_path)}" all all\npause\n'
    with open(bat_path, "w", encoding="cp1251") as f: f.write(cnt)

def is_file_a_template(path, ext):
    try:
        if ext == '.docx':
            doc = Document(path)
            for p in doc.paragraphs:
                if re.search(r'\{\{.*?\}\}', p.text): return True
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        if re.search(r'\{\{.*?\}\}', c.text): return True
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if isinstance(cell, str) and re.search(r'\{\{.*?\}\}', cell): return True
    except: pass
    return False

def get_structure_key(path, ext):
    try:
        if ext == '.docx':
            doc = Document(path)
            return ("word", len(doc.paragraphs), len(doc.tables))
        elif ext == '.xlsx':
            wb = openpyxl.load_workbook(path, read_only=True)
            return ("excel", len(wb.sheetnames))
    except: return None

def sort_files_by_complexity(files, ext):
    def get_complexity(f):
        try:
            if ext == '.docx':
                doc = Document(f)
                return sum(len(p.text) for p in doc.paragraphs) + sum(len(c.text) for t in doc.tables for r in t.rows for c in r.cells)
            elif ext == '.xlsx':
                wb = openpyxl.load_workbook(f, data_only=True)
                return sum(len(str(cell.value or "")) for sheet in wb.worksheets for row in sheet.iter_rows() for cell in row)
        except:
            return 0
    return sorted(files, key=get_complexity, reverse=True)

def run_compare_two(f1, f2):
    ext = os.path.splitext(f1)[1].lower()
    if ext not in ['.docx', '.xlsx'] or os.path.splitext(f2)[1].lower() != ext:
        print("Помилка: Файли повинні мати однакове розширення (.docx або .xlsx).")
        return
    files = sort_files_by_complexity([f1, f2], ext)
    f1_a = os.path.abspath(files[0])
    f_dir, f_name = os.path.dirname(f1_a), os.path.splitext(os.path.basename(f1_a))[0]
    t_o = os.path.join(f_dir, f"template_{f_name}{ext}")
    m_x = os.path.join(f_dir, f"{f_name}_config.xlsx")
    b_o = os.path.join(f_dir, f"{f_name}_run_all.bat")
    print(f"Аналіз 2 файлів (Режим: Порівняння)...")
    try:
        if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
        else: t_p, data, n_v = compare_excel_group(files, t_o)
        if n_v == 0:
            print(f"\nКонфіг та BAT-файл не створено, оскільки змінних не знайдено.")
            print(f"Шаблон (копія): {t_p}")
        else:
            m_p = save_single_config(m_x, t_p, data, n_v, files)
            create_bat_file(b_o, m_p)
            print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")
    except Exception as e: print(f"Помилка: {e}")

def run_package(sample, folder):
    ext = os.path.splitext(sample)[1].lower()
    files = [sample]
    for root, _, fs in os.walk(folder):
        for f in fs:
            if f.lower().endswith(ext):
                if "template" in f.lower(): continue
                fp = os.path.join(root, f)
                if os.path.abspath(fp) == os.path.abspath(sample): continue
                if is_file_a_template(fp, ext): continue
                files.append(fp)
    if len(files) < 2:
        print("Не знайдено подібних файлів для порівняння.")
        return
    files = sort_files_by_complexity(files, ext)
    sample_a = os.path.abspath(files[0])
    f_dir, f_name = os.path.dirname(sample_a), os.path.splitext(os.path.basename(sample_a))[0]
    t_o = os.path.join(f_dir, f"template_{f_name}{ext}")
    m_x = os.path.join(f_dir, f"{f_name}_config.xlsx")
    b_o = os.path.join(f_dir, f"{f_name}_run_all.bat")
    print(f"Аналіз {len(files)} файлів (Режим: Пакетний)...")
    try:
        if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
        else: t_p, data, n_v = compare_excel_group(files, t_o)
        if n_v == 0:
            print(f"\nКонфіг та BAT-файл не створено, оскільки змінних не знайдено.")
            print(f"Шаблон (копія): {t_p}")
        else:
            m_p = save_single_config(m_x, t_p, data, n_v, files)
            create_bat_file(b_o, m_p)
            print(f"\nУспішно!\nШаблон: {t_p}\nКонфіг: {m_p}")
    except Exception as e: print(f"Помилка: {e}")

def run_full_auto(folder):
    print(f"Сканування папки: {folder} (Режим: Повний автопілот)...")
    groups = {}
    for root, _, fs in os.walk(folder):
        for f in fs:
            ext = os.path.splitext(f)[1].lower()
            if ext not in ['.docx', '.xlsx']: continue
            if "template" in f.lower(): continue
            fp = os.path.join(root, f)
            if is_file_a_template(fp, ext): continue
            key = get_structure_key(fp, ext)
            if key:
                if key not in groups: groups[key] = []
                groups[key].append(fp)
    if not groups:
        print("Не знайдено підходящих документів.")
        return
    results = []
    out_dir = os.getcwd()
    for key, files in groups.items():
        if len(files) < 2:
            print(f"Група {key[0]} ({key[1]}) має лише один файл, пропускаємо.")
            continue
        ext = os.path.splitext(files[0])[1].lower()
        files = sort_files_by_complexity(files, ext)
        base_name = os.path.splitext(os.path.basename(files[0]))[0]
        t_o = os.path.join(out_dir, f"template_{base_name}{ext}")
        print(f"Обробка групи: {base_name} ({len(files)} файлів)...")
        try:
            if ext == '.docx': t_p, data, n_v = compare_word_group(files, t_o)
            else: t_p, data, n_v = compare_excel_group(files, t_o)
            print(f"  - Знайдено змінних: {n_v}")
            if n_v > 0:
                results.append((files, t_p, data, n_v))
            else:
                print(f"  - Аркуш пропущено: змінних не знайдено.")
        except Exception as e:
            import traceback
            print(f"  - Помилка: {e}")
            traceback.print_exc()
    if not results:
        print("Не вдалося створити жодного шаблону (або змінні відсутні).")
        return
    m_p = save_master_config(os.path.join(out_dir, "Auto_Config.xlsx"), results)
    create_bat_file(os.path.join(out_dir, "Auto_Run_All.bat"), m_p)
    print(f"\nГотово!\nСтворено конфіг: {m_p}\nОброблено груп: {len(results)}")

# ==========================================
# ЧАСТИНА 2: ГЕНЕРАЦІЯ ДОКУМЕНТІВ
# ==========================================

def render_string_template(template_str, variables):
    result = template_str
    for key, val in variables.items():
        pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
        result = re.sub(pattern, lambda _: str(val), result)
    return result

def get_now_vars():
    now = datetime.now()
    return {
        "YYYY": now.strftime("%Y"), "MM": now.strftime("%m"), "DD": now.strftime("%d"),
        "hh": now.strftime("%H"), "mm": now.strftime("%M"), "ss": now.strftime("%S")
    }

def process_word(template_path, output_path, variables):
    try:
        doc = DocxTemplate(template_path)
        processed_vars = {}
        for key, value in variables.items():
            if isinstance(value, str) and '\n' in value: processed_vars[key] = Listing(value)
            else: processed_vars[key] = value
        doc.render(processed_vars)
        doc.save(output_path)
        print(f"  [Word] Створено: {os.path.basename(output_path)}")
    except Exception as e: print(f"  [Помилка Word]: {e}")

def process_excel(template_path, output_path, variables):
    try:
        wb = openpyxl.load_workbook(template_path)
        str_vars = {k: str(v) for k, v in variables.items()}
        for sheet in wb.worksheets:
            new_title = sheet.title
            title_modified = False
            for key, val in str_vars.items():
                pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
                if re.search(pattern, new_title):
                    new_title = re.sub(pattern, lambda _: val, new_title)
                    title_modified = True
            if title_modified:
                clean_title = re.sub(r'[\\/*?:\[\]]', "", new_title)[:31].strip()
                if not clean_title: clean_title = "Sheet"
                base_title = clean_title
                counter = 1
                while clean_title in wb.sheetnames and clean_title != sheet.title:
                    suffix = f"_{counter}"
                    limit = 31 - len(suffix)
                    clean_title = base_title[:limit] + suffix
                    counter += 1
                sheet.title = clean_title
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        original_value = cell.value
                        exact_key = None
                        for key in variables.keys():
                            pattern = r"^\s*\{\{\s*" + re.escape(key) + r"\s*\}\}\s*$"
                            if re.match(pattern, original_value):
                                exact_key = key
                                break
                        if exact_key is not None:
                            new_value = variables[exact_key]
                            source_type = variables.get(f"{exact_key}_type_code", "s")
                            template_type = cell.data_type
                            target_type = source_type if source_type != "s" else template_type
                            if isinstance(new_value, str):
                                if target_type == "n":
                                    try:
                                        if "." in new_value or "e" in new_value.lower(): new_value = float(new_value)
                                        else: new_value = int(new_value)
                                    except: pass
                                elif target_type == "d":
                                    try: new_value = datetime.fromisoformat(new_value)
                                    except: pass
                                elif target_type == "b":
                                    if new_value.lower() in ["true", "1", "yes"]: new_value = True
                                    elif new_value.lower() in ["false", "0", "no"]: new_value = False
                            cell.value = new_value
                            if isinstance(cell.value, str) and '\n' in cell.value: cell.alignment = Alignment(wrapText=True)
                            continue
                        new_value = original_value
                        modified = False
                        for key, val in str_vars.items():
                            pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
                            if re.search(pattern, new_value):
                                new_value = re.sub(pattern, lambda _: val, new_value)
                                modified = True
                        if modified:
                            cell.value = new_value
                            if isinstance(new_value, str) and '\n' in new_value: cell.alignment = Alignment(wrapText=True)
        wb.save(output_path)
        print(f"  [Excel] Створено: {os.path.basename(output_path)}")
    except Exception as e: print(f"  [Помилка Excel]: {e}")

def run_generation(excel_file, sheet_selector="all", row_selector="all"):
    excel_dir = os.path.dirname(os.path.abspath(excel_file))
    try: wb = openpyxl.load_workbook(excel_file, data_only=True)
    except Exception as e:
        print(f"Помилка при відкритті Excel: {e}")
        return
    now_vars = get_now_vars()
    for k in list(now_vars.keys()): now_vars[f"{k}_type"] = "s"
    TYPE_MAP = {'s': 'string', 'n': 'number', 'd': 'date', 'b': 'boolean', 'f': 'formula', 'e': 'error'}
    sheets_to_process = []
    if sheet_selector.lower() == "all": sheets_to_process = wb.worksheets
    else:
        try:
            idx = int(sheet_selector) - 1
            if 0 <= idx < len(wb.worksheets): sheets_to_process = [wb.worksheets[idx]]
        except ValueError:
            if sheet_selector in wb.sheetnames: sheets_to_process = [wb[sheet_selector]]
    if not sheets_to_process:
        print(f"Аркуш '{sheet_selector}' не знайдено.")
        return

    for sheet in sheets_to_process:
        print(f"Обробка аркуша: {sheet.title}")
        template_rel_path = sheet.cell(row=1, column=1).value
        if not template_rel_path:
            print(f"  Пропуск: Не вказано шаблон у комірці A1.")
            continue
        template_path = resolve_path(excel_dir, str(template_rel_path))
        if not os.path.exists(template_path):
            print(f"  Помилка: Шаблон не знайдено: {template_path}")
            continue
        output_pattern = sheet.cell(row=2, column=1).value
        if not output_pattern:
            print(f"  Пропуск: Не вказано шлях результату в A2.")
            continue
        headers = []
        for col in range(1, sheet.max_column + 1):
            val = sheet.cell(row=4, column=col).value
            if val is not None:
                header_name = str(val).strip()
                if header_name.startswith("{{") and header_name.endswith("}}"): header_name = header_name[2:-2].strip()
                headers.append(header_name)
            else: headers.append(None)
        while headers and headers[-1] is None: headers.pop()
        if not any(h is not None for h in headers):
            print(f"  Пропуск: Не знайдено заголовків у рядку 4.")
            continue
        print(f"  Знайдені змінні: {', '.join([h for h in headers if h is not None])}")
        data_rows = []
        max_row = sheet.max_row
        if row_selector.lower() == "all":
            for r in range(5, max_row + 1):
                row_vals = []
                for c in range(1, len(headers) + 1):
                    cell = sheet.cell(row=r, column=c)
                    row_vals.append((cell.value, cell.data_type))
                if any(v[0] is not None for v in row_vals): data_rows.append((r, row_vals))
        else:
            try:
                r_idx = int(row_selector)
                if 5 <= r_idx <= max_row:
                    row_vals = []
                    for c in range(1, len(headers) + 1):
                        cell = sheet.cell(row=r_idx, column=c)
                        row_vals.append((cell.value, cell.data_type))
                    if any(v[0] is not None for v in row_vals): data_rows.append((r_idx, row_vals))
            except ValueError: print(f"  Невірний формат номера рядка: {row_selector}")

        for r_num, row_vals in data_rows:
            variables = {**now_vars}
            for i, h in enumerate(headers):
                if h is None: continue
                val, v_type = row_vals[i] if i < len(row_vals) else (None, "s")
                if val is None: val = ""
                variables[h] = val
                variables[f"{h}_type"] = TYPE_MAP.get(v_type, v_type)
                variables[f"{h}_type_code"] = v_type
            rendered_out_path = render_string_template(str(output_pattern), variables)
            ext = os.path.splitext(template_path)[1].lower()
            if not rendered_out_path.lower().endswith(ext): rendered_out_path += ext
            rendered_out_path = rendered_out_path.replace('/', os.sep).replace('\\', os.sep)
            final_output_path = resolve_path(excel_dir, rendered_out_path)
            final_output_path = get_unique_path(final_output_path)
            print(f"  Рядок {r_num}:")
            if ext == '.docx': process_word(template_path, final_output_path, variables)
            elif ext == '.xlsx': process_excel(template_path, final_output_path, variables)
            else: print(f"  [Помилка] Непідтримуваний формат шаблону: {ext}")

# ==========================================
# ГОЛОВНИЙ РОУТЕР
# ==========================================

def main():
    args = sys.argv[1:]
    if len(args) == 0:
        print("Використання СИСТЕМИ _templates_machine_:")
        print("  [СТВОРЕННЯ ШАБЛОНІВ ТА КОНФІГІВ]")
        print("  1. Повний автопілот: python _templates_machine_.py <папка>")
        print("  2. Пакетний режим:   python _templates_machine_.py <файл_зразок> <папка>")
        print("  3. Порівняння двох:  python _templates_machine_.py <файл_1> <файл_2>\n")
        print("  [ГЕНЕРАЦІЯ ГОТОВИХ ДОКУМЕНТІВ]")
        print("  4. Запуск генерації: python _templates_machine_.py <файл_конфігу.xlsx> [аркуш] [рядок]")
        return
        
    arg1 = args[0]
    
    if len(args) == 1 and os.path.isdir(arg1):
        run_full_auto(arg1)
        return
        
    if len(args) == 2 and os.path.isfile(arg1) and os.path.isdir(args[1]):
        run_package(arg1, args[1])
        return
        
    if len(args) == 2 and os.path.isfile(arg1) and os.path.isfile(args[1]):
        run_compare_two(arg1, args[1])
        return
        
    if arg1.lower().endswith('.xlsx') and os.path.isfile(arg1):
        sheet = args[1] if len(args) > 1 else "all"
        row = args[2] if len(args) > 2 else "all"
        run_generation(arg1, sheet, row)
        return
        
    print("Помилка: Невідома комбінація параметрів.")
    print("Запустіть скрипт без параметрів для виклику довідки.")

if __name__ == "__main__":
    main()
